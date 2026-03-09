
import streamlit as st
from google import genai
from docx import Document
from docx.shared import Pt
import requests
import io
import time
import re 
import os
import pandas as pd
from docx.enum.text import WD_ALIGN_PARAGRAPH
from collections import defaultdict
from streamlit_quill import st_quill
from docx.shared import RGBColor
from htmldocx import HtmlToDocx
from docx.shared import Inches

try:
    from htmldocx import HtmlToDocx
except ImportError:
    HtmlToDocx = None

def auditar_tablas_maestro(doc_maestro):
    datos_auditoria = []
    
    for i, tabla in enumerate(doc_maestro.tables):
        titulo_final = "No se detectó etiqueta 'Tabla' arriba"
        
        # --- LÓGICA DE BÚSQUEDA HACIA ARRIBA ---
        elemento = tabla._element.getprevious()
        # Buscamos hasta 10 elementos hacia arriba por si hay mucha separación
        for _ in range(10): 
            if elemento is not None:
                if elemento.tag.endswith('p'):
                    from docx.text.paragraph import Paragraph
                    p_temp = Paragraph(elemento, doc_maestro)
                    texto = p_temp.text.strip()

                    # Capturamos el estilo (usualmente 'Heading 6' o 'Título 6')
                    nombre_estilo = p_temp.style.name.lower()
                    if "tabla" in texto.lower() or any(n in nombre_estilo for n in ["5", "6", "7"]):
                        titulo_final = texto if texto else f"Detectado por estilo: {p_temp.style.name}"
                        break # Encontramos el título, dejamos de buscar arriba
                elemento = elemento.getprevious()
        
        # Muestra del contenido para confirmar que es la tabla correcta
        try:
            muestra = [celda.text[:30] for celda in tabla.rows[0].cells[:2]]
        except:
            muestra = "No accesible"
            
        datos_auditoria.append({
            "Índice": i,
            "Título Identificado": titulo_final,
            "Dimensiones": f"{len(tabla.rows)} filas x {len(tabla.columns)} col.",
            "Contenido inicial": muestra
        })
    
    if datos_auditoria:
        st.dataframe(datos_auditoria, use_container_width=True)
    else:
        st.warning("No se encontraron tablas físicas en el documento.")
        

def mapear_todas_las_tablas(archivo_dm):
    from docx import Document
    import io

    # Validamos si es un objeto cargado (Streamlit) o una ruta
    if not isinstance(archivo_dm, str):
        archivo_dm.seek(0)
        doc_maestro = Document(archivo_dm)
    else:
        doc_maestro = Document(archivo_dm)

    mapa_tablas = {}

    for i, tabla in enumerate(doc_maestro.tables):
        titulo_detectado = f"Tabla sin título {i+1}"
        elemento = tabla._element.getprevious()
        
        # Buscamos hasta 10 elementos hacia arriba (más persistente)
        for _ in range(10): 
            if elemento is not None and elemento.tag.endswith('p'):
                from docx.text.paragraph import Paragraph
                p_temp = Paragraph(elemento, doc_maestro)
                texto = p_temp.text.strip()
                nombre_estilo = p_temp.style.name.lower()
                
                # Detectamos por palabra "Tabla" o por Estilo Título 5, 6 o 7
                if "tabla" in texto.lower() or any(n in nombre_estilo for n in ["5", "6", "7"]):
                    
                    # Lógica para Títulos en 2 líneas: "Tabla xx." (arriba) + "Título" (abajo)
                    # Si el texto es muy corto (ej: "Tabla 1.") buscamos el párrafo de abajo
                    if len(texto) < 15 and "tabla" in texto.lower():
                        elemento_siguiente = elemento.getnext()
                        if elemento_siguiente is not None and elemento_siguiente.tag.endswith('p'):
                            from docx.text.paragraph import Paragraph
                            p_sig = Paragraph(elemento_siguiente, doc_maestro)
                            texto_sig = p_sig.text.strip()
                            if texto_sig:
                                titulo_detectado = f"{texto} {texto_sig}"
                            else:
                                titulo_detectado = texto
                        else:
                            titulo_detectado = texto
                    else:
                        # Si ya viene todo en una línea: "Tabla 1. Plan de formación"
                        titulo_detectado = texto if texto else f"Tabla detectada por estilo: {p_temp.style.name}"
                    
                    break
            
            if elemento is not None:
                elemento = elemento.getprevious()
        
        mapa_tablas[titulo_detectado] = tabla
        
    return mapa_tablas


def insertar_tabla_automatica(doc_destino, placeholder, keyword_titulo):
    from docx.shared import Pt
    import copy
    import unicodedata

    # Función interna para quitar tildes y normalizar
    def normalizar(texto):
        if not texto: return ""
        return "".join(
            c for c in unicodedata.normalize('NFD', texto)
            if unicodedata.category(c) != 'Mn'
        ).lower().strip()

    mapa = st.session_state.get("mapa_tablas", {})
    
    # Preparamos las palabras clave (separadas por espacios)
    palabras_busqueda = normalizar(keyword_titulo).split()

# --- REEMPLAZO DE LÓGICA DE BÚSQUEDA ---
    tabla_fuente = None
    # Separamos por el carácter "|" para permitir búsquedas alternativas
    opciones_busqueda = keyword_titulo.split("|") 
    
    for titulo_maestro, tabla in mapa.items():
        titulo_maestro_norm = normalizar(titulo_maestro)
        
        # Probamos cada una de las opciones del "ó"
        for opcion in opciones_busqueda:
            palabras_clave = normalizar(opcion).split()
            
            # Si todas las palabras de esta opción están en el título, ¡bingo!
            if all(p in titulo_maestro_norm for p in palabras_clave):
                tabla_fuente = tabla
                break
        
        if tabla_fuente: 
            break
    
    if not tabla_fuente:
        return False
    
    #tabla_fuente = None
    # Buscamos en el mapa de tablas
    #for titulo_maestro, tabla in mapa.items():
     #   titulo_maestro_norm = normalizar(titulo_maestro)
      #  # Verificamos que TODAS las palabras clave estén en el título del maestro
       # if all(p in titulo_maestro_norm for p in palabras_busqueda):
        #    tabla_fuente = tabla
         #   break
    
    #if not tabla_fuente:
     #   return False

    # --- PROCESO DE INSERCIÓN ---
    for paragraph in doc_destino.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, "")
            
            new_tbl = doc_destino.add_table(rows=0, cols=len(tabla_fuente.columns))
            try: new_tbl.style = 'Table Grid'
            except: pass
            
            for row in tabla_fuente.rows:
                # Condición de parada (Fuente)
                contenido_fila = " ".join([cell.text for cell in row.cells])
                if "fuente" in contenido_fila.lower():
                    break
                
                new_row = new_tbl.add_row()
                for j, cell in enumerate(row.cells):
                    new_cell = new_row.cells[j]
                    p = new_cell.paragraphs[0]
                    p.clear() # Evita el error de la celda vacía
                    run = p.add_run(cell.text)
                    run.font.size = Pt(10)
                    
                    # Copiar sombreado/color
                    shd_elements = cell._tc.xpath('.//w:shd')
                    if shd_elements:
                        shd_copy = copy.deepcopy(shd_elements[0])
                        tcPr = new_cell._tc.get_or_add_tcPr()
                        tcPr.append(shd_copy)

            paragraph._p.addnext(new_tbl._element)
            return True
    return False


def insertar_tabla_seleccionada(doc_destino, placeholder, titulo_seleccionado):
    from docx.shared import Pt
    import copy
    
    mapa = st.session_state.get("mapa_tablas", {})
    tabla_fuente = mapa.get(titulo_seleccionado)
    
    if not tabla_fuente:
        return False

    for paragraph in doc_destino.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, "")
            
            # 1. Crear la tabla
            new_tbl = doc_destino.add_table(rows=0, cols=len(tabla_fuente.columns))
            try: new_tbl.style = 'Table Grid'
            except: pass
            
            # 2. Copiar filas
            for row in tabla_fuente.rows:
                contenido_fila = " ".join([cell.text for cell in row.cells])
                if "fuente" in contenido_fila.lower():
                    break
                
                new_row = new_tbl.add_row()
                
                for j, cell in enumerate(row.cells):
                    new_cell = new_row.cells[j]
                    
                    # --- MANEJO SEGURO DE TEXTO Y TAMAÑO ---
                    # En lugar de cell.text = "", usamos el párrafo existente
                    p = new_cell.paragraphs[0]
                    p.clear() # Esto borra el contenido sin eliminar el párrafo
                    run = p.add_run(cell.text)
                    run.font.size = Pt(10)
                    
                    # --- COPIAR COLOR DE FONDO ---
                    shd_elements = cell._tc.xpath('.//w:shd')
                    if shd_elements:
                        shd_copy = copy.deepcopy(shd_elements[0])
                        tcPr = new_cell._tc.get_or_add_tcPr()
                        tcPr.append(shd_copy)

            # 3. Mover la tabla
            paragraph._p.addnext(new_tbl._element)
            return True
    return False



def insertar_tabla_desde_maestro(doc_destino, doc_maestro, placeholder, titulo_tabla):
    """
    Busca una tabla en el documento maestro por su título y la inserta
    en el lugar del placeholder en el documento de destino.
    """
    tabla_encontrada = None
    
    # 1. Buscar la tabla en el maestro
    for i, tbl in enumerate(doc_maestro.tables):
        # Buscamos si el párrafo anterior a la tabla contiene el título
        # O si la primera celda contiene parte del nombre
        if titulo_tabla in doc_maestro.paragraphs[i].text or titulo_tabla in tbl.cell(0,0).text:
            tabla_encontrada = tbl
            break
    
    if not tabla_encontrada:
        return False

    # 2. Buscar el placeholder en el destino para insertar ahí
    for paragraph in doc_destino.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, "")
            # Insertar la tabla después de este párrafo
            new_tbl = doc_destino.add_table(rows=len(tabla_encontrada.rows), 
                                            cols=len(tabla_encontrada.columns))
            new_tbl.style = tabla_encontrada.style # Intentar mantener el estilo
            
            # Copiar contenido celda por celda
            for r in range(len(tabla_encontrada.rows)):
                for c in range(len(tabla_encontrada.columns)):
                    new_tbl.cell(r, c).text = tabla_encontrada.cell(r, c).text
            
            # Añadir la fuente debajo
            doc_destino.add_paragraph("Fuente: Elaboración propia")
            return True
    return False

def insertar_imagen_en_placeholder(doc_destino, placeholder, archivo_imagen, ancho_pulgadas=6.0):
    """
    Busca un placeholder en el documento y lo reemplaza por una imagen cargada.
    """
    from docx.shared import Inches
    import os

    for paragraph in doc_destino.paragraphs:
        if placeholder in paragraph.text:
            # Limpiar el texto del placeholder
            paragraph.text = paragraph.text.replace(placeholder, "")
            run = paragraph.add_run()
            
            # Nombre temporal único para evitar conflictos
            temp_name = f"temp_{placeholder.replace('{','').replace('}','')}.png"
            
            try:
                # Escribir el buffer de Streamlit a un archivo real
                with open(temp_name, "wb") as f:
                    f.write(archivo_imagen.getbuffer())
                
                # Insertar en el documento
                run.add_picture(temp_name, width=Inches(ancho_pulgadas))
                
                # Borrar rastro temporal
                if os.path.exists(temp_name):
                    os.remove(temp_name)
                return True
            except Exception as e:
                st.error(f"Error técnico al insertar imagen: {e}")
                return False
    return False
    

# Función para Insertar DEBAJO de un párrafo específico
def insertar_lista_bajo_titulo(documento, texto_titulo, lista_items):
    """
    Busca el párrafo que contenga 'texto_titulo'.
    Si lo encuentra, inserta los items de la lista justo debajo.
    """
    encontrado = False  
    for i, paragraph in enumerate(documento.paragraphs):
                # Buscamos el título (ignorando mayúsculas/minúsculas para asegurar)
            if texto_titulo.lower() in paragraph.text.lower():
                    
                    # Truco técnico: Para insertar "despues", nos paramos en el párrafo SIGUIENTE                   
                    # Verificamos si hay un párrafo siguiente
                 if i + 1 < len(documento.paragraphs):
                    p_siguiente = documento.paragraphs[i + 1]
                                                                       
                        #Insertamos antes del siguiente párrafo
                    for item in lista_datos:
                        p_siguiente.insert_paragraph_before(item)
                    encontrado = True
                    break
                        
    if not encontrado:
            doc.add_heading("1.2. Generalidades del programa", level=2)
            for item in lista_datos:
                doc.add_paragraph(item)

def reemplazar_etiqueta_por_imagen(doc, etiqueta, imagen_st, ancho_pulgadas=6.0):
    """
    Busca una etiqueta en el doc y la reemplaza por una imagen cargada desde Streamlit.
    """
    if imagen_st is None:
        return
        
    for paragraph in doc.paragraphs:
        if etiqueta in paragraph.text:
            # Limpiar el texto del párrafo (quitar la etiqueta)
            paragraph.text = paragraph.text.replace(etiqueta, "")
            run = paragraph.add_run()
            # Insertar la imagen (imagen_st es el archivo subido por file_uploader)
            run.add_picture(imagen_st, width=Inches(ancho_pulgadas))

# SECCIÓN: HELPERS

def nested_dict():
    """Crea un diccionario infinito para guardar la estructura del Word."""
    return defaultdict(nested_dict)

def is_noise(title):
    """Detecta si un Heading es ruido (tablas, figuras, etc.)."""
    title = title.strip().lower()
    if not title:
        return True
    # Filtramos leyendas comunes que Word a veces confunde con títulos
    ruido = ["tabla", "figura", "imagen", "ilustración", "gráfico", "anexo"]
    return any(title.startswith(r) for r in ruido)

def clean_dict(d):
    """Convierte defaultdict a dict normal y elimina secciones vacías."""
    if not isinstance(d, dict):
        return d
    cleaned = {}
    for k, v in d.items():
        if k == "_content":
            if v.strip():
                cleaned[k] = v.strip()
            continue
        child = clean_dict(v)
        if child:
            cleaned[k] = child
    return cleaned

def docx_to_clean_dict(path):
    """Analiza el Documento Maestro y crea un mapa jerárquico por estilos."""
    doc = Document(path)
    estructura = nested_dict()
    stack = []

    for para in doc.paragraphs:
        text = para.text.strip()
        style = para.style.name

        # Buscamos estilos que empiecen por 'Heading' o 'Título'
        if "Heading" in style or "Título" in style:
            if is_noise(text):
                if stack:
                    current = estructura
                    for item in stack: current = current[item]
                    current["_content"] += text + "\n"
                continue

            try:
                # Intentamos extraer el nivel (ej: 'Heading 2' -> 2)
                level = int(''.join(filter(str.isdigit, style)))
            except:
                level = 1 # Por defecto si no tiene número

            stack = stack[:level-1]
            stack.append(text)

            current = estructura
            for item in stack:
                current = current[item]
            if "_content" not in current: current["_content"] = ""
            if "_nodes" not in current: current["_nodes"] = []

        else:
            # Es un párrafo normal: se añade al contenido de la sección actual
            if stack and text:
                current = estructura
                for item in stack:
                    current = current[item]
                if "_content" not in current:
                    current["_content"] = ""
                current["_content"] += text + "\n"

    return clean_dict(estructura)

    #Fundamentación epistemológica
# Fundamentación epistemológica
def extraer_fundamentacion(diccionario):
    # Claves de inicio optimizadas
    claves = ["onceptualiza", "teor", "epistemol"]
    # Clave de parada específica (para evitar que 3.5 detenga a 3.4)
    freno = "3.5. mecanismos de evaluación"
    
    texto_completo = ""
    seccion_encontrada = False

    def obtener_texto_profundo(nodo):
        texto = ""
        if isinstance(nodo, dict):
            if "_content" in nodo:
                texto += str(nodo["_content"]) + "\n"
            for k, v in nodo.items():
                if k != "_content":
                    texto += f"\n{k}\n" + obtener_texto_profundo(v)
        elif isinstance(nodo, str):
            texto += nodo + "\n"
        return texto

    # EL BUCLE FOR DEBE ESTAR ALINEADO A LA IZQUIERDA
    for titulo_real, contenido in diccionario.items():
        titulo_min = titulo_real.lower().strip()
        
        # 1. LÓGICA DE PARADA (Prioridad: Si vemos el freno, salimos)
        if seccion_encontrada and freno in titulo_min:
            break

        # 2. LÓGICA DE INICIO (Detectar si es el título 3.4)
        if not seccion_encontrada:
            coincidencias = sum(1 for c in claves if c in titulo_min)
            # Si tiene al menos 2 de las palabras clave (ej: Conceptualización + Epistemológica)
            if coincidencias >= 2:
                seccion_encontrada = True
                texto_completo += f"{titulo_real}\n"
                texto_completo += obtener_texto_profundo(contenido)
                continue

        # 3. LÓGICA DE CAPTURA (Solo si ya encontramos el inicio y NO es el freno)
        if seccion_encontrada:
            texto_completo += f"\n{titulo_real}\n"
            texto_completo += obtener_texto_profundo(contenido)

    return texto_completo


def extraer_area_especifica(diccionario):  
    # Buscamos por áreas de formación o fundamentación específica
    claves = ["fundamentac", "espec"]
    excluir = ["basica", "epistemol"]
    
    def obtener_texto_profundo(nodo):
        texto = ""
        if isinstance(nodo, dict):
            contenido_nodo = nodo.get("_content", "")
            
            # LÓGICA DE PARADA (Tabla/Figura)
            contenido_min = contenido_nodo.lower()
            if "tabla" in contenido_min or "figura" in contenido_min:
                # Cortamos en el primer indicio de tabla o figura
                puntos = [i for i in [contenido_min.find("tabla"), contenido_min.find("figura")] if i != -1]
                texto += contenido_nodo[:min(puntos)]
                return texto, True 
            
            texto += contenido_nodo + "\n"
            
            for k, v in nodo.items():
                if k != "_content":
                    if "tabla" in k.lower() or "figura" in k.lower():
                        return texto, True
                    sub_texto, bandera = obtener_texto_profundo(v)
                    texto += f"\n{k}\n" + sub_texto
                    if bandera: return texto, True
        return texto, False
 
    for titulo_real, contenido in diccionario.items():
        titulo_min = titulo_real.lower()
        
        # FILTRO CRÍTICO: Debe tener las claves Y NO tener las palabras de exclusión
        if all(c in titulo_min for c in ["fundament", "espec"]):
            if not any(e in titulo_min for e in excluir):
                texto_final, _ = obtener_texto_profundo(contenido)
                return texto_final
        
        # Búsqueda recursiva
        if isinstance(contenido, dict):
            res = extraer_area_especifica(contenido)
            if res: return res
    return ""
               
def extraer_justificacion_programa(diccionario):
    import re
    nodos_totales = []
    seccion_encontrada = False
    
    # Esta sub-función ahora es más "agresiva" buscando texto
    def obtener_nodos_recursivos(nodo):
        nodos_locales = []
        if isinstance(nodo, dict):
            # 1. Sacamos los nodos de este nivel
            objetos = nodo.get("_nodes", [])
            for obj in objetos:
                # IMPORTANTE: Solo agregamos si tiene texto real (ignora el 'Enter')
                if obj.text.strip():
                    nodos_locales.append(obj)
            
            # 2. Entramos a las subsecciones (2.1, 2.2...) sin importar los huecos
            for k, v in nodo.items():
                if k not in ["_content", "_nodes", "_tables"]:
                    nodos_locales.extend(obtener_nodos_recursivos(v))
        return nodos_locales

    for titulo_real, contenido in diccionario.items():
        titulo_limpio = titulo_real.strip()
        
        # Detectamos el inicio (Capítulo 2)
        if re.match(r'^2(\.|\s|$)', titulo_limpio):
            seccion_encontrada = True
            nodos_totales.extend(obtener_nodos_recursivos(contenido))
            continue

        if seccion_encontrada:
            # Si llegamos al 3, paramos
            if re.match(r'^3(\.|\s|$)', titulo_limpio):
                break
            
            # Seguimos acumulando todo lo que encontremos en el medio
            nodos_totales.extend(obtener_nodos_recursivos(contenido))

    # Filtro final de seguridad: quitar cualquier cosa que no tenga texto
    return [n for n in nodos_totales if n.text.strip()]
    


def extraer_resultados_aprendizaje(diccionario):  
    claves = ["resultados", "aprendizaje", "rapa"]
    palabras_omision = ["tabla", "figura", "fuente:"]
    
    def obtener_texto_profundo(nodo, estado_salto={"omitido": False}):
        texto = ""
        if isinstance(nodo, dict):
            contenido_nodo = nodo.get("_content", "")
            
            # Dividimos el contenido en párrafos por saltos de línea
            lineas = [l.strip() for l in contenido_nodo.split('\n') if l.strip()]
            
            for linea in lineas:
                # 1. Verificamos si es una tabla o figura (estas se ignoran siempre)
                if any(p in linea.lower() for p in palabras_omision):
                    continue
                
                # 2. Lógica para ignorar el PRIMER párrafo válido encontrado
                if not estado_salto["omitido"]:
                    estado_salto["omitido"] = True
                    continue # Saltamos este párrafo y marcamos como omitido
                
                # 3. Si ya omitimos el primero, acumulamos el resto
                texto += linea + "\n\n"
            
            for k, v in nodo.items():
                if k != "_content" and k != "_tables":
                    if any(p in k.lower() for p in palabras_omision):
                        continue 
                    # Pasamos el estado de salto a las subsecciones
                    texto += obtener_texto_profundo(v, estado_salto)
        return texto

    for titulo_real, contenido in diccionario.items():
        titulo_min = titulo_real.lower()
        if all(c in titulo_min for c in ["resultados", "aprendizaje"]):
            # Iniciamos la extracción
            return obtener_texto_profundo(contenido, {"omitido": False})
        
        if isinstance(contenido, dict):
            res = extraer_resultados_aprendizaje(contenido)
            if res: return res
    return ""

def extraer_perfil_generico(diccionario, claves_busqueda):
    """
    Función versátil para extraer perfiles omitiendo tablas/figuras.
    """
    palabras_omision = ["tabla", "figura", "fuente:"]
    
    def obtener_texto_profundo(nodo):
        texto = ""
        if isinstance(nodo, dict):
            contenido_nodo = nodo.get("_content", "")
            # Omisión de líneas con Tablas o Figuras
            lineas = contenido_nodo.split('\n')
            for linea in lineas:
                if not any(p in linea.lower() for p in palabras_omision):
                    texto += linea + "\n"
            
            for k, v in nodo.items():
                if k != "_content" and k != "_tables":
                    # Si el subtítulo es una Tabla/Figura, lo saltamos
                    if any(p in k.lower() for p in palabras_omision):
                        continue
                    texto += f"\n{k}\n" + obtener_texto_profundo(v)
        return texto

    for titulo_real, contenido in diccionario.items():
        titulo_min = titulo_real.lower()
        # Verificamos que todas las claves estén en el título (ej: 'perfil', 'ocupacional')
        if all(c.lower() in titulo_min for c in claves_busqueda):
            return obtener_texto_profundo(contenido)
        
        if isinstance(contenido, dict):
            res = extraer_perfil_generico(contenido, claves_busqueda)
            if res: return res
    return ""

def obtener_solo_estructura(d):
    """
    Crea una copia del diccionario que contiene solo los títulos, 
    eliminando las claves '_content'.
    """
    if not isinstance(d, dict):
        return d
    # Filtramos para dejar solo las llaves que no son '_content'
    return {k: obtener_solo_estructura(v) for k, v in d.items() if k != "_content"}                

#def reemplazar_en_todo_el_doc(doc, diccionario_reemplazos):
   # """
  #  Busca y reemplaza texto en párrafos y tablas, aplicando color naranja.
    #"""
    # 1. Buscar en párrafos normales
   # for paragraph in doc.paragraphs:
    #    for key, value in diccionario_reemplazos.items():
     #       if key in paragraph.text:
                # Realizamos el reemplazo de texto plano
      #          paragraph.text = paragraph.text.replace(key, str(value))
                
                # Aplicamos el color naranja oscuro a los fragmentos (runs)
       #         for run in paragraph.runs:
        #            run.font.color.rgb = RGBColor(255, 140, 0)

def reemplazar_en_todo_el_doc(doc, diccionario_reemplazos):
    """
    Busca y reemplaza texto en párrafos y tablas, aplicando color naranja institucional.
    """
    from docx.shared import RGBColor
    
    # Color Naranja (RGB: 255, 140, 0 o el institucional 227, 108, 9)
    naranja = RGBColor(227, 108, 9)

    # 1. Función interna para procesar párrafos (evita repetir código)
    def procesar_parrafo(p):
        for key, value in diccionario_reemplazos.items():
            if key in p.text:
                # Reemplazamos el texto
                p.text = p.text.replace(key, str(value))
                # Aplicamos color a cada fragmento del párrafo
                for run in p.runs:
                    run.font.color.rgb = naranja

    # 2. Buscar en párrafos normales del documento
    for paragraph in doc.paragraphs:
        procesar_parrafo(paragraph)

    # 3. Buscar en todas las TABLAS del documento (Crucial para Perfiles/Justificación)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    procesar_parrafo(paragraph)
    
    # 2. Buscar dentro de Tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in diccionario_reemplazos.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, str(value))
                            for run in paragraph.runs:
                                run.font.color.rgb = RGBColor(255, 140, 0)
    return ""
    
def insertar_tabla_desde_maestro(doc_destino, doc_maestro, placeholder, patron_busqueda):
    tabla_encontrada = None
    indice_tabla = -1
    
    # 1. Localizar la tabla por su título (Regex)
    for i, tbl in enumerate(doc_maestro.tables):
        texto_previo = ""
        # Buscamos en los párrafos cercanos a la tabla
        for p_idx in range(max(0, i-2), min(len(doc_maestro.paragraphs), i+2)):
            texto_previo += doc_maestro.paragraphs[p_idx].text

        if re.search(patron_busqueda, texto_previo, re.IGNORECASE):
            tabla_encontrada = tbl
            indice_tabla = i
            break
    
    if not tabla_encontrada:
        return False

    # 2. Insertar en el destino
    for paragraph in doc_destino.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, "")
            
            # Creamos la tabla en el PEP
            new_tbl = doc_destino.add_table(rows=0, cols=len(tabla_encontrada.columns))
            new_tbl.style = 'Table Grid'
            
            # 3. COPIAR FILAS CON CONDICIÓN DE PARADA
            for row in tabla_encontrada.rows:
                # Revisamos si en alguna celda de esta fila dice "Fuente"
                contenido_fila = " ".join([cell.text for cell in row.cells])
                
                if "Fuente" in contenido_fila:
                    break # Detenemos la copia de esta tabla inmediatamente
                
                # Si no dice Fuente, añadimos la fila al destino
                new_row = new_tbl.add_row()
                for idx, cell in enumerate(row.cells):
                    new_row.cells[idx].text = cell.text
            
            # Opcional: Agregar un párrafo vacío después para separar
            doc_destino.add_paragraph("")
            return True
    return False


def limpiar_completamente(texto):
    if not texto:
        return ""
    import re
    t = texto.replace("</p>", "\n").replace("<br>", "\n")
    return re.sub(r'<[^>]+>', '', t).strip()                   

# 1. FUNCIONES (El cerebro)
# 1.1 Leer DM
def extraer_secciones_dm(archivo_word, mapa_claves):
    """archivo_word: El archivo subido por st.file_uploader. mapa_claves: Un diccionario que dice {'TITULO EN WORD': 'key_de_streamlit'}"""
    doc = Document(archivo_word)
    resultados = {}

# 1. Extraer todos los párrafos del documento
    todos_los_parrafos = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    # --- BUSCAR EL PUNTO DE PARTIDA ---
    indice_inicio_real = 0
    punto_partida = "BREVE RESEÑA HISTÓRICA DEL PROGRAMA"
    
    for i, texto in enumerate(todos_los_parrafos):
        if punto_partida in texto.upper():
            indice_inicio_real = i
            break # Encontramos el inicio real, dejamos de buscar
            
    # Creamos una nueva lista que solo contiene lo que hay desde la Reseña en adelante
    parrafos_validos = todos_los_parrafos[indice_inicio_real:]
    
    # --- PROCESO DE EXTRACCIÓN SOBRE LOS PÁRRAFOS VÁLIDOS ---
    for titulo_buscado, key_st in mapa_claves.items():
        contenido_seccion = []
        for i, texto in enumerate(parrafos_validos):
            texto_upper = texto.upper()
            
            # Buscamos el título (asegurándonos de que no sea una línea gigante)
            if titulo_buscado.upper() in texto_upper and len(texto) < 120:
                
                for j in range(i + 1, len(parrafos_validos)):
                    siguiente_p = parrafos_validos[j]
                    sig_upper = siguiente_p.upper()
                    
                   # Parar SOLO si encontramos un título principal (Ej: 3. o 4.)
                    # Bajamos el límite a 60 caracteres para no confundir párrafos con títulos
                    es_nuevo_capitulo = re.match(r'^\d+[\.\s]', siguiente_p.strip())
                    es_otro_titulo_mapa = any(t.upper() == sig_upper for t in mapa_claves.keys())

                    if (es_nuevo_capitulo or es_otro_titulo_mapa) and len(siguiente_p) < 60:
                        break
                        
                    contenido_seccion.append(siguiente_p)
                
                # 1. Guardamos TODO el texto en una variable "secreta" para el Word final
                texto_completo = "\n\n".join(contenido_seccion).strip()
                st.session_state[f"full_{key_st}"] = texto_completo
                
                # 2. Preparamos la VISTA PREVIA para el cuadro de texto
                parrafos_lista = texto_completo.split("\n\n")
                if len(parrafos_lista) > 2:
                    # Mostramos primer párrafo + aviso + último párrafo
                    resumen = f"{parrafos_lista[0]}\n\n[... {len(parrafos_lista)-2} PÁRRAFOS INTERMEDIOS CARGADOS TOTALMENTE ...]\n\n{parrafos_lista[-1]}"
                    resultados[key_st] = resumen
                else:
                    resultados[key_st] = texto_completo
                
                break

    #  PARTE 2: BUSCAR EN TABLAS
    for tabla in doc.tables:
        for fila in tabla.rows:
            # Verificamos que la fila tenga al menos 2 celdas
            if len(fila.cells) >= 2:
                texto_izq = fila.cells[0].text.strip().upper()
                texto_der = fila.cells[1].text.strip()
                
                # Comparamos la celda izquierda con nuestras palabras clave
                for titulo_buscado, key_st in mapa_claves.items():
                    if titulo_buscado.upper() in texto_izq:
                    # SIMPLIFICACIÓN: Guardamos el texto crudo del Word.
                    # La lógica de conversión la haremos en el widget (selectbox)
                        resultados[key_st] = texto_der

    return resultados

#1.2 Cargar BD
@st.cache_data # Esto hace que el Excel se lea una sola vez y no cada que muevas un botón
#def cargar_base_datos():
 #   try:
        # Puedes usar pd.read_csv("programas.csv") si prefieres CSV
     #   df = pd.read_excel("Programas.xlsx", dtype={'snies_input': str}) 
        # Convertimos el DataFrame en un diccionario donde la llave es el SNIES
  #      return df.set_index("snies_input").to_dict('index')
   # except Exception as e:
    #    st.warning(f"No se pudo cargar la base de datos de Excel: {e}")
     #   return {}


def cargar_base_datos():
    try:
        # 1. Leemos el archivo asegurando que SNIES sea tratado como texto
        df = pd.read_excel("Programas.xlsx", dtype={'snies_input': str}) 
        
        # 2. LIMPIEZA CRÍTICA: Quitamos espacios y posibles decimales .0
        df['snies_input'] = df['snies_input'].astype(str).str.strip().str.replace(r'\.0$', '', regex=True)
        df['facultad_limpia'] = df.iloc[:, 13].astype(str).str.strip()
        df['depto_limpio'] = df.iloc[:, 14].astype(str).str.strip()
        
        # 3. Eliminamos filas donde el SNIES esté vacío por error
        df = df.dropna(subset=['snies_input'])
        
        # Convertimos en diccionario
        return df.set_index("snies_input").to_dict('index')
    except Exception as e:
        st.warning(f"No se pudo cargar la base de datos de Excel: {e}")
        return {}

#1.3 Carga de datos inicial
BD_PROGRAMAS = cargar_base_datos()

#2. MAPEO Y ESTRUCTURA (DICCIONARIO)
# Mapeo de: "Título exacto en el DM" -> "Key en App Streamlit"
MAPA_EXTRACCION = {
    "OBJETO DE CONOCIMIENTO": "obj_nombre_input",
    "JUSTIFICACIÓN": "justificacion_input",
    "Conceptualización teórica y epistemológica del programa": "input_epi_p1",
    "Mecanismos de evaluación": "input_mec_p1",
    "IDENTIDAD DISCIPLINAR": "input_epi_p2",
    "ITINERARIO FORMATIVO": "input_itinerario",
    "Justificación del Programa": "input_just",
    "JUSTIFICACIÓN DEL PROGRAMA": "input_just"
    

}

#3. DICCIONARIO / ESTRUCTURA
# Agregamos 'key_dm' para que el extractor sepa qué título buscar en el Word
estructura_pep = {
    "1. Información del Programa": {
        "1.1. Historia del Programa": {"tipo": "especial_historia"},
        "1.2. Generalidades del Programa": {"tipo": "directo"}
    },
    "2. Referentes Conceptuales": {
        "2.1. Naturaleza del Programa": {
            "tipo": "directo",
            "key_dm": "OBJETO DE CONOCIMIENTO", # Palabra clave para buscar en el DM
            "campos": [
                {
                    "label": "Objeto de conocimiento del Programa", 
                    "req": True, 
                    "key": "obj_nombre_input",
                    "help": "¿Qué conoce, investiga y transforma este programa?"
                }
            ]
        },
        "2.2. Fundamentación epistemológica": {
            "tipo": "directo",
            "key_dm": "Conceptualización teórica y epistemológica del programa",
            "campos": [
                {"label": "Naturaleza epistemológica e identidad académica", "req": True, "key": "input_epi_p1"},
                {"label": "Campo del saber y relación con ciencia/tecnología", "req": True, "key": "input_epi_p2"}
            ]
        },
        "2.3. Fundamentación académica": {
            "tipo": "especial_pascual", 
            "campos": [] 
        }
    }
}


st.markdown("---")

#  CONFIGURACIÓN DE PÁGINA 
st.set_page_config(page_title="Generador Proyecto Educativo", layout="wide")
try:
    st.image("logopascual.png", width=450)
except Exception:
    pass
st.title("Generador Proyecto Educativo del Programa")
st.markdown("""
Esta herramienta permite generar el PEP de dos formas:
1. **Manual:** Completa los campos en las secciones de abajo.
2. **Semiautomatizada:** Sube el Documento Maestro (DM) y el sistema pre-llenará algunos campos.
""")

   
# SELECTOR DE MODALIDAD
# Usamos un radio button estilizado para elegir el método
metodo_trabajo = st.radio(
    "Selecciona cómo deseas trabajar hoy:",
    ["Manual (Desde cero)", "Semiautomatizado (Cargar Documento Maestro)"],
    horizontal=True,
    help="La opción semiautomatizada intentará pre-llenar los campos usando un archivo Word."
)

# Botón DM
if metodo_trabajo == "Semiautomatizado (Cargar Documento Maestro)":
    st.subheader("2. Carga de Documento Maestro")
    archivo_dm = st.file_uploader("Sube el archivo .docx del Documento Maestro", type=["docx"])
    
    if archivo_dm:
        # --- PERSISTENCIA DEL DICCIONARIO ---
        if "dict_maestro" not in st.session_state:
            with st.spinner("Escaneando Documento Maestro..."):
                st.session_state["dict_maestro"] = docx_to_clean_dict(archivo_dm)
        
        dict_m = st.session_state["dict_maestro"]

        if "mapa_tablas" not in st.session_state:
                st.session_state["mapa_tablas"] = mapear_todas_las_tablas(archivo_dm)

        # Ejecutamos las funciones que buscan en el diccionario recién creado
        texto_fund = extraer_fundamentacion(st.session_state["dict_maestro"])
        nodos_just = extraer_justificacion_programa(st.session_state["dict_maestro"])

        #  EL EXPANDER DE AUDITORÍA 
        with st.expander("🔍 Auditoría de Títulos (Jerarquía Detectada)"):
            if not dict_m:
                st.error("No se detectaron estilos de Título en el Word.")
            else:
                # 1. Mostrar Estructura
                estructura_limpia = obtener_solo_estructura(dict_m)
                st.write("Jerarquía detectada:")
                st.json(estructura_limpia)
                st.divider()

                # --- VALIDACIÓN INTELIGENTE (Basada en el resultado de la extracción) ---
                # En lugar de comparar el nombre del título, preguntamos si la función trajo algo
                if texto_fund and len(texto_fund) > 50: # Verificamos que traiga contenido real
                    st.success(f"✅ Conceptualización detectada y extraída ({len(texto_fund)} caracteres).")
                    # Guardamos para el Word
                    st.session_state["fund_epi_manual"] = texto_fund
                else:
                    # Si el título aparece en el JSON de arriba pero aquí sale error, 
                    # es porque el 'freno' (mecanismos) está actuando antes de tiempo.
                    st.error("❌ No se encontró contenido para 'Conceptualización'.")
                    st.info("Nota: Si ves el título en el JSON de arriba, revisa que el título de la siguiente sección no contenga la palabra 'mecanismos' antes de tiempo.")
                
                #RESULTADOS DE CONCEPTUALIZACIÓN
               # if texto_fund:
                #    st.success(f"✅ Conceptualización: {len(texto_fund)} caracteres detectados.")
                 #   st.session_state["fund_epi_manual"] = texto_fund
                #else:
                 #   st.error("❌ No se encontró 'Conceptualización teórica y epistemológica'.")

                # RESULTADOS DE JUSTIFICACIÓN
                if nodos_just:
                    st.session_state["justificacion_manual"] = nodos_just
                    st.success(f"✅ Justificación extraída con {len(nodos_just)} párrafos.")
                else:
                    st.session_state["justificacion_manual"] = [] 
                    st.error("❌ La extracción de Justificación devolvió una lista vacía.")
                        
        #  EL EXPANDER DE AUDITORÍA DE TABLAS 
        with st.expander("🔍 Auditoría de Tablas (Búsqueda por Texto Plano)"):
            # Usamos la variable que ya tienes definida en tu flujo
            if archivo_dm:
                try:
                    # 1. Resetear el puntero para lectura limpia
                    archivo_dm.seek(0)
                    doc_para_auditar = Document(archivo_dm)
                    
                    # 2. Llamamos a la función que mapea y muestra TODO
                    # Esta función es la que definimos para ver todos los títulos y tablas
                    auditar_tablas_maestro(doc_para_auditar)
                    
                except Exception as e:
                    st.error(f"Error al auditar tablas: {e}")

                  # 2. Ejecutar Extracciones (Usando tu nomenclatura)
                texto_fund = extraer_fundamentacion(st.session_state["dict_maestro"])
                #texto_fund = extraer_fundamentacion(dict_m)
                texto_especifica = extraer_area_especifica(dict_m)
                texto_just = extraer_justificacion_programa(st.session_state["dict_maestro"])
                texto_prof_exp = extraer_perfil_generico(dict_m, ["perfil", "profesional", "experiencia"])
                texto_prof_egr = extraer_perfil_generico(dict_m, ["perfil", "profesional", "egresado"])
                texto_ocupacional = extraer_perfil_generico(dict_m, ["perfil", "ocupacional"])
                texto_rapa = extraer_resultados_aprendizaje(dict_m)
                
                


                #RESULTADOS DE ESPECÍFICA
                if texto_especifica:
                    st.success(f"✅ Fund. Específica: {len(texto_especifica)} caracteres detectados.")
                    st.session_state["fund_especifica_txt"] = texto_especifica
                else:
                    st.error("❌ No se encontró 'Fundamentación específica del programa'.") 

               
                # RESULTADOS DE PERFILES
                # Perfil Profesional con Experiencia
                if texto_prof_exp:
                    st.success(f"✅ Perfil Profesional con Experiencia: {len(texto_prof_exp)} caracteres.")
                    st.session_state["perfil_profesional_experiencia_txt"] = texto_prof_exp
                else:
                    st.error("❌ No se encontró 'Perfil Profesional con Experiencia (Rediseño)'.")

                # Perfil Profesional del Egresado
                if texto_prof_egr:
                    st.success(f"✅ Perfil Profesional del Egresado: {len(texto_prof_egr)} caracteres.")
                    st.session_state["perfil_profesional_egresado_txt"] = texto_prof_egr
                else:
                    st.error("❌ No se encontró 'Perfil Profesional del Egresado (Rediseño)'.")

                # Perfil Ocupacional
                if texto_ocupacional:
                    st.success(f"✅ Perfil Ocupacional: {len(texto_ocupacional)} caracteres.")
                    st.session_state["perfil_ocupacional_txt"] = texto_ocupacional
                else:
                    st.error("❌ No se encontró 'Perfil Ocupacional (Rediseño)'.")

                #RESULTADOS ACADÉMICOS
                if texto_rapa:
                    st.session_state["resultados_aprendizaje_txt"] = texto_rapa
                    st.success(f"✅ RAPA detectado (Primer párrafo omitido): {len(texto_ocupacional)} caracteres.")
                else:
                    st.error("❌ No se encontró 'Resultados Académicos'.")
            
# LÓGICA DE MODALIDAD

with st.expander("Buscador Información general del Programa por SNIES", expanded=True):
    st.subheader("1. Búsqueda del Programa por SNIES")
    
    col_busq, col_btn = st.columns([3, 1])
    
    with col_busq:
        snies_a_buscar = st.text_input("Ingresa el código SNIES:", placeholder="Ej: 54862", key="search_snies_tmp")
        
    with col_btn:
        st.write(" ")
        st.write(" ")
        if st.button("🔍 Consultar Base de Datos"):
            if snies_a_buscar in BD_PROGRAMAS:
                datos_encontrados = BD_PROGRAMAS[snies_a_buscar]

                # 1. Borramos las llaves viejas para que el formulario no se bloquee
                llaves_a_limpiar = ["denom_input", "titulo_input", "snies_input", "acuerdo_input", "instancia_input", "reg1", "cred", "periodo_idx", "estudiantes_input", "acred1", "lugar"
]
                for k in llaves_a_limpiar:
                    if k in st.session_state:
                        del st.session_state[k]
                
                # 2. Inyectamos los nuevos datos del Excel
                for key, valor in datos_encontrados.items():
                    st.session_state[key] = valor
                
                # 3. Guardamos el SNIES que acabamos de buscar
                st.session_state["snies_input"] = snies_a_buscar
                
                st.success(f"✅ Programa encontrado: {datos_encontrados.get('denom_input')}")
                st.rerun()
            else:
                st.error("❌ Código SNIES no registrado en el sistema.")

    st.markdown("---")

# BOTÓN DE DATOS DE EJEMPLO
if st.button("Llenar con datos de ejemplo"):
    for k in ["denom_input", "titulo_input", "snies_input"]:
        if k in st.session_state:
            del st.session_state[k]
    st.session_state.ejemplo = {
        "denom_input": "Ingeniería de Sistemas",
        "titulo_input": "Ingeniero de Sistemas",
        "nivel_idx": 2, # Profesional universitario
        "area_input": "Ingeniería, Arquitectura y Urbanismo",
        "modalidad_input": 4, # Presencial y Virtual
        "acuerdo_input:": "Acuerdo 012 de 2015",
        "instancia_input": "Consejo Académico",
        "reg1": "Res. 4567 de 2016",
        "reg2": "Res. 8901 de 2023",
        "acred1": "Res. 00234 de 2024",
        "cred": "165",
        "estudiantes_input":"40",
        "periodo_idx": 0, # Semestral
        "lugar": "Sede Principal (Cali)",
        "snies": "54321",
        "motivo": "La creación del Programa se fundamenta en la necesidad de formar profesionales capaces de liderar la transformación digital, diseñar y desarrollar soluciones de software de alta complejidad, gestionar sistemas de información y responder de manera innovadora a los retos tecnológicos, organizacionales y sociales del entorno local, nacional e internacional.",
        "p1_nom": "EO1", "p1_fec": "Acuerdo 012-2015",
        "p2_nom": "EO2", "p2_fec": "Acuerdo 088-2020",
        "p3_nom": "EO3", "p3_fec": "Acuerdo 102-2024",
        #DATOS CAPÍTULO 2
        "objeto_nombre": "Sistemas de información",
        "objeto_concep": "Los sistemas de información son conjuntos organizados de personas, datos, procesos, tecnologías y recursos que interactúan de manera integrada para capturar, almacenar, procesar, analizar y distribuir información, con el fin de apoyar la toma de decisiones, la gestión operativa, el control organizacional y la generación de conocimiento. Estos sistemas permiten transformar los datos en información útil y oportuna, facilitando la eficiencia, la innovación y la competitividad en organizaciones de distintos sectores. Su diseño y gestión consideran aspectos técnicos, organizacionales y humanos, garantizando la calidad, seguridad, disponibilidad y uso ético de la información.",        
        "fund_epi": "El programa se inscribe en el racionalismo crítico y el pragmatismo tecnológico, vinculando la ciencia de la computación con la ingeniería aplicada.",
        # DATOS PARA LAS TABLAS (Se guardan como listas de diccionarios)
        "recon_data": [
            {"Año": "2024", "Nombre del premio": "Excelencia Académica", "Nombre del Ganador": "Juan Pérez", "Cargo": "Docente"}
        ],
        "tabla_cert_ej": [
            {"Nombre": "Desarrollador Web Junior", "Curso 1": "Programación I", "Créditos 1": 3, "Curso 2": "Bases de Datos", "Créditos 2": 4},
            {"Nombre": "Analista de Datos", "Curso 1": "Estadística", "Créditos 1": 4, "Curso 2": "Python para Ciencia", "Créditos 2": 4}
        ], #         
        "referencias_data": [
            {
                "Año": "2021", 
                "Autor(es)": "Sommerville, I.", 
                "Revista": "Computer science", 
                "Título del artículo/Libro": "Engineering Software Products"
            },
            {
                "Año": "2023", 
                "Autor(es)": "Pressman, R. & Maxim, B.", 
                "Revista": "Software Engineering Journal", 
                "Título del artículo/Libro": "A Practitioner's Approach"
            }
        ],
    }

    
    st.rerun()

# --- FORMULARIO DE ENTRADA ---
with st.form("pep_form"):
    # 1. Recuperamos datos de ejemplo si existen
    ej = st.session_state.get("ejemplo", {})

    st.markdown("### 1. Identificación General")
    col1, col2 = st.columns(2)
    
    with col1:
        # Denominación del programa
        # 1. Lógica de inicialización (asegura que siempre tenga un valor inicial si está vacío)
        if "denom_input" not in st.session_state:
            st.session_state["denom_input"] = ej.get("denom_input", "")

        if "facultad" not in st.session_state:
            st.session_state["facultad"] = ej.get("facultad", "")
            
        if "departamento" not in st.session_state:
            st.session_state["departamento"] = ej.get("departamento", "")
        
        denom = st.text_input(
            "Denominación del programa :red[•]", 
            key="denom_input"
        )

        # Título otorgado (Ahora bien indentado dentro de col1)
        titulo = st.text_input(
            "Título otorgado :red[•]", 
            value=st.session_state.get("titulo_input", ej.get("titulo_input", "")),
            key="titulo_input"
        )
    
        niveles_opciones = ["Técnico", "Tecnológico", "Profesional universitario", "Especialización", "Maestría", "Doctorado"]
        val_nivel = st.session_state.get("nivel_idx", st.session_state.get("ejemplo", {}).get("nivel_idx", 2))
        try:
            idx_final = int(val_nivel)
        except (ValueError, TypeError):
            idx_final = 2 # Por defecto Profesional
        
        nivel = st.selectbox(
            "Nivel de formación :red[•]", 
            options=niveles_opciones, 
            index=idx_final,
            key="nivel_formacion_widget"
        )
        # Código SNIES
        snies = st.text_input(
            "Código SNIES", 
            value=st.session_state.get("snies_input", ej.get("snies_input", "")),
            key="snies_input"
            )
        # 5. Número de Semestres 
        semestres = st.text_input(
            "Número de semestres (actuales) :red[•]",
            value=st.session_state.get("semestres_input", ej.get("semestres_input", "")),
            placeholder="Ej: 10",
            key="semestres_input"
        )
    
    with col2:
        idx_mod = st.session_state.get("modalidad_idx", 0)
        modalidad = st.selectbox(
            "Modalidad de oferta :red[•]", 
            ["Presencial", "Virtual", "A Distancia", "Dual", "Presencial y Virtual", "Presencial y a Distancia", "Presencial y Dual"],
            index=int(idx_mod) if isinstance(idx_mod, (int, float)) else 0,
            key="modalidad_input"
        )
        
        acuerdo = st.text_input(
            "Acuerdo de creación / Norma interna :red[•]", 
            key="acuerdo_input"
        )

        # Instancia interna
        instancia = st.text_input(
            "Instancia interna que aprueba :red[•]", 
            key="instancia_input"
        )

        # --- Fila 5: Periodicidad y Créditos ---
        col5_1, col5_2 = st.columns(2)
        
        with col5_1:
            periodicidad = st.selectbox(
                "Periodicidad de admisión :red[•]",
                ["Semestral", "Anual", "Trimestral", "Cuatrimestral"],
                index=0,
                key="periodicidad_input"
            )
    
        with col5_2:
            # --- TRUCO DE LIMPIEZA ---
            # Si "cred" ya existe en session_state y no es texto, lo convertimos a la fuerza
            if "cred" in st.session_state and not isinstance(st.session_state["cred"], str):
                st.session_state["cred"] = str(st.session_state["cred"])
            
            # Ahora sí, extraemos el valor inicial con seguridad
            valor_inicial_creditos = str(st.session_state.get("cred", ej.get("cred", "")))
            
            creditos = st.text_input(
                "Créditos académicos (actuales) :red[•]",
                value=valor_inicial_creditos,
                placeholder="Ej: 160",
                key="cred"
            )
    
        # --- Fila 6: Lugar y Estudiantes ---
        col6_1, col6_2 = st.columns(2)
        
        with col6_1:
            lugar = st.text_input(
                "Lugar de desarrollo :red[•]",
                value=st.session_state.get("lugar_input", ej.get("lugar_input", "Medellín - Campus Robledo")),
                key="lugar_input"
            )
    
        with col6_2:
            # --- PROTECCIÓN CONTRA TYPEERROR ---
            # Si el valor en session_state no es string, lo convertimos ahora mismo
            if "estudiantes_input" in st.session_state and not isinstance(st.session_state["estudiantes_input"], str):
                st.session_state["estudiantes_input"] = str(st.session_state["estudiantes_input"])
            
            # Aseguramos que el valor inicial sea string también desde el diccionario 'ej'
            valor_estudiantes = str(st.session_state.get("estudiantes_input", ej.get("estudiantes_input", "")))
            
            estudiantes_primer = st.text_input(
                "Número de estudiantes en primer periodo :red[•]",
                value=valor_estudiantes,
                placeholder="Ej: 40",
                key="estudiantes_input"
            )

    st.markdown("---")
    st.markdown("### 2. Registros y Acreditaciones")
    def forzar_texto(key, fuente):
        # 1. Recuperamos el valor (de la sesión o del ejemplo)
        valor = st.session_state.get(key, fuente.get(key, ""))
        
        # 2. Si es None, lo convertimos a vacío
        if valor is None:
            valor = ""
        
        # 3. Lo convertimos a String (texto) sí o sí, y actualizamos la sesión
        # Esto sobreescribe cualquier "basura" (números o nulos) que haya quedado en memoria
        st.session_state[key] = str(valor)
   
    with st.container(border=True):
        col_reg, col_acred = st.columns(2)

        with col_reg:
            st.markdown("#### **Registros Calificados**")
                              
            forzar_texto("reg1", ej)
            st.text_input(
                "Resolución Registro Calificado 1 :red[•]", 
                placeholder="Ej: Resolución 12345 de 2023",
                key="reg1"
            )
            
            # --- REGISTRO 2 ---
            forzar_texto("reg2", ej)
            st.text_input(
                "Resolución Registro Calificado 2", 
                placeholder="Ej: Resolución 67890 de 2023",
                key="reg2"
            )

            # --- REGISTRO 3 ---
            forzar_texto("reg3", ej)
            st.text_input(
                "Resolución Registro Calificado 3", 
                placeholder="Dejar vacío si no aplica",
                key="reg3"
            )
            
        with col_acred:
            st.markdown("#### **Acreditaciones**")
            
            # --- ACREDITACIÓN 1 ---
            forzar_texto("acred1", ej)
            st.text_input(
                "Resolución Acreditación Alta Calidad 1", 
                placeholder="Ej: Resolución 012345 de 2022",
                key="acred1"
            )
            
            # --- ACREDITACIÓN 2 ---
            forzar_texto("acred2", ej)
            st.text_input(
                "Resolución Acreditación Alta Calidad 2", 
                placeholder="Dejar vacío si no aplica",
                key="acred2"
            )
    
    frase_auto = f"La creación del Programa {denom} se fundamenta en la necesidad de "
    val_motivo = ej.get("motivo", frase_auto)
    motivo = st.text_area("Motivo de creación :red[•]", value=val_motivo, height=150, key="motivo_input")
      
    st.markdown("---")
    st.markdown("### 3. Modificaciones al Plan de Estudios")
    p_col1, p_col2, p_col3 = st.columns(3)
    with p_col1:
        p1_nom = st.text_input("Nombre Plan v1:red[•]", value=ej.get("p1_nom", ""), key="p1_nom")
        p1_fec = st.text_input("Acuerdo aprobación Plan v1 :red[•]", value=ej.get("p1_fec", ""), key="p1_fec")
        p1_cred = st.text_input("Número de créditos Plan v1 :red[•]", value=ej.get("p1_cred", ""), key="p1_cred")
        p1_sem = st.text_input("Número de semestres Plan v1:red[•]", value=ej.get("p1_sem", ""), key="p1_sem")
    with p_col2:
        p2_nom = st.text_input("Nombre Plan v2 (Opcional)", value=ej.get("p2_nom", ""), key="p2_nom")
        p2_fec = st.text_input("Acuerdo aprobación Plan v2 (Opcional)", value=ej.get("p2_fec", ""), key="p2_fec")
        p2_cred = st.text_input("Número de créditos Plan v2 (Opcional) :red[•]", value=ej.get("p2_cred", ""),key="p2_cred")
        p2_sem = st.text_input("Número de semestres Plan v2 (Opcional):red[•]",value=ej.get("p2_sem", ""),key="p2_sem")
    with p_col3:
        p3_nom = st.text_input("Nombre Plan v3 (Opcional)", value=ej.get("p3_nom", ""), key="p3_nom")
        p3_fec = st.text_input("Acuerdo aprobación Plan v3 (Opcional)", value=ej.get("p3_fec", ""), key="p3_fec")
        p3_cred = st.text_input("Número de créditos Plan v3 (Opcional)", value=ej.get("p3_cred", ""),key="p3_cred")
        p3_sem = st.text_input("Número de semestresPlan v3(Opcional)", value=ej.get("p3_sem", ""), key="p3_sem")

    st.markdown("---")
    st.markdown("### 🏆 4. Reconocimientos (Opcional)")
    recon_data = st.data_editor(
        ej.get("recon_data", [{"Año": "", "Nombre del premio": "", "Nombre del Ganador": "", "Cargo": "Estudiante"}]),
        num_rows="dynamic",
        key="editor_recon", # Es vital tener una key única
        column_config={
            "Cargo": st.column_config.SelectboxColumn(options=["Docente", "Líder", "Decano", "Estudiante","Docente Investigador", "Investigador"])
        },
        use_container_width=True
    )
    st.session_state["recon_data"] = recon_data
    
#CAPÍTULO 2
    st.markdown("---")
    st.header("2. Referentes Conceptuales")
   # 2. MODO MANUAL Objeto de conocimiento del Programa
    val_obj_nombre = ej.get("objeto_nombre", "")
    
    objeto_nombre = st.text_input(
        "1. ¿Cuál es el Objeto de conocimiento del Programa? :red[•]",
        value=st.session_state.get("obj_nombre_input", val_obj_nombre),
        placeholder="Ejemplo: Sistemas de información",
        key="obj_nombre_input"  #
    )
    # 2. Definición del Objeto (Lo que llenará {{def_oc}})
    if metodo_trabajo == "Manual":
    
        st.write("---")
        st.write("**Definición del Objeto de Conocimiento**")
        
    # Si es manual, el usuario escribe directamente la definición
        st.text_area(
                "Escriba la definición del Objeto de Conocimiento:",
                value=st.session_state.get("def_oc_manual", ""),
                placeholder="Ingrese el texto aquí...",
                key="def_oc_manual",
                height=200
            )  
         
    # 3. REFERENCIAS (Esto sigue igual para ambos casos)
    if metodo_trabajo != "Semiautomatizado (Cargar Documento Maestro)":
        st.write(" ")
        st.write("**Referencias bibliográficas**")
        st.caption("Sustentan la conceptualización del Objeto de Conocimiento.")
        referencias_previa = ej.get("referencias_data", [
            {"Año": "", "Autor(es) separados por coma": "", "Revista": "", "Título del artículo/Libro": ""}
        ])
        
        referencias_data = st.data_editor(
            referencias_previa,
            num_rows="dynamic",
            key="editor_referencias",
            use_container_width=True,
            column_config={
                "Año": st.column_config.TextColumn("Año", width="small"),
                "Autor(es)": st.column_config.TextColumn("Autor(es)", width="medium"),
                "Revista": st.column_config.TextColumn("Revista", width="medium"),
                "Título del artículo/Libro": st.column_config.TextColumn("Título del artículo/Libro", width="large"),
            }
        )

  # 2.2. Fundamentación epistemológica en Pestañas ---
    st.markdown("---")
    st.subheader("2.2. Fundamentación epistemológica")
    if metodo_trabajo != "Semiautomatizado (Cargar Documento Maestro)":
        
        # CASO 1: MODO MANUAL 
        st.info("Utilice las pestañas para completar los tres párrafos de la Fundamentación epistemológica.")
        
        #  Creamos las tabs SOLO si es manual 
        tab1, tab2, tab3 = st.tabs(["Párrafo 1", "Párrafo 2", "Párrafo 3"])

        # Configuración de columnas para referencias
        config_columnas_ref = {
            "Año": st.column_config.TextColumn("Año", width="small"),
            "Autor(es) separados por coma": st.column_config.TextColumn("Autor(es)", width="medium"),
            "Revista": st.column_config.TextColumn("Revista", width="medium"),
            "Título del artículo/Libro": st.column_config.TextColumn("Título del artículo/Libro", width="large"),
        }

        # Bloque Párrafo 1
        with tab1:
            st.markdown("### Párrafo 1: Marco filósofico")
            st.text_area(
                "¿Cuál es la postura filosófica predominante? :red[•]",
                value=ej.get("fund_epi_p1", ""), 
                height=200,
                key="input_epi_p1",
                placeholder="Ejemplo: El programa se fundamenta en el paradigma de la complejidad..."
            )
            st.write("Referencias bibliográficas (Párrafo 1):")
            st.data_editor(
                ej.get("referencias_epi_p1", [{"Año": "", "Autor(es) separados por coma": "", "Revista": "", "Título del artículo/Libro": ""}]),
                num_rows="dynamic",
                key="editor_refs_p1",
                use_container_width=True,
                column_config=config_columnas_ref
            )

        # Bloque Párrafo 2
        with tab2:
            st.markdown("### Párrafo 2: Identidad disciplinar")
            st.text_area(
                "Origen etimológico y teorías conceptuales :red[•]",
                value=ej.get("fund_epi_p2", ""), 
                height=200,
                key="input_epi_p2",
                placeholder="Ejemplo: La identidad de este programa se define desde..."
            )
            st.write("Referencias bibliográficas (Párrafo 2):")
            st.data_editor(
                ej.get("referencias_epi_p2", [{"Año": "", "Autor(es) separados por coma": "", "Revista": "", "Título del artículo/Libro": ""}]),
                num_rows="dynamic",
                key="editor_refs_p2",
                use_container_width=True,
                column_config=config_columnas_ref
            )

        # Bloque Párrafo 3
        with tab3:
            st.markdown("### Párrafo 3: Intencionalidad social")
            st.text_area(
                "¿Intervención ética y transformadora? :red[•]",
                value=ej.get("fund_epi_p3", ""), 
                height=200,
                key="input_epi_p3",
                placeholder="Ejemplo: Finalmente, la producción de conocimiento..."
            )
            st.write("Referencias bibliográficas (Párrafo 3):")
            st.data_editor(
               ej.get("referencias_epi_p3", [{"Año": "", "Autor(es) separados por coma": "", "Revista": "", "Título del artículo/Libro": ""}]),
                num_rows="dynamic",
                key="editor_refs_p3",
                use_container_width=True,
                column_config=config_columnas_ref
            )

    else:
        
        # CASO 2: MODO AUTOMATIZADO (SIN pestañas)
        st.success("✅ Modo Estructurado: El sistema extraerá automáticamente el contenido de la sección 'Conceptualización teórica y epistemológica del programa' desde el Documento Maestro.")
        # No hay col_ini ni col_fin aquí
    
        #st.info("Configuración de Extracción:  Indique dónde inicia y termina la Conceputalización Teórica y Epistemológica en el Documento Maestro. Fundamentación Epistemológica")
        
        # Aquí NO usamos st.tabs, usamos columnas directas
        #with st.container(border=True):
         #   col_inicio, col_fin = st.columns(2)
            
          #  with col_inicio:
           #     st.text_input(
            #        "Texto de inicio :red[•]", 
             #       placeholder="Ej: 2.2 Fundamentación Epistemológica",
              #      help="Copia y pega las primeras palabras del capítulo en el Word.",
               #     key="txt_inicio_fund_epi"
                #)
            
            #with col_fin:
             #   st.text_input(
              #      "Texto final :red[•]", 
               #     placeholder="Ej: 2.3 Justificación",
                #    help="Copia y pega las primeras palabras del SIGUIENTE capítulo o donde termina este.",
                 #   key="txt_fin_fund_epi"
                #)


    
  # --- 2.3. Fundamentación Académica ---
    if metodo_trabajo != "Semiautomatizado (Cargar Documento Maestro)":
        st.markdown("---")
        st.subheader("2.3. Fundamentación Académica")
        
        # 2.3.1 MICROCREDENCIALES (Siempre visible)
    
        st.write("***2.3.1. Microcredenciales***")
        st.info("Agregue filas según sea necesario para listar las microcredenciales.")
        
        datos_micro = ej.get("tabla_micro", [
            {"Nombre de la Certificación": "", "Nombre del Curso": "", "Créditos": 0}
        ])
        
        st.data_editor(
            datos_micro,
            num_rows="dynamic", 
            key="editor_microcredenciales",
            use_container_width=True,
            column_config={
                "Nombre de la Certificación": st.column_config.TextColumn("Certificación", width="medium"),
                "Nombre del Curso": st.column_config.TextColumn("Curso Asociado", width="medium"),
                "Créditos": st.column_config.NumberColumn("Créditos", min_value=0, step=1, width="small")
            }
        )
    
        st.write(" ") 
    
        # 2.3.2 MACROCREDENCIALES 
        st.write("***2.3.2. Macrocredenciales***")
        st.info("Cada fila representa una Certificación (Macrocredencial). Complete los cursos que la componen (máx 3).")
    
        datos_macro = ej.get("tabla_macro", [
            {
                "Certificación": "", 
                "Curso 1": "", "Créditos 1": 0,
                "Curso 2": "", "Créditos 2": 0,
                "Curso 3": "", "Créditos 3": 0
            }
        ])
    
        columnas_config = {
            "Certificación": st.column_config.TextColumn(
                "Nombre Macrocredencial", 
                width="medium",
                help="Nombre de la certificación global (ej: Diplomado en Big Data)",
                required=True
            ),
            "Curso 1": st.column_config.TextColumn("Curso 1", width="medium"),
            "Créditos 1": st.column_config.NumberColumn("Créd. 1", width="small", min_value=0, step=1),
            "Curso 2": st.column_config.TextColumn("Curso 2", width="medium"),
            "Créditos 2": st.column_config.NumberColumn("Créd. 2", width="small", min_value=0, step=1),
            "Curso 3": st.column_config.TextColumn("Curso 3", width="medium"),
            "Créditos 3": st.column_config.NumberColumn("Créd. 3", width="small", min_value=0, step=1),
        }
    
        st.data_editor(
            datos_macro,
            num_rows="dynamic", 
            key="editor_macrocredenciales",
            use_container_width=True,
            column_config=columnas_config
        )
    else:
    # MENSAJE PARA EL MODO SEMIAUTOMATIZADO
        st.markdown("---")
        st.subheader("2.3. Fundamentación Académica")

        # Recuperamos el mapa del estado de sesión
        mapa = st.session_state.get("mapa_tablas", {})

        if not mapa:
            st.error("❌ No se detectaron tablas en el Documento Maestro. Revise la carga del archivo.")
        else:
            # En lugar de selectores, mostramos un resumen de éxito
            st.success("✅ Modo Estructurado: El sistema extraerá automáticamente el contenido de la sección 'Micro y Macro credenciales' desde el Documento Maestro.")          
            
    # 2.3.3 ÁREAS DE FORMACIÓN (Condicional)
    st.write("") 
    st.write("**2.3.3. Áreas de formación**")
    
    # CASO MANUAL
    if metodo_trabajo != "Semiautomatizado (Cargar Documento Maestro)":
        area_especifica = st.text_area(
            "Descripción del Área de Fundamentación Específica :red[•]",
            value=ej.get("fund_especifica_desc", ""),
            height=150,
            placeholder="Desarrolla competencias técnicas y profesionales específicas del programa...",
            key="input_area_especifica"
        )
    # CASO AUTOMATIZADO
    else:
        st.info("**Configuración de Extracción Automática: Área Específica**")
    
        with st.container(border=True):
            st.success("✅ Modo Estructurado: El sistema extraerá automáticamente el contenido de la sección 'Áreas de formación' desde el Documento Maestro.")
            
        # Opcional: Un verificador rápido de esta sección específica
            if st.session_state.get("dict_maestro"):
            # Supongamos que ya tienes una función extraer_area_especifica
                texto_area = extraer_area_especifica(st.session_state["dict_maestro"])
                if texto_area:
                    with st.expander("Ver contenido detectado para Área Específica"):
                        st.write(texto_area)
                else:
                    st.error("⚠️ No se encontró la sección 'Área Específica' en el DM cargado.")

   
# 2.3.4 CURSOS POR ÁREA
    # CONDICIONAL: Solo se muestra si el método es estrictamente Manual
    if metodo_trabajo != "Semiautomatizado (Cargar Documento Maestro)":
        
        st.write("---") # Divisor para mantener orden visual
        st.write("***2.3.4. Cursos por área de formación***")
        
        # Lista de áreas en el orden solicitado
        areas_formacion = [
            "Formación Humanística",
            "Fundamentación Básica",
            "Formación Básica Profesional",
            "Fundamentación Específica del Programa",
            "Formación Flexible o Complementaria"
        ]

        # Solo mostramos la información informativa del modo Manual
        st.info("En el documento final, asegúrese de incluir las tablas de cursos organizadas por:")
        for area in areas_formacion:
            st.write(f"- {area}")
            
    # Si es Semiautomatizado, el bloque se ignora por completo y no aparece nada en el front.

# Itinerario formativo
    st.write("") 
    st.subheader("3. Itinerario formativo")
    
    # Texto de fundamento original intacto
    st.write("Teniendo como fundamento que, en torno a un objeto de conocimiento se pueden estructurar varios programas a diferentes niveles de complejidad, es importante expresar si el programa en la actualidad es único en torno al objeto de conocimiento al que está adscrito o hay otros de mayor o de menor complejidad. :red[•]")
    
    # Nota simple del límite
    st.caption("Nota: Máximo 500 palabras.")

    # 1. Asegurar que la clave exista en el estado
    if "input_itinerario" not in st.session_state:
        st.session_state["input_itinerario"] = ej.get("fund_especifica_desc", "")

    # 2. EL EDITOR CON BOTONES (Sin lógica de conteo en pantalla)
    contenido_quill = st_quill(
        value=st.session_state["input_itinerario"],
        placeholder="Ejemplo si el PEP es de Ingeniería Mecánica, determinar si hay otro programa de menor complejidad como una tecnología Mecánica o uno de mayor complejidad como una especialización o una maestría. Este itinerario debe considerar posibles programas de la misma naturaleza que se puedan desarrollar en el futuro.",
        key="quill_itinerario_final", 
        toolbar=["bold", "italic"],
        html=True
    )

    # 3. Guardar el resultado para el Word
    if contenido_quill is not None:
        st.session_state["input_itinerario"] = contenido_quill
    

# Justificación del Programa
    st.write("") 
    st.subheader("4. Justificación del Programa")
    
    # CONDICIONAL: Manual vs Semiautomatizado
    if metodo_trabajo != "Semiautomatizado (Cargar Documento Maestro)":
        
        # CASO 1: MODO MANUAL
        st.write("**Redacción Manual de la Justificación**")
        
        # Inicialización de la variable en session_state si no existe
        if "input_just_manual" not in st.session_state:
            st.session_state["input_just_manual"] = ej.get("justificacion_desc", "")

        st.text_area(
            "Demostrar la relevancia del programa en el contexto actual, resaltando su impacto en la solución de problemáticas sociales y productivas. Se debe enfatizar cómo la formación impartida contribuye al desarrollo del entorno local, regional y global... :red[•]",
            height=250,
            placeholder="Fundamentar la relevancia del programa con datos actualizados, referencias normativas y estudios sectoriales...",
            key="input_just_manual"
        )

    else:
        # CASO 2: MODO SEMIAUTOMATIZADO
         st.success("✅Modo Estructurado: El sistema extraerá automáticamente el contenido de la sección 'Justificación del Programa' desde el Documento Maestro.")
   
#  SECCIÓN 5: ESTRUCTURA CURRICULAR 
    st.markdown("---")
    st.subheader("5. Estructura Curricular")
    
    # CONDICIONAL: Manual vs Semiautomatizado
    if metodo_trabajo != "Semiautomatizado (Cargar Documento Maestro)":
        
        # CASO 1: MODO MANUAL
        st.info("Defina el objeto de conocimiento y relacione las perspectivas de intervención con sus respectivas competencias.")

        # 1. Sector social y/o productivo
        with st.container(border=True):
            st.write("***Sector Social y/o Productivo***")
            st.text_area(
                " Sector Social y/o Productivo en el que interviene el Programa :red[•]",
                placeholder="Ejemplo: Sector manufactura...",
                key="sector",
                height=50
            )

        st.write("") # Espacio
        st.write("***Perspectivas de Intervención y Competencias***")
        st.markdown("Complete los cuadros paralelos a continuación:")

        # 2. Generación de los 6 Cuadros Paralelos
        for i in range(1, 7):
            with st.container(border=True):
                st.markdown(f"**Relación de Desempeño #{i}**")
                col_izq, col_der = st.columns(2)
                
                with col_izq:
                    st.text_area(
                        f"Objeto de Formación / Perspectiva de intervención {i}",
                        placeholder=f"Defina la perspectiva {i}...",
                        key=f"objeto_formacion_{i}",
                        height=100
                    )
                    
                with col_der:
                    st.text_area(
                        f"Competencia de Desempeño Profesional {i}",
                        placeholder=f"Defina la competencia {i}...",
                        key=f"competencia_desempeno_{i}",
                        height=100
                    )

        st.caption("Nota: No es obligatorio llenar los 6 campos. El sistema procesará solo aquellos que contengan información.")

    else:
        # CASO 2: MODO SEMIAUTOMATIZADO
        st.success("✅ **Modo Estructurado:** El sistema extraerá automáticamente la 'Pertinencia social y académica del Programa' desde el Documento Maestro.")

    
    #-- 5.2. Pertinencia Académica ---
    # Al poner el IF aquí, si es Semiautomatizado, Streamlit ignora todo lo que sigue
    if metodo_trabajo != "Semiautomatizado (Cargar Documento Maestro)":
        st.markdown("---")
        st.write("***5.2. Pertinencia Académica***")

        # Contenido exclusivo del Modo Manual
        st.info("En el modo manual, redacte la pertinencia académica directamente en su documento final.")
        st.text_area(
            "Descripción de la Pertinencia Académica (Opcional)",
            placeholder="Describa cómo el programa se alinea con las tendencias académicas actuales...",
            key="input_pertinencia_manual",
            height=150
        )

    # --- 5.3. Plan de Estudios ---
    st.markdown("---")
    st.write("***5.3. Plan de Estudios***")
    
    st.info("Por favor, cargue la imagen del Plan de Estudios (Malla Curricular) para ser incluida en el documento.")

    # Contenedor de carga de archivo
    with st.container(border=True):
        archivo_imagen = st.file_uploader(
            "Seleccionar imagen del Plan de Estudios :red[•]",
            type=["png", "jpg", "jpeg"],
            help="Soporta formatos PNG, JPG y JPEG. Esta imagen se insertará en la sección 5.3 del Word.",
            key="upload_plan_estudios"
        )

        # Si el usuario sube un archivo, mostrar una vista previa
        if archivo_imagen is not None:
            st.success("✅ Imagen cargada correctamente.")
            
            # Mostramos una vista previa pequeña/mediana
            st.image(
                archivo_imagen, 
                caption="Vista previa del Plan de Estudios cargado", 
                use_container_width=True
            )
            
            # Opción para que el usuario añada un título o fuente a la imagen
            st.text_input(
                "Título/Nota de la imagen (Opcional):",
                value="Gráfico: Plan de Estudios del Programa",
                key="caption_plan_estudios"
            )
        else:
            st.warning("⚠️ No se ha cargado ninguna imagen aún.")

     # --- 5.4 PERFILES ---
    if metodo_trabajo != "Semiautomatizado (Cargar Documento Maestro)":
        st.write("") 
        st.markdown("---")
        st.header("5.4. Perfiles")
        
        # ==========================================
        # CASO: MODO MANUAL
        # ==========================================
        st.info("Defina manualmente los perfiles que caracterizan al programa")

        with st.container(border=True):
            col_prof, col_egr, col_ocup = st.columns(3)
            
            with col_prof:
                st.markdown("### **Perfil Profesional con Experiencia.**")
                st.text_area(
                    "Defina el perfil del profesional con experiencia :red[•]",
                    placeholder="Describa las capacidades y trayectoria que se esperan del profesional...",
                    key="perfil_profesional_exp",
                    height=300
                )
                
            with col_egr:
                st.markdown("### **Perfil Profesional del Egresado.**")
                st.text_area(
                    "Defina el perfil profesional del egresado :red[•]",
                    placeholder="Describa las competencias y conocimientos con los que sale el estudiante...",
                    key="perfil_profesional_egresado",
                    height=300
                )
                
            with col_ocup:
                st.markdown("### **Perfil Ocupacional.**")
                st.text_area(
                    "Defina el perfil ocupacional :red[•]",
                    placeholder="Mencione los cargos, sectores y áreas donde podrá desempeñarse...",
                    key="perfil_ocupacional",
                    height=300
                )

        # Nota de ayuda para la redacción
        with st.expander("💡 Tips para redactar los perfiles"):
            st.markdown("""
            * **Profesional con experiencia:** Declaración que hace el programa académico acerca del resultado esperado de la formación para toda la vida.
            * **Egresado:** Promesa de valor que la institución hace a los estudiantes y a la sociedad en general.
            * **Ocupacional:** Conjunto de conocimientos, habilidades, destrezas y actitudes que desarrollará el futuro profesional...
            """)


    
# --- 7. RECURSOS ACADÉMICOS ---
    st.markdown("---")
    st.subheader("7. Recursos Académicos")
    
    # 7.1 Entornos académicos
    st.subheader("7.1. Entornos académicos")
    
    st.info("""
        Describa los espacios físicos y virtuales que soportan el programa. 
        Incluya laboratorios, bases de datos, plataformas de aprendizaje (LMS), 
        aulas especializadas y software técnico.
    """)

    with st.container(border=True):
        st.caption("Nota: Máximo 1000 palabras. Use los botones para dar formato.")

        # 1. Asegurar que la clave exista en el estado
        if "input_entornos_academicos" not in st.session_state:
            st.session_state["input_entornos_academicos"] = ej.get("entornos_academicos_desc", "")

        # 2. EL EDITOR CON BOTONES
        entornos_quill = st_quill(
            value=st.session_state["input_entornos_academicos"],
            placeholder="""Ejemplo: El programa cuenta con acceso a laboratorios de última generación...""",
            key="quill_entornos_final",
            toolbar=["bold", "italic"],
            html=True
        )

        # 3. CAPTURA Y VALIDACIÓN INVISIBLE
        if entornos_quill is not None:
            st.session_state["input_entornos_academicos"] = entornos_quill
            
            # Conteo interno
            import re
            texto_limpio = re.sub('<[^<]+?>', '', str(entornos_quill))
            num_palabras = len(texto_limpio.split())
            
            # Solo muestra el error si se excede el límite
            if num_palabras > 1000:
                st.error(f"⚠️ El texto es demasiado largo ({num_palabras} palabras). El límite para esta sección es de 1000 palabras.")
        
    
# --- 7.2. TALENTO HUMANO ---
    st.write("") 
    st.subheader("7.2. Talento Humano")
    
    st.info("""
        Describa el perfil del equipo docente requerido (formación académica, 
        experiencia profesional e investigativa) para garantizar el desarrollo 
        de las funciones de docencia, investigación y extensión del Programa.
    """)

    with st.container(border=True):
        # Nota visual del límite actualizado a 500
        st.caption("Nota: Máximo 500 palabras. Use los botones para dar formato (Negrita/Cursiva).")

        # 1. Asegurar que la clave exista
        if "input_perfil_docente" not in st.session_state:
            st.session_state["input_perfil_docente"] = ej.get("talento_humano_desc", "")

        # 2. EL EDITOR
        talento_quill = st_quill(
            value=st.session_state["input_perfil_docente"], # <-- Cambiado
            placeholder="""Ejemplo...""",
            key="quill_talento_final",
            toolbar=["bold", "italic"],
            html=True
        )

        # 3. CAPTURA
        if talento_quill is not None:
            st.session_state["input_perfil_docente"] = talento_quill
            
            import re
            texto_limpio = re.sub('<[^<]+?>', '', str(talento_quill))
            num_palabras = len(texto_limpio.split())
            
            # Alerta roja solo si se pasa de 500
            if num_palabras > 500:
                st.error(f"⚠️ El texto es demasiado largo ({num_palabras} palabras). El límite para esta sección es de 500 palabras.")
    
# --- 8. INVESTIGACIÓN, TECNOLOGÍA E INNOVACIÓN ---
    st.markdown("---")
    st.header("8. Investigación, Tecnología e Innovación")
    
    # Corregido el st.info para evitar SyntaxError
    st.info("**Indicaciones:** Describa la organización de la investigación en el programa. "
            "Especifique las líneas y grupos de investigación, destacando "
            "objetivos y su articulación con el proceso formativo.")

    with st.container(border=True):
        # 1. Descripción General y Grupos
        st.subheader("Estructura de Investigación")
        st.caption("Nota: Máximo 1000 palabras. Use los botones para dar formato (Negrita/Cursiva).")
        
        # Manejo de Session State
        if "input_investigacion_general" not in st.session_state:
            st.session_state["input_investigacion_general"] = ej.get("investigacion_desc", "")
        
        # 2. EL EDITOR CON BOTONES (Solo Negrita y Cursiva)
        investigacion_quill = st_quill(
            value=st.session_state["input_investigacion_general"],
            placeholder="Ejemplo: El programa se articula con el Grupo de Investigación (Nombre)...",
            key="quill_investigacion_final_v8", # Key única para evitar conflictos
            toolbar=["bold", "italic"], 
            html=True
        )

        # 3. CAPTURA Y VALIDACIÓN INVISIBLE (Límite 1000)
        if investigacion_quill is not None:
            st.session_state["input_investigacion_general"] = investigacion_quill
            
            import re
            # Limpieza para conteo real
            texto_limpio = re.sub('<[^<]+?>', '', str(investigacion_quill))
            num_palabras = len(texto_limpio.split())
            
            # Alerta roja solo si se pasa de 1000
            if num_palabras > 1000:
                st.error(f"⚠️ El texto es demasiado largo ({num_palabras} palabras). El límite para esta sección es de 1000 palabras.")
        

        
    # --- 9. VINCULACIÓN NACIONAL E INTERNACIONAL ---
    st.markdown("---")
    st.header("9. Vinculación Nacional e Internacional")
    
    # 9.1 Estrategias de internacionalización
    st.subheader("9.1. Estrategias de internacionalización")
    
    st.info("""
        **Indicaciones:** Describa las acciones que permiten la visibilidad nacional e internacional del programa. 
        Incluya estrategias como: movilidad académica (estudiantes y docentes), convenios de doble titulación, 
        participación en redes académicas, internacionalización del currículo (COIL, invitados internacionales) 
        y bilingüismo.
    """)

    with st.container(border=True):
        st.caption("Nota: Máximo 1000 palabras. Use los botones para dar formato (Negrita/Cursiva).")

        # 1. Asegurar que la clave exista en el estado
        if "input_internacionalizacion" not in st.session_state:
            st.session_state["input_internacionalizacion"] = ej.get("internacionalizacion_desc", "")

        # 2. EL EDITOR CON BOTONES (Solo Negrita y Cursiva)
        internacionalizacion_quill = st_quill(
            value=st.session_state["input_internacionalizacion"],
            placeholder="""Ejemplo: El programa fomenta la internacionalización a través de convenios marco con universidades de España y México para movilidad estudiantil. 
Se implementa la metodología COIL en las asignaturas de... 
Además, el programa participa activamente en la red (Nombre de la Red) y promueve el bilingüismo...""",
            key="quill_internacionalizacion_final",
            toolbar=["bold", "italic"],
            html=True
        )

        # 3. CAPTURA Y VALIDACIÓN INVISIBLE (Límite 1000)
        if internacionalizacion_quill is not None:
            st.session_state["input_internacionalizacion"] = internacionalizacion_quill
            
            import re
            # Limpieza de etiquetas HTML para el conteo real de palabras
            texto_limpio = re.sub('<[^<]+?>', '', str(internacionalizacion_quill))
            num_palabras = len(texto_limpio.split())
            
            # Alerta roja solo si se excede el límite
            if num_palabras > 1000:
                st.error(f"⚠️ El texto es demasiado largo ({num_palabras} palabras). El límite para esta sección es de 1000 palabras.")

    # Tabla complementaria opcional para convenios específicos
    #with st.expander("📋 Listado de Convenios y Aliados (Opcional)"):
     #   st.write("Si desea tabular los convenios vigentes, lístelos aquí:")
      #  datos_convenios = ej.get("tabla_convenios", [
       #     {"Institución/Aliado": "", "País": "Colombia", "Tipo de Alianza": "Movilidad"}
       # ])
        
   #     st.data_editor(
    #        datos_convenios,
     #       num_rows="dynamic",
      #      use_container_width=True,
       #     key="editor_convenios",
        #    column_config={
         #       "Tipo de Alianza": st.column_config.SelectboxColumn(
          #          "Tipo de Alianza",
           #         options=["Movilidad Académica", "Doble Titulación", "Investigación Conjunta", "Prácticas Profesionales", "Otro"],
            #        required=True
             #   )
            #}
        #)

# --- 10. BIENESTAR UNIVERSITARIO ---
    st.markdown("---")
    st.header("10. Bienestar en el Programa")
    
    st.info("""
        **Indicaciones:** Describa las acciones, programas y servicios de bienestar que 
        impactan directamente a los estudiantes y docentes del programa. 
        Enfoque su respuesta en la **permanencia académica**, el desarrollo humano, 
        la salud, el deporte, la cultura y los apoyos socioeconómicos.
    """)

    with st.container(border=True):
        # Nota visual del límite
        st.caption("Nota: Máximo 500 palabras. Use los botones para dar formato (Negrita/Cursiva).")

        # 1. Asegurar que la clave exista en el estado
        if "input_bienestar" not in st.session_state:
            st.session_state["input_bienestar"] = ej.get("bienestar_desc", "")

        # 2. EL EDITOR CON BOTONES (Solo Negrita y Cursiva)
        bienestar_quill = st_quill(
            value=st.session_state["input_bienestar"],
            placeholder="""Ejemplo: El programa se articula con la Política de Bienestar Institucional a través de estrategias de acompañamiento docente (tutorías) para mitigar el riesgo de deserción... 
Se cuenta con programas de apoyo psicosocial, becas socioeconómicas y fomento de la cultura y el deporte...""",
            key="quill_bienestar_final",
            toolbar=["bold", "italic"],
            html=True
        )

        # 3. CAPTURA Y VALIDACIÓN INVISIBLE (Límite 500)
        if bienestar_quill is not None:
            st.session_state["input_bienestar"] = bienestar_quill
            
            import re
            # Limpieza para el conteo real
            texto_limpio = re.sub('<[^<]+?>', '', str(bienestar_quill))
            num_palabras = len(texto_limpio.split())
            
            # Alerta roja solo si se excede el límite
            if num_palabras > 500:
                st.error(f"⚠️ El texto es demasiado largo ({num_palabras} palabras). El límite para esta sección es de 500 palabras.")

    # Tabla opcional para programas de apoyo específicos
    #with st.expander("📋 Programas Específicos de Apoyo (Opcional)"):
     #   st.write("Si el programa cuenta con apoyos específicos (ej: tutorías especializadas, bonos, convenios), lístelos aquí:")
      #  datos_apoyo = [
       #     {"Programa/Estrategia": "Tutorías Académicas", "Objetivo": "Reducir la pérdida académica"},
        #    {"Programa/Estrategia": "Acompañamiento Psicológico", "Objetivo": "Salud mental y estabilidad"}
        #]
        
      #  st.data_editor(
       #     datos_apoyo,
        #    num_rows="dynamic",
         #   use_container_width=True,
          #  key="editor_apoyos_bienestar"
        #)
        
    # --- 11. ESTRUCTURA ADMINISTRATIVA ---
    st.markdown("---")
    st.header("11. Estructura Administrativa")
    
    # 11.1 Imagen de la Estructura
    st.subheader("11.1. Estructura Administrativa del Programa")
    st.info("""
        **Indicaciones:** Cargue la representación gráfica de la estructura organizativa del programa. 
        Recuerde que debe visualizarse la jerarquía desde la **Vicerrectoría de Enseñanza y Aprendizaje** hacia el Programa.
    """)

    with st.container(border=True):
        img_estructura = st.file_uploader(
            "Cargar Organigrama del Programa (PNG, JPG) :red[•]",
            type=["png", "jpg", "jpeg"],
            key="upload_estructura_admin"
        )
        
        if img_estructura:
            st.image(img_estructura, caption="Vista previa: Estructura Administrativa", use_container_width=True)

    st.write("")

# 11.2 Órganos de decisión (Cuadros Paralelos)
    st.subheader("11.2. Órganos de decisión")
    st.markdown("Describa la conformación y dinámica de los cuerpos colegiados:")

    with st.container(border=True):
        col_comite, col_consejo = st.columns(2)
        
        # --- COLUMNA 1: COMITÉ CURRICULAR ---
        with col_comite:
            st.markdown("### **Comité Curricular**")
            st.caption("Máximo 500 palabras (Negrita/Cursiva)")

            # Inicializar estado si no existe
            if "desc_comite_curricular" not in st.session_state:
                st.session_state["desc_comite_curricular"] = ""

            # Editor Quill
            comite_quill = st_quill(
                value=st.session_state["desc_comite_curricular"],
                placeholder="Conformación (Director, docentes, egresados...), periodicidad de reuniones y funciones...",
                key="quill_comite_final",
                toolbar=["bold", "italic"],
                html=True
            )

            # Validación Invisible
            if comite_quill is not None:
                st.session_state["desc_comite_curricular"] = comite_quill
                import re
                txt_c = re.sub('<[^<]+?>', '', str(comite_quill))
                if len(txt_c.split()) > 500:
                    st.error(f"⚠️ Comité: Límite excedido ({len(txt_c.split())}/500)")
            
        # --- COLUMNA 2: CONSEJO DE FACULTAD ---
        with col_consejo:
            st.markdown("### **Consejo de Facultad**")
            st.caption("Máximo 500 palabras (Negrita/Cursiva)")

            # Inicializar estado si no existe
            if "desc_consejo_facultad" not in st.session_state:
                st.session_state["desc_consejo_facultad"] = ""

            # Editor Quill
            consejo_quill = st_quill(
                value=st.session_state["desc_consejo_facultad"],
                placeholder="Conformación (Decano, representantes...), periodicidad y rol en la toma de decisiones...",
                key="quill_consejo_final",
                toolbar=["bold", "italic"],
                html=True
            )

            # Validación Invisible
            if consejo_quill is not None:
                st.session_state["desc_consejo_facultad"] = consejo_quill
                import re
                txt_f = re.sub('<[^<]+?>', '', str(consejo_quill))
                if len(txt_f.split()) > 500:
                    st.error(f"⚠️ Consejo: Límite excedido ({len(txt_f.split())}/500)")

    # Nota de recordatorio institucional
    st.caption("Nota: Estas descripciones deben estar alineadas con el Estatuto General y los reglamentos internos de la I.U. Pascual Bravo.")

# --- 12. EVALUACIÓN Y MEJORAMIENTO CONTINUO ---
    st.markdown("---")
    st.header("12. Evaluación y Mejoramiento Continuo")
    
    # 12.1 Sistema de Aseguramiento de la Calidad
    st.subheader("12.1. Sistema de Aseguramiento de la Calidad del Programa")
    
    st.info("""
        **Indicaciones:** Describa los procesos específicos del programa para garantizar la calidad académica. 
        Debe evidenciar cómo se evalúa el desempeño, cómo se identifican oportunidades de mejora 
        y la ejecución de planes de acción alineados con la I.U. Pascual Bravo.
    """)

    with st.container(border=True):
        st.caption("Nota: Máximo 1000 palabras. Use los botones para dar formato (Negrita/Cursiva).")

        # 1. Asegurar que la clave exista en el estado
        if "input_aseguramiento_calidad" not in st.session_state:
            st.session_state["input_aseguramiento_calidad"] = ej.get("calidad_mejora_desc", "")

        # 2. EL EDITOR CON BOTONES (Solo Negrita y Cursiva)
        calidad_quill = st_quill(
            value=st.session_state["input_aseguramiento_calidad"],
            placeholder="""Ejemplo: El programa implementa el Modelo de Autoevaluación Institucional, realizando jornadas semestrales de revisión de indicadores de... 
Se recolecta información de fuentes primarias (estudiantes, docentes, egresados y empleadores) para alimentar el Plan de Mejoramiento Continuo (PMC)...""",
            key="quill_calidad_final",
            toolbar=["bold", "italic"],
            html=True
        )

        # 3. CAPTURA Y VALIDACIÓN INVISIBLE (Límite 1000)
        if calidad_quill is not None:
            st.session_state["input_aseguramiento_calidad"] = calidad_quill
            
            import re
            # Limpieza para el conteo real
            texto_limpio = re.sub('<[^<]+?>', '', str(calidad_quill))
            num_palabras = len(texto_limpio.split())
            
            # Alerta roja solo si se excede el límite
            if num_palabras > 1000:
                st.error(f"⚠️ El texto es demasiado largo ({num_palabras} palabras). El límite para esta sección es de 1000 palabras.")

    # Bloque de apoyo conceptual
    with st.expander("Puntos clave para esta sección"):
        st.markdown("""
        Para una redacción robusta, asegúrese de mencionar:
        * **Autoevaluación:** Periodicidad y actores involucrados.
        * **Planes de Mejoramiento:** Cómo se transforman los hallazgos en acciones concretas.
        """)

    
    generar = st.form_submit_button("GENERAR DOCUMENTO PEP", type="primary")

#  LÓGICA DE GENERACIÓN DEL WORD 
if generar:
    # --- 1. GENERALIDADES DEL PROGRAMA ---
    #datos_prog = BD_PROGRAMAS.get(snies_seleccionado, {})
    denom = str(st.session_state.get("denom_input", "")).strip()
    titulo = str(st.session_state.get("titulo_input", "")).strip()
    snies = str(st.session_state.get("snies_input", "")).strip()
    facultad = str(st.session_state.get("facultad", "FACULTAD NO DEFINIDA")).strip()
    departamento = str(st.session_state.get("departamento", "DEPARTAMENTO NO DEFINIDO")).strip()
    semestres = str(st.session_state.get("semestres_input", "")).strip()
    lugar = str(st.session_state.get("lugar_input", "")).strip()
    creditos = str(st.session_state.get("cred", "")).strip()
    estudiantes = str(st.session_state.get("estudiantes_input", "")).strip()
    acuerdo = str(st.session_state.get("acuerdo_input", "")).strip()
    instancia = str(st.session_state.get("instancia_input", "")).strip()
    periodicidad = str(st.session_state.get("periodicidad_input", "")).strip()
    area = str(st.session_state.get("area", "")).strip()
    nivel = str(st.session_state.get("nivel_formacion_widget", "")).strip()
    modalidad = str(st.session_state.get("modalidad_input", "")).strip()

    # --- 2. REGISTROS CALIFICADOS Y ACREDITACIONES ---
    reg1 = str(st.session_state.get("reg1", "")).strip()
    reg2 = str(st.session_state.get("reg2", "")).strip()
    reg3 = str(st.session_state.get("reg3", "")).strip()
    acred1 = str(st.session_state.get("acred1", "")).strip()
    acred2 = str(st.session_state.get("acred2", "")).strip()
    
    # Cálculo automático del registro calificado más reciente para la tabla resumen
    reg_final = reg3 if reg3 else (reg2 if reg2 else reg1)

    # --- 3. PLANES DE ESTUDIO (Versiones 1, 2 y 3) ---
    # Versión 1 (Actual)
    p1_nom = str(st.session_state.get("p1_nom", "")).strip()
    p1_fec = str(st.session_state.get("p1_fec", "")).strip()
    p1_cred = str(st.session_state.get("p1_cred", "")).strip()
    p1_sem = str(st.session_state.get("p1_sem", "")).strip()
    
    # Versión 2 (Anterior)
    p2_nom = str(st.session_state.get("p2_nom", "")).strip()
    p2_fec = str(st.session_state.get("p2_fec", "")).strip()
    p2_cred = str(st.session_state.get("p2_cred", "")).strip()
    p2_sem = str(st.session_state.get("p2_sem", "")).strip()

    # Versión 3 (Antiguo)
    p3_nom = str(st.session_state.get("p3_nom", "")).strip()
    p3_fec = str(st.session_state.get("p3_fec", "")).strip()
    p3_cred = str(st.session_state.get("p3_cred", "")).strip()
    p3_sem = str(st.session_state.get("p3_sem", "")).strip()

    motivo_final = str(st.session_state.get("motivo_input", "")).strip()
    iti_formativo_final = str(st.session_state.get("input_itinerario", "")).strip()
    entornos_academicos_final = str(st.session_state.get("input_entornos_academicos", "")).strip()
    perfil_docente_final = str(st.session_state.get("input_perfil_docente", "")).strip()
    investigacion_raw = str(st.session_state.get("input_investigacion_general", ""))
    internacional_raw = str(st.session_state.get("input_internacionalizacion", ""))
    bienestar_raw = str(st.session_state.get("input_bienestar", ""))
    comite_raw = str(st.session_state.get("desc_comite_curricular", ""))
    consejo_raw = str(st.session_state.get("desc_consejo_facultad", ""))
    calidad_raw = str(st.session_state.get("input_aseguramiento_calidad", ""))
    

  #  LIMPIEZA DE HTML 
    # Procesamos la variable antes de meterla al diccionario
    iti_formativo_limpio = limpiar_completamente(iti_formativo_final)
    entornos_academicos_limpio = limpiar_completamente(entornos_academicos_final)
    perfil_docente_limpio = limpiar_completamente(perfil_docente_final)
    investigacion_limpio = limpiar_completamente(investigacion_raw)
    internacional_limpio = limpiar_completamente(internacional_raw)
    bienestar_limpio = limpiar_completamente(bienestar_raw)
    comite_limpio = limpiar_completamente(comite_raw)
    consejo_limpio = limpiar_completamente(consejo_raw)
    calidad_limpio = limpiar_completamente(calidad_raw)

    mapa_general_tablas = mapear_todas_las_tablas(archivo_dm)
    
   
    # VALIDACIÓN INICIAL
    if not denom or not reg1:
        st.error("⚠️ Falta información obligatoria (Denominación o Registro Calificado 1).")
    else:      
        ruta_plantilla = "PlantillaPEP.docx" 
        if not os.path.exists(ruta_plantilla):
                st.error(f"❌ No encuentro el archivo '{ruta_plantilla}'.")
        else:
                doc = Document(ruta_plantilla)

                if metodo_trabajo == "Semiautomatizado (Cargar Documento Maestro)":
                    # Recuperamos lo que el usuario eligió en los selectores del formulario
                    seleccion_micro = st.session_state.get("sel_micro")
                    seleccion_macro = st.session_state.get("sel_macro")
    
                    if seleccion_micro:
                        insertar_tabla_seleccionada(doc, "{{certificaciones_micro}}", seleccion_micro)
                    
                    if seleccion_macro:
                        insertar_tabla_seleccionada(doc, "{{certificaciones_macro}}", seleccion_macro) 

                    areas_mapeo = {
                        "{{certificaciones_micro}}": "micro", 
                        "{{certificaciones_micro2}}": "micro_1", 
                        "{{certificaciones_micro3}}": "micro_2", 
                        "{{certificaciones_macro}}": "macro",
                        "{{certificaciones_macro2}}": "macro_1",
                        "{{certificaciones_macro3}}": "macro_2",
                        "{{area_human}}": "formación humanística",
                        # Añadimos "formación básica" como alternativa
                        "{{area_basica}}": "Fundamentación básica | formación básica",
                        "{{area_bp}}": "formación básica profesional",
                        # Unificamos las variantes de electivas
                        "{{area_elec}}": "Cursos electivos | formación flexible o complementaria (electivas) | formación flexible o complementaria",
                        # Unificamos las variantes de profundización/optativas
                        "{{area_prof}}": "Cursos de profundización | formación flexible o complementaria (optativas) | formación de profundización",
                        # Añadimos la variante de específica
                        "{{area_esp}}": "Cursos de fundamentación específica | formación específica",
                        "{{pertinencia_social}}": "objeto de conocimiento perspectivas de intervención",
                        "{{pertinencia_academica}}": "pertinencia academica | cursos academicos"
                    }
                
                    for p_holder, k_word in areas_mapeo.items():
                        exito = insertar_tabla_automatica(doc, p_holder, k_word)
                        if not exito:
                            st.warning(f"⚠️ No se pudo auto-detectar la tabla para: {k_word}")
                        insertar_tabla_automatica(doc, p_holder, k_word)

                        
 
             # 1. CREACIÓN
                texto_base = (
                        f"El Programa de {denom} fue creado mediante el {acuerdo} del {instancia} y aprobado mediante la {reg1} del Ministerio de Educación Nacional con código SNIES {snies}")
                if reg3:
                    texto_historia = f"{texto_base}, posteriormente recibe la renovación del registro calificado a través de la {reg2} y la {reg3}."
                elif reg2:
                    texto_historia = f"{texto_base}, posteriormente recibe la renovación del registro calificado a través de la {reg2}."
                else:
                    texto_historia = f"{texto_base}."
        
                # MOTIVO CREACIÓN
                if motivo_final and motivo_final.strip():
                    parrafo_motivo = motivo_final
                else:
                    parrafo_motivo = motivo_final if motivo_final else "No se suministró información sobre el motivo de creación."

                # MODIFICACIONES CURRICULARES
                intro_planes = (f"El plan de estudios del Programa de {denom} ha sido objeto de procesos periódicos de evaluación,con el fin de asegurar su pertinencia académica y su alineación con los avances tecnológicos y las demandas del entorno. Como resultado")
        
                if p1_nom and p2_nom:
                    # CASO 3 PLANES: Menciona P1 (Viejo) -> P2 (Medio) -> P3 (Actual)
                    parrafo_planes = (f"{intro_planes}se han realizado las modificaciones curriculares al plan {p1_nom} aprobado mediante {p1_fec}, con {p1_cred} créditos y {p1_sem} semestres, posteriormente se actualiza al plan {p2_nom} mediante {p2_fec}, con {p2_cred} créditos y {p2_sem} semestres y por último al plan de estudio vigente {p3_nom} mediante {p3_fec}, con {p3_cred} créditos y {p3_sem} semestres.")
                    
                elif p2_nom: 
                    # CASO 2 PLANES: Asumimos que P2 es el anterior y P3 el actual
                    # (P2 -> P3)
                    parrafo_planes = (f"{intro_planes}se han realizado las modificaciones curriculares al plan {p2_nom} aprobado mediante {p2_fec}, con {p2_cred} créditos y {p2_sem} semestres, posteriormente se actualiza al plan de estudio vigente {p3_nom} mediante {p3_fec}, con {p3_cred} créditos y {p3_sem} semestres.")
        
                elif p1_nom:
                    # CASO ALTERNATIVO 2 PLANES: Solo llenaron P1 (Viejo) y P3 (Actual), saltándose el P2
                    # (P1 -> P3)
                    parrafo_planes = (f"{intro_planes}se han realizado las modificaciones curriculares al plan {p1_nom} aprobado mediante {p1_fec}, con {p1_cred} créditos y {p1_sem} semestres, posteriormente se actualiza al plan de estudio vigente {p3_nom} mediante {p3_fec}, con {p3_cred} créditos y {p3_sem} semestres.")  
                
                else:
                    # CASO 1 PLAN (Solo existe el actual P3)
                    # Preparamos variables por si faltan datos para que no salga vacío
                    nom = p3_nom if p3_nom else "[FALTA NOMBRE PLAN VIGENTE]"
                    fec = p3_fec if p3_fec else "[FALTA FECHA]"
                    
                    parrafo_planes = (
                        f"{intro_planes}se estableció el plan de estudios vigente {nom} "
                        f"aprobado mediante {fec}, con {p3_cred} créditos y {p3_sem} semestres."
                    )
          
                # ACREDITACIÓN
                texto_acred = "" 
                
                acred1 = str(st.session_state.get("acred1", "")).strip()
                acred2 = str(st.session_state.get("acred2", "")).strip()
                
                if acred1 and acred2:
                    # Caso: Dos acreditaciones
                    texto_acred = (
                        f"El programa obtuvo por primera vez la Acreditación en alta calidad otorgada por el "
                        f"Consejo Nacional de Acreditación (CNA) a través de la resolución {acred1}, "
                        f"esta le fue renovada mediante resolución {acred2}, reafirmando la solidez "
                        f"académica, administrativa y de impacto social del Programa."
                    )
                elif acred1:
                     # Caso: Solo una acreditación 
                    texto_acred = (
                        f"El programa obtuvo la Acreditación en alta calidad otorgada por el "
                        f"Consejo Nacional de Acreditación (CNA) a través de la resolución {acred1}, "
                        f"como reconocimiento a su solidez académica, administrativa y de impacto social."
                    )
               
               # RECONOCIMIENTOS
                texto_recons = ""
                recon_data = st.session_state.get("recon_data", [])
                
                # Filtramos los vacíos
                recons_validos = [
                    r for r in recon_data 
                    if isinstance(r, dict) and str(r.get("Nombre del premio", "")).strip()
                ]
                
                if recons_validos:
                    # Encabezado del párrafo de reconocimientos
                    intro_recon = (
                        f"Adicionalmente, el Programa de {denom} ha alcanzado importantes logros académicos e institucionales "
                        f"que evidencian su calidad y compromiso con la excelencia. Entre ellos se destacan:"
                    )
                    lista_items = []        
                    for r in recons_validos:
                        premio = str(r.get("Nombre del premio", "Premio")).strip()
                        anio = str(r.get("Año", "")).strip()
                        ganador = str(r.get("Nombre del Ganador", "")).strip()
                        cargo = str(r.get("Cargo", "")).strip()
                                    
                        item = f"• {premio} ({anio}): Otorgado a {ganador}, en su calidad de {cargo}."
                        lista_items.append(item)
                    
                    texto_recons = intro_recon + "\n" + "\n".join(lista_items)
        
                #LINEA DE TIEMPO
                texto_timeline = ""
                eventos = []
        
                # Función auxiliar para sacar el año (busca 19XX o 20XX en cualquier lado)
                def obtener_anio(texto):
                    if not texto: return 9999 # Si no hay fecha, lo mandamos al final
                    match = re.search(r'\b(19|20)\d{2}\b', str(texto))
                    return int(match.group(0)) if match else 9999
        
                # --- A. Agregamos Resoluciones ---
                if reg1: eventos.append((obtener_anio(reg1), f"Creación y Registro Calificado inicial ({reg1})."))
                if reg2: eventos.append((obtener_anio(reg2), f"Renovación del Registro Calificado ({reg2})."))
                if reg3: eventos.append((obtener_anio(reg3), f"Segunda Renovación Registro Calificado ({reg3})."))
        
                # --- B. Agregamos Planes (P1=Viejo, P2=Medio, P3=Actual) ---
                # Solo agregamos si hay fecha válida
                if p1_fec: eventos.append((obtener_anio(p1_fec), f"Inicio Plan de Estudios {p1_nom}."))
                if p2_fec: eventos.append((obtener_anio(p2_fec), f"Actualización Curricular - Plan {p2_nom}."))
                if p3_fec: eventos.append((obtener_anio(p3_fec), f"Implementación Plan Vigente {p3_nom}."))
        
                # --- C. Agregamos Acreditaciones ---
                if acred1: eventos.append((obtener_anio(acred1), f"Obtención Acreditación de Alta Calidad ({acred1})."))
                if acred2: eventos.append((obtener_anio(acred2), f"Renovación Acreditación de Alta Calidad ({acred2})."))
        
                # --- D. Agregamos Reconocimientos (Solo los destacados) ---
                if recons_validos:
                    for r in recons_validos:
                        anio_r = obtener_anio(r.get("Año", ""))
                        nom_r = r.get("Nombre del premio", "Premio")
                        # Solo agregamos si encontramos un año válido para no ensuciar la línea
                        if anio_r != 9999:
                             eventos.append((anio_r, f"Reconocimiento: {nom_r}."))
        
                # --- E. Ordenar y Construir Texto ---
                # Ordenamos la lista por el año (el primer elemento de la tupla)
                eventos.sort(key=lambda x: x[0])
        
                if eventos:
                    # Creamos un "título" visual en negrita o separado
                    lines = ["Hitos relevantes en la línea de tiempo del programa:"]
                    
                    last_year = 0
                    for anio, desc in eventos:
                        if anio != 9999:
                            lines.append(f"• {anio}: {desc}")
                    
                    texto_timeline = "\n".join(lines)
        
           
                # UNIÓN FINAL E INSERCIÓN
                partes = [
                    texto_historia,  # 1. Creación
                    parrafo_motivo,  # 2. Motivo
                    parrafo_planes,  # 3. Planes
                    texto_acred,     # 4. Acreditación
                    texto_recons,    # 5. Reconocimientos
                    texto_timeline   # 6. Línea de Tiempo (¡Aquí va!)
                ]
                
                # Unimos todo en un solo bloque de texto grande
                texto_final_completo = "\n\n".join([p for p in partes if p and p.strip()])
                
                # DICCIONARIO DE REEMPLAZOS: Definimos los datos que queremos meter en el Word
                mis_reemplazos = {
                    "{{historia_programa}}": texto_final_completo,
                    "{{fundamentacion_epistemologica}}": st.session_state.get("fund_epi_manual", ""),
                    "{{itinerario_formativo}}": iti_formativo_limpio,
                    "{{entornos_academicos}}": entornos_academicos_limpio,
                    "{{perfil_equipo_docente}}": perfil_docente_limpio,
                    "{{descripcion_investigacion}}": investigacion_limpio,
                    "{{estrategias_internacionalizacion}}": internacional_limpio,
                    "{{estrategias_bienestar}}": bienestar_limpio,
                    "{{descripcion_comite_curricular}}": comite_limpio,
                    "{{descripcion_concejo_facultad}}": consejo_limpio,
                    "{{asegu_calidad}}": calidad_limpio,
                    "{{facultad}}": facultad,      
                    "{{departamento}}": departamento,
                }
                
                #st.write(mis_reemplazos)
                reemplazar_en_todo_el_doc(doc, mis_reemplazos)

            
            #IMAGEN: PLAN DE ESTUDIOS
                img_plan = st.session_state.get("upload_plan_estudios")
                if img_plan is not None:
                    reemplazar_etiqueta_por_imagen(
                        doc, 
                        "{{plan_estudios}}", 
                        img_plan, 
                        ancho_pulgadas=6.5 # Ajusta el tamaño según tu margen
                    )

            # IMAGEN: ESTRUCTURA ADMINISTRATIVA 
                img_admin = st.session_state.get("upload_estructura_admin")
                if img_admin:
                    # Usamos el placeholder exacto que indicas
                    exito_admin = insertar_imagen_en_placeholder(
                        doc, 
                        "{{estruc_admin}}", 
                        img_admin,
                        ancho_pulgadas=6.0
                    )
                    if not exito_admin:
                        st.warning("⚠️ No se encontró el placeholder {{estructura_administrativa}} en la sección 11 de la plantilla.")

                # Reemplazos en Portada/Encabezados
                datos_portada = {
                        "{{DENOMINACION}}": denom.upper(),
                        "{{DENOMINACION2}}": denom.title(),
                        "{{SNIES}}": snies,
                        "{{facultad}}": facultad,      
                        "{{departamento}}": departamento,
                }
                reemplazar_en_todo_el_doc(doc, datos_portada)    
                
                reemplazos_especiales = {
                        "{{facultad}}": facultad,
                        "{{departamento}}": departamento
                    }


                for p in doc.paragraphs:
                    for clave, valor in reemplazos_especiales.items():
                        if clave in p.text:
                            # Borramos el placeholder
                            p.text = p.text.replace(clave, "")
                            # Insertamos el texto con el formato deseado
                            run = p.add_run(str(valor))
                            run.bold = True
                            run.font.name = 'Arial'
                            run.font.size = Pt(14)
                            # Aseguramos que el párrafo esté centrado (típico de portadas)
                            p.alignment = 1 
            
            # 3. Hacemos lo mismo por si acaso están dentro de una tabla en la portada
                for tabla in doc.tables:
                    for fila in tabla.rows:
                        for celda in fila.cells:
                            for p in celda.paragraphs:
                                for clave, valor in reemplazos_especiales.items():
                                    if clave in p.text:
                                        p.text = p.text.replace(clave, "")
                                        run = p.add_run(str(valor))
                                        run.bold = True
                                        run.font.name = 'Arial'
                                        run.font.size = Pt(14)
                                        p.alignment = 1


            
                    
                    # Lista de datos para la sección 1.2
                lista_datos = [
                        f"● Denominación del programa: {denom}",
                        f"● Título otorgado: {titulo}",
                        f"● Nivel de formación: {nivel}",
                        f"● Área de formación: {area}",
                        f"● Modalidad de oferta: {modalidad}",
                        f"● Acuerdo de creación: {acuerdo}",
                        f"● Registro calificado: {reg_final}",
                        f"● Créditos académicos: {creditos}",
                        f"● Periodicidad de admisión: {periodicidad}",
                        f"● Lugares de desarrollo: {lugar}",
                        f"● SNIES: {snies}"
                    ]
        
                    # Inserción en el documento
                insertar_lista_bajo_titulo(doc, "Generalidades del programa", lista_datos)   

                if not texto_final_completo.strip():
                    st.warning("La historia está vacía. Verifica que los campos de registros y planes tengan datos.")

                # Paso 2: Ejecutar los reemplazos finales (si no se han hecho)
                reemplazar_en_todo_el_doc(doc, mis_reemplazos)
                reemplazar_en_todo_el_doc(doc, datos_portada)

                if metodo_trabajo == "Semiautomatizado (Cargar Documento Maestro)":
                    # Recuperamos las selecciones que el usuario hizo en los selectboxes
                    seleccion_micro = st.session_state.get("sel_micro")
                    seleccion_macro = st.session_state.get("sel_macro")

                    if seleccion_micro:
                        # Buscamos el placeholder y pegamos la tabla
                        insertar_tabla_seleccionada(doc, "{{certificaciones_micro}}", seleccion_micro)
                    
                    if seleccion_macro:
                        # Buscamos el placeholder y pegamos la tabla
                        insertar_tabla_seleccionada(doc, "{{certificaciones_macro}}", seleccion_macro)
        

        # CAPÍTULO 2: REFERENTES CONCEPTUALES
        #2.1 NATURALEZA DEL PROGRAMA
       
        v_obj_nombre = str(st.session_state.get("obj_nombre_input", "")).strip()
        texto_para_pegar = "" # Contendrá la definición extensa

        if metodo_trabajo == "Semiautomatizado (Cargar Documento Maestro)" and archivo_dm is not None:
            try:
                doc_m = Document(archivo_dm)
                t_inicio = str(st.session_state.get("inicio_def_oc", "")).strip().lower()
                t_fin = str(st.session_state.get("fin_def_oc", "")).strip().lower()
                
                p_extraidos_21 = []
                capturando_21 = False

                for p_m in doc_m.paragraphs:
                    # Usamos el texto original para el recorte final, 
                    # pero una versión limpia para la búsqueda
                    p_text_raw = p_m.text
                    p_text_low = p_text_raw.lower()
                    busqueda_ini = t_inicio.lower()
                    busqueda_fin = t_fin.lower()
                    
                    # CASO: Encontrar el inicio
                    if busqueda_ini in p_text_low and not capturando_21:
                        capturando_21 = True
                        idx_start = p_text_low.find(busqueda_ini)
                        
                        # Verificamos si el final está en este mismo párrafo
                        if busqueda_fin in p_text_low[idx_start + len(busqueda_ini):]:
                            # Si ambos están en el mismo párrafo, cortamos ambos extremos
                            idx_end = p_text_low.find(busqueda_fin, idx_start) + len(busqueda_fin)
                            p_extraidos_21.append(p_text_raw[idx_start:idx_end])
                            capturando_21 = False
                            break
                        else:
                            # Si no está el final, guardamos desde el inicio hasta el final del párrafo
                            p_extraidos_21.append(p_text_raw[idx_start:])
                        continue
                    
                    # CASO: Estamos capturando párrafos intermedios
                    if capturando_21:
                        if busqueda_fin in p_text_low:
                            # Encontramos el cierre: cortamos hasta donde termina el marcador final
                            idx_end = p_text_low.find(busqueda_fin) + len(busqueda_fin)
                            p_extraidos_21.append(p_text_raw[:idx_end])
                            capturando_21 = False
                            break
                        else:
                            # Párrafo intermedio completo
                            p_extraidos_21.append(p_text_raw)

                texto_para_pegar = "\n\n".join(p_extraidos_21)
            except Exception as e:
                st.error(f"Error en la extracción: {e}")

        # 2. INSERCIÓN EN PLACEHOLDERS {{oc}} y {{def_oc}}
        texto_nombre_completo = f"Objeto de conocimiento del programa: {v_obj_nombre}"

        for p_plan in doc.paragraphs:
            # Reemplazo del Nombre del Objeto
            if "{{oc}}" in p_plan.text:
                p_plan.text = p_plan.text.replace("{{oc}}", texto_nombre_completo)
            
            # Reemplazo de la Definición (Estricta)
            if "{{def_oc}}" in p_plan.text:
                if texto_para_pegar:
                    p_plan.text = p_plan.text.replace("{{def_oc}}", texto_para_pegar)
                    p_plan.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                else:
                    p_plan.text = p_plan.text.replace("{{def_oc}}", "")
    
#FUNDAMENTACIÓN EPISTEMOLÓGICA
       # fundamentacion_txt = st.session_state.get("fund_epi_manual", "")
        fundamentacion_txt = str(st.session_state.get("fund_epi_manual", "")).strip()
        if not fundamentacion_txt or fundamentacion_txt == "None":
            fundamentacion_txt = "[ERROR: El texto extraído llegó vacío al generador]"
        marca_epi = "{{fundamentacion_epistemologica}}"

        # 2. Reemplazo en Párrafos de texto libre
        for p in doc.paragraphs:
            if marca_epi in p.text:
                p.text = p.text.replace(marca_epi, str(fundamentacion_txt))
                
                # Le damos formato básico para que no se vea desordenado
                p.alignment = 3  # Justificado
                if p.runs:
                    p.runs[0].font.name = 'Arial' # O la fuente de tu plantilla

    #FUNDAMENTACIÓN ESPECÍFICA
        fund_especifica_txt = st.session_state.get("fund_especifica_txt", "")
        marca_especifica = "{{fundamentación_especifica_programa}}"

        # Reemplazo en Párrafos de texto libre
        for p in doc.paragraphs:
            if marca_especifica in p.text:
                # Limpiamos el párrafo y ponemos el contenido extraído del DM
                p.text = p.text.replace(marca_especifica, st.session_state.get("fund_especifica_txt", ""))
                
                # Formato: Justificado y Fuente Arial
                p.alignment = 3  # WD_ALIGN_PARAGRAPH.JUSTIFY
                if p.runs:
                    p.runs[0].font.name = 'Arial'

    # JUSTIFICACIÓN DEL PROGRAMA
      
        justificacion_nodos = st.session_state.get("justificacion_manual", []) 
        
        marca_justificacion = "{{justificacion_programa}}"
        
        for p in doc.paragraphs:
            if marca_justificacion in p.text:
                # Si encontramos la marca, limpiamos el párrafo original
                p.text = p.text.replace(marca_justificacion, "")
                cursor = p
                
                # Si no hay nada en la lista, ponemos un aviso para saber qué pasó
                if not justificacion_nodos:
                    p.text = "[Aviso: No se extrajo contenido de la Justificación del Maestro]"
                    break
        
                if isinstance(justificacion_nodos, list):
                    for nodo_origen in justificacion_nodos:
                        # Si es un objeto de párrafo (con estilo/runs)
                        if hasattr(nodo_origen, 'runs'):
                            nuevo_p = doc.add_paragraph()
                            cursor._element.addnext(nuevo_p._element)
                            
                            for run_origen in nodo_origen.runs:
                                # Insertamos solo el texto
                                nuevo_run = nuevo_p.add_run(run_origen.text)
                                
                                # --- FORZAMOS TEXTO NORMAL (Sin negritas ni cursivas) ---
                                nuevo_run.bold = False
                                nuevo_run.italic = False
                                
                                # Formato Arial 11
                                nuevo_run.font.name = 'Arial'
                                nuevo_run.font.size = Pt(11)
                            
                            nuevo_p.alignment = 3 # Justificado
                            cursor = nuevo_p
                        else:
                            # Si el nodo es solo una cadena de texto (string)
                            nuevo_p = doc.add_paragraph(str(nodo_origen))
                            cursor._element.addnext(nuevo_p._element)
                            for run in nuevo_p.runs:
                                run.bold = False
                                run.font.name = 'Arial'
                                run.font.size = Pt(11)
                            nuevo_p.alignment = 3
                            cursor = nuevo_p
                else:
                    # Si por alguna razón es un solo string largo
                    p.text = str(justificacion_nodos)
                    for run in p.runs: run.bold = False
                
                break

        perfiles_mapeo = {
            "{{perfil_profesional_experiencia}}": st.session_state.get("perfil_profesional_experiencia_txt", ""),
            "{{perfil_profesional_egresado}}": st.session_state.get("perfil_profesional_egresado_txt", ""),
            "{{perfil_ocupacional}}": st.session_state.get("perfil_ocupacional_txt", "")
        }

        for marca, contenido in perfiles_mapeo.items():
            # Limpieza si es tupla
            texto_limpio = contenido[0] if isinstance(contenido, tuple) else contenido
            
            for p in doc.paragraphs:
                if marca in p.text:
                    p.text = p.text.replace(marca, str(texto_limpio))
                    p.alignment = 3 # Justificado
                    for run in p.runs:
                        run.font.name = 'Arial'

    #RESULTADOS ACADÉMICOS
        resultados_aprendizaje_txt = st.session_state.get("resultados_aprendizaje_txt", "")
        
        # Blindaje por si llega como tupla
        if isinstance(resultados_aprendizaje_txt, tuple):
            resultados_aprendizaje_txt = resultados_aprendizaje_txt[0]

            marca_rapa = "{{resultados_aprendizaje}}"

            for p in doc.paragraphs:
                    if marca_rapa in p.text:
                        # Reemplazo seguro convirtiendo a string
                        p.text = p.text.replace(marca_rapa, str(resultados_aprendizaje_txt))
                        
                        # Formato: Justificado y Arial
                        p.alignment = 3 
                        for run in p.runs:
                            run.font.name = 'Arial'


    #GUARDAR ARCHIVO
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
        
    st.success("✅ ¡Documento PEP generado!")
    st.download_button(
                label="📥 Descargar Documento PEP en Word",
                data=bio.getvalue(),
                file_name=f"PEP_Modulo1_{denom.replace(' ', '_')}.docx",
                 mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                  )
    
