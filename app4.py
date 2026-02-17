import streamlit as st
from google import genai
from docx import Document
from docx.shared import Pt
import requests
import io
import time
import re 
import os
import pandas as pd
from docx.enum.text import WD_ALIGN_PARAGRAPH

st.set_page_config(page_title="Generador Proyecto Educativo", layout="wide")
#LEER DM
def extraer_secciones_dm(archivo_word, mapa_claves):
    """archivo_word: El archivo subido por st.file_uploader. mapa_claves: Un diccionario que dice {'TITULO EN WORD': 'key_de_streamlit'}"""
    doc = Document(archivo_word)
    resultados = {}
    todos_los_parrafos = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
# BUSCAR EL PUNTO DE PARTIDA
    indice_inicio_real = 0
    punto_partida = "BREVE RESEÑA HISTÓRICA DEL PROGRAMA"
    
    for i, texto in enumerate(todos_los_parrafos):
        if punto_partida in texto.upper():
            indice_inicio_real = i
            break # Encontramos el inicio real, dejamos de buscar
            
    parrafos_validos = todos_los_parrafos[indice_inicio_real:]
    
#PROCESO DE EXTRACCIÓN SOBRE LOS PÁRRAFOS VÁLIDOS 
    for titulo_buscado, key_st in mapa_claves.items():
        contenido_seccion = []
        for i, texto in enumerate(parrafos_validos):
            texto_upper = texto.upper()
            
            # Buscamos el título (asegurándonos de que no sea una línea gigante)
            if titulo_buscado.upper() in texto_upper and len(texto) < 120:
                
                for j in range(i + 1, len(parrafos_validos)):
                    siguiente_p = parrafos_validos[j]
                    sig_upper = siguiente_p.upper()
                    
                   # Parar SOLO si encontramos un título principal (Ej: 3. o 4.)
                    # Bajamos el límite a 60 caracteres para no confundir párrafos con títulos
                    es_nuevo_capitulo = re.match(r'^\d+[\.\s]', siguiente_p.strip())
                    es_otro_titulo_mapa = any(t.upper() == sig_upper for t in mapa_claves.keys())

                    if (es_nuevo_capitulo or es_otro_titulo_mapa) and len(siguiente_p) < 60:
                        break
                        
                    contenido_seccion.append(siguiente_p)
                
                # Guardamos TODO el texto en una variable "secreta" para el Word final
                texto_completo = "\n\n".join(contenido_seccion).strip()
                st.session_state[f"full_{key_st}"] = texto_completo
                
                # Preparamos la VISTA PREVIA para el cuadro de texto
                parrafos_lista = texto_completo.split("\n\n")
                if len(parrafos_lista) > 2:
                    # Mostramos primer párrafo + aviso + último párrafo
                    resumen = f"{parrafos_lista[0]}\n\n[... {len(parrafos_lista)-2} PÁRRAFOS INTERMEDIOS CARGADOS TOTALMENTE ...]\n\n{parrafos_lista[-1]}"
                    resultados[key_st] = resumen
                else:
                    resultados[key_st] = texto_completo
                
                break

#CARGAR BD
@st.cache_data # Esto hace que el Excel se lea una sola vez y no cada que muevas un botón
def cargar_base_datos():
    try:
        # Puedes usar pd.read_csv("programas.csv") si prefieres CSV
        df = pd.read_excel("Programas.xlsx", dtype={'snies_input': str}) 
        # Convertimos el DataFrame en un diccionario donde la llave es el SNIES
        return df.set_index("snies_input").to_dict('index')
    except Exception as e:
        st.warning(f"No se pudo cargar la base de datos de Excel: {e}")
        return {}

#CARGA DE DATOS INICIAL
BD_PROGRAMAS = cargar_base_datos()

#2. MAPEO Y ESTRUCTURA (DICCIONARIO)
# Mapeo de: "Título exacto en el DM" -> "Key en App Streamlit"
MAPA_EXTRACCION = {
    "OBJETO DE CONOCIMIENTO": "obj_nombre_input",
    "JUSTIFICACIÓN": "justificacion_input",
    "Conceptualización teórica y epistemológica del programa": "input_epi_p1",
    "Mecanismos de evaluación": "input_mec_p1",
    "IDENTIDAD DISCIPLINAR": "input_epi_p2",
    "ITINERARIO FORMATIVO": "input_itinerario",
    "Justificación del Programa": "input_just",
    "JUSTIFICACIÓN DEL PROGRAMA": "input_just",
    "FUNDAMENTACIÓN ACADÉMICA": "input_acad"

}

# ESRUCTURA PARA EXTRACCIÓN GUIADA (Cap2)
if "config_cap2" not in st.session_state:
    st.session_state.config_cap2 = [
                    {
                        "id": "concOC_input", 
                        "nombre": "2.1 Conceptualización del objeto de conocimiento del Programa", 
                        "inicio": " ", 
                        "fin": " "
                    },
                    {
                        "id": "input_epi_p1", 
                        "nombre": "2.2. Fundamentación Epistemológica", 
                        "inicio": "FUNDAMENTACIÓN EPISTEMOLÓGICA", 
                        "fin": "ESTADO DE LA OCUPACIÓN"
                    },
                    {
                        "id": "input_acad", 
                        "nombre": "2.3. Fundamentación Académica (Certificaciones Temáticas Tempranas", 
                        "inicio": "FUNDAMENTACIÓN ACADÉMICA", 
                        "fin": "CERTIFICACIONES TEMATICAS"
                    },
              ]    
                     # --- Definición de la estructura Capítulo 4 ---
if "config_cap4" not in st.session_state:
    st.session_state.config_cap4 = [
                    {
                        "id": "input_justificacion", 
                        "nombre": "4.1. Justificación del Programa", 
                        "inicio": "JUSTIFICACIÓN", 
                        "fin": "OBJETIVOS" # O la sección que siga en tu documento
                    }
                 
                                ]

#  CONFIGURACIÓN DE PÁGINA
st.title("Generador PEP - Módulo 1: Información del Programa")
st.markdown("""
Esta herramienta permite generar el PEP de dos formas:
1. **Manual:** Completa los campos en las secciones de abajo.
2. **Automatizada:** Sube el Documento Maestro (DM) y el sistema pre-llenará algunos campos.
""")

# SELECTOR DE MODALIAD
metodo_trabajo = st.radio(
    "Selecciona cómo deseas trabajar hoy:",
    ["Manual (Desde cero)", "Automatizado (Cargar Documento Maestro)"],
    horizontal=True
)

# LÓGICA DE CARGA
if metodo_trabajo == "Automatizado (Cargar Documento Maestro)":
    st.subheader("2. Carga de Documento Maestro")
    archivo_dm = st.file_uploader("Sube el archivo .docx del Documento Maestro", type=["docx"])
        
    if archivo_dm:
        # Pestañas para elegir el tipo de automatización
        tab_auto, tab_guiado = st.tabs([
            "Automatizado (Extracción)", 
            "Guiado (Definir Inicio/Fin)"
        ])
        
        #PESTAÑA 1: AUTOMÁTICO 
        with tab_auto:
            st.info("El sistema buscará títulos estándar (ej: 'JUSTIFICACIÓN', 'MISIÓN') y extraerá el contenido automáticamente.")
            
            # Usamos un key único para evitar conflictos
            if st.button("Procesar y Pre-llenar Todo", key="btn_procesar_auto"):
                with st.spinner("Analizando la estructura del documento..."):
                    try:
                        # 1. Llamamos a la función que definimos arriba (Sección 3)
                        datos_capturados = extraer_secciones_dm(archivo_dm, MAPA_EXTRACCION)   
                        
                        # 2. Guardamos los resultados en la memoria (Session State)
                        contador = 0
                        for key, valor in datos_capturados.items():
                            if valor: # Solo guardamos si encontró algo
                                st.session_state[key] = valor
                                contador += 1
                        
                        # 3. Feedback y Recarga
                        if contador > 0:
                            st.success(f"✅ Éxito: Se extrajeron {contador} secciones correctamente.")
                            st.rerun() # Recarga la página para mostrar los datos en el formulario de abajo
                        else:
                            st.warning("⚠️ No se encontraron coincidencias exactas con los títulos estándar.")
                            
                    except Exception as e:
                        st.error(f"Ocurrió un error al procesar el archivo: {e}")

        # PESTAÑA 2: GUIADO
        with tab_guiado:
            st.info("Configura las frases de inicio y fin para ambos capítulos y luego ejecuta la extracción masiva.")
            
            # Verificamos que existan las configuraciones en memoria
            if "config_cap2" in st.session_state and "config_cap4" in st.session_state:
                
                # --- BLOQUE VISUAL 1: CAPÍTULO 2 ---
                st.markdown("#### 📘 Capítulo 2: Referentes Conceptuales")
                st.caption("Define los límites para: Objeto, Epistemología y Fundamentación Académica.")
                
                for i, item in enumerate(st.session_state.config_cap2):
                    with st.expander(f"Configurar: {item['nombre']}", expanded=False):
                        c1, c2 = st.columns(2)
                        # Nota: Usamos keys únicos (g2_...)
                        item["inicio"] = c1.text_input("Empieza con la frase...", value=item["inicio"], key=f"g2_start_{i}")
                        item["fin"] = c2.text_input("Termina antes de...", value=item["fin"], key=f"g2_end_{i}")

                st.markdown("---") # Separador visual

                # --- BLOQUE VISUAL 2: CAPÍTULO 4 ---
                st.markdown("#### 📙 Capítulo 4: Justificación")
                st.caption("Define los límites para la Justificación del programa.")

                for i, item in enumerate(st.session_state.config_cap4):
                    with st.expander(f"Configurar: {item['nombre']}", expanded=False):
                        c1, c2 = st.columns(2)
                        # Nota: Usamos keys únicos (g4_...)
                        item["inicio"] = c1.text_input("Empieza con la frase...", value=item["inicio"], key=f"g4_start_{i}")
                        item["fin"] = c2.text_input("Termina antes de...", value=item["fin"], key=f"g4_end_{i}")

                st.markdown("---")

                # --- EL ÚNICO BOTÓN DE EJECUCIÓN ---
                if st.button("Ejecutar Extracción Guiada", key="btn_guiado_total", type="primary"):
                    with st.spinner("Leyendo documento y extrayendo secciones..."):
                        try:
                            # 1. Rebobinamos el archivo (CRÍTICO)
                            archivo_dm.seek(0)
                            doc_obj = Document(archivo_dm)
                            
                            # 2. Unimos ambas configuraciones en una sola lista de tareas
                            plan_completo = st.session_state.config_cap2 + st.session_state.config_cap4
                            
                            exitos = 0
                            
                            # 3. Iteramos sobre cada configuración
                            for item in plan_completo:
                                contenido = []
                                capturando = False
                                # Limpiamos espacios y mayúsculas para comparar mejor
                                marcador_inicio = item["inicio"].strip().lower()
                                marcador_fin = item["fin"].strip().lower()
                                
                                # Si el usuario dejó algo vacío, saltamos esa sección
                                if not marcador_inicio or not marcador_fin:
                                    continue
                                
                                # Barrido del documento
                                for para in doc_obj.paragraphs:
                                    texto_limpio = para.text.strip().lower()
                                    if not texto_limpio: continue
                                    
                                    # Detectar inicio
                                    if marcador_inicio in texto_limpio and not capturando:
                                        capturando = True
                                        continue # Saltamos el título mismo
                                    
                                    # Detectar fin
                                    if marcador_fin in texto_limpio and capturando:
                                        capturando = False
                                        break # Salimos del bucle de párrafos para esta sección
                                    
                                    # Guardar contenido
                                    if capturando:
                                        contenido.append(para.text)
                                
                                # Si encontramos algo, lo guardamos en Session State
                                if contenido:
                                    texto_final = "\n\n".join(contenido)
                                    st.session_state[item["id"]] = texto_final
                                    st.session_state[f"full_{item['id']}"] = texto_final # Respaldo
                                    exitos += 1
                            
                            # 4. Resultado final
                            if exitos > 0:
                                st.success(f"✅ ¡Éxito! Se extrajeron {exitos} secciones y están listas en el formulario de abajo.")
                                st.rerun() # Recarga para ver los datos abajo
                            else:
                                st.error("❌ No se pudo extraer nada. Verifica que las frases de inicio y fin estén escritas EXACTAMENTE igual (tildes, espacios) que en el Word.")

                        except Exception as e:
                            st.error(f"Error técnico leyendo el archivo: {e}")
            
            else:
                st.error("⚠️ Error interno: No se cargó la configuración inicial (config_cap2/4). Revisa la Sección 4 de tu código.")
        




#3. DICCIONARIO / ESTRUCTURA
# Agregamos 'key_dm' para que el extractor sepa qué título buscar en el Word
estructura_pep = {
    "1. Información del Programa": {
        "1.1. Historia del Programa": {"tipo": "especial_historia"},
        "1.2. Generalidades del Programa": {"tipo": "directo"}
    },
    "2. Referentes Conceptuales": {
        "2.1. Naturaleza del Programa": {
            "tipo": "directo",
            "key_dm": "OBJETO DE CONOCIMIENTO", # Palabra clave para buscar en el DM
            "campos": [
                {
                    "label": "Objeto de conocimiento del Programa", 
                    "req": True, 
                    "key": "obj_nombre_input",
                    "help": "¿Qué conoce, investiga y transforma este programa?"
                }
            ]
        },
        "2.2. Fundamentación epistemológica": {
            "tipo": "directo",
            "key_dm": "FUNDAMENTACIÓN EPISTEMOLÓGICA",
            "campos": [
                {"label": "Naturaleza epistemológica e identidad académica", "req": True, "key": "input_epi_p1"},
                {"label": "Campo del saber y relación con ciencia/tecnología", "req": True, "key": "input_epi_p2"}
            ]
        },
        "2.3. Fundamentación académica": {
            "tipo": "especial_pascual", 
            "campos": [] 
        }
    }
}


st.markdown("---")


# LÓGICA DE MODALIDAD

with st.expander("Buscador Información general del Programa por SNIES", expanded=True):
    st.subheader("1. Búsqueda del Programa por SNIES")
    
    col_busq, col_btn = st.columns([3, 1])
    
    with col_busq:
        snies_a_buscar = st.text_input("Ingresa el código SNIES:", placeholder="Ej: 102345", key="search_snies_tmp")
        
    with col_btn:
        st.write(" ")
        st.write(" ")
        if st.button("🔍 Consultar Base de Datos"):
            if snies_a_buscar in BD_PROGRAMAS:
                datos_encontrados = BD_PROGRAMAS[snies_a_buscar]

                # 1. Borramos las llaves viejas para que el formulario no se bloquee
                llaves_a_limpiar = ["denom_input", "titulo_input", "snies_input", "acuerdo_input", "instancia_input", "reg1", "Creditos", "periodo_idx", "acred1", "lugar"
]
                for k in llaves_a_limpiar:
                    if k in st.session_state:
                        del st.session_state[k]
                
                # 2. Inyectamos los nuevos datos del Excel
                for key, valor in datos_encontrados.items():
                    st.session_state[key] = valor
                
                # 3. Guardamos el SNIES que acabamos de buscar
                st.session_state["snies_input"] = snies_a_buscar
                
                st.success(f"✅ Programa encontrado: {datos_encontrados.get('denom_input')}")
                st.rerun()
            else:
                st.error("❌ Código SNIES no registrado en el sistema.")

    st.markdown("---")

# --- FORMULARIO DE ENTRADA ---
with st.form("pep_form"):
    # 1. Recuperamos datos de ejemplo si existen
    ej = st.session_state.get("ejemplo", {})

    st.markdown("### 1. Identificación General")
    col1, col2 = st.columns(2)
    
    with col1:
        # Denominación del programa
        denom = st.text_input(
            "Denominación del programa :red[•]", 
            value=st.session_state.get("denom_input", ej.get("denom_input", "")),
            key="denom_input"
        )

        # Título otorgado (Ahora bien indentado dentro de col1)
        titulo = st.text_input(
            "Título otorgado :red[•]", 
            value=st.session_state.get("titulo_input", ej.get("titulo_input", "")),
            key="titulo_input"
        )
    
    # Nivel de formación (Protección contra errores de índice)
    niveles_opciones = ["Técnico", "Tecnológico", "Profesional universitario", "Especialización", "Maestría", "Doctorado"]
    
    # Intentamos obtener el valor del extractor o del ejemplo
    val_nivel = st.session_state.get("nivel_idx", st.session_state.get("ejemplo", {}).get("nivel_idx", 2))
    
    # Aseguramos que sea un número para el selectbox
    try:
        idx_final = int(val_nivel)
    except (ValueError, TypeError):
        idx_final = 2 # Por defecto Profesional
    
    nivel = st.selectbox(
        "Nivel de formación :red[•]", 
        options=niveles_opciones, 
        index=idx_final,
        key="nivel_formacion_widget"
    )
    with col2:
        idx_mod = st.session_state.get("modalidad_idx", 0)
        modalidad = st.selectbox(
            "Modalidad de oferta :red[•]", 
            ["Presencial", "Virtual", "A Distancia", "Dual", "Presencial y Virtual", "Presencial y a Distancia", "Presencial y Dual"],
            index=int(idx_mod) if isinstance(idx_mod, (int, float)) else 0,
            key="modalidad_input"
        )
        
        acuerdo = st.text_input(
            "Acuerdo de creación / Norma interna :red[•]", 
            key="acuerdo_input"
        )

        # Instancia interna
        instancia = st.text_input(
            "Instancia interna que aprueba :red[•]", 
            key="instancia_input"
        )

        # Código SNIES
        snies = st.text_input(
            "Código SNIES", 
            key="snies_input"
        )

    st.markdown("---")
    st.markdown("### 2. Registros y Acreditaciones")
    col3, col4 = st.columns(2)
    with col3:
        reg1 = st.text_input(
            label="Resolución Registro calificado 1 :red[•]", 
            value=st.session_state.get("reg1", ej.get("reg1", "")), 
            placeholder="Ej: Resolución 12345 de 2023",
            key="reg1"
        )
        reg2 = st.text_input("Registro calificado 2 (Opcional)", value=ej.get("reg2", ""))
        acred1 = st.text_input(
            label="Resolución Acreditación en alta calidad 1 (Opcional)", 
            value=st.session_state.get("acred1", ej.get("acred1", "")),
            placeholder="Ej: Resolución 012345 de 2022 (Dejar vacío si no aplica)",
            key="acred1"
        )
        acred2 = st.text_input("Resolución Acreditación en alta calidad 2 (Opcional)", value="")

    with col4:
        st.text_input(
            "Créditos Académicos :red[•]",
            value=str(st.session_state.get("Creditos", ej.get("Creditos", ""))),
            placeholder="Ej: 160",
            key="creditos"
        )
        periodicidad = st.selectbox("Periodicidad de admisión :red[•]", ["Semestral", "Anual"], index=ej.get("periodo_idx", 0))
        
        st.text_input(
            "Lugares de desarrollo :red[•]",
            value=st.session_state.get("lugar", ej.get("lugar", "")),
            placeholder="Ej: Medellín, Bogotá, Virtual",
            key="lugar"
        )

    frase_auto = f"La creación del Programa {denom} se fundamenta en la necesidad de "
    val_motivo = ej.get("motivo", frase_auto)
    motivo = st.text_area("Motivo de creación :red[•]", value=val_motivo, height=150)
      
    st.markdown("---")
    st.markdown("### 3. Modificaciones al Plan de Estudios")
    p_col1, p_col2, p_col3 = st.columns(3)
    with p_col1:
        p1_nom = st.text_input("Nombre Plan v1:red[•]", value=ej.get("p1_nom", ""))
        p1_fec = st.text_input("Acuerdo aprobación Plan v1 :red[•]", value=ej.get("p1_fec", ""))
    with p_col2:
        p2_nom = st.text_input("Nombre Plan v2 (Opcional)", value=ej.get("p2_nom", ""))
        p2_fec = st.text_input("Acuerdo aprobación Plan v2 (Opcional)", value=ej.get("p2_fec", ""))
    with p_col3:
        p3_nom = st.text_input("Nombre Plan v3 (Opcional)", value=ej.get("p3_nom", ""))
        p3_fec = st.text_input("Acuerdo aprobación Plan v3 (Opcional)", value=ej.get("p3_fec", ""))

    st.markdown("---")
    st.markdown("### 🏆 4. Reconocimientos (Opcional)")
    recon_data = st.data_editor(
        ej.get("recon_data", [{"Año": "", "Nombre del premio": "", "Nombre del Ganador": "", "Cargo": "Estudiante"}]),
        num_rows="dynamic",
        key="editor_recon", # Es vital tener una key única
        column_config={
            "Cargo": st.column_config.SelectboxColumn(options=["Docente", "Líder", "Decano", "Estudiante,Docente Investigador, Investigador"])
        },
        use_container_width=True
        )  
        





# BOTÓN DE DATOS DE EJEMPLO
    if st.button("Llenar con datos de ejemplo"):
        for k in ["denom_input", "titulo_input", "snies_input"]:
            if k in st.session_state:
                del st.session_state[k]
        st.session_state.ejemplo = {
        "denom_input": "Ingeniería de Sistemas",
        "titulo_input": "Ingeniero de Sistemas",
        "nivel_idx": 2, # Profesional universitario
        "area_input": "Ingeniería, Arquitectura y Urbanismo",
        "modalidad_input": 4, # Presencial y Virtual
        "acuerdo_input:": "Acuerdo 012 de 2015",
        "instancia_input": "Consejo Académico",
        "reg1": "Res. 4567 de 2016",
        "reg2": "Res. 8901 de 2023",
        "acred1": "Res. 00234 de 2024",
        "creditos": "165",
        "periodo_idx": 0, # Semestral
        "lugar": "Sede Principal (Cali)",
        "snies": "54321",
        "motivo": "La creación del Programa se fundamenta en la necesidad de formar profesionales capaces de liderar la transformación digital, diseñar y desarrollar soluciones de software de alta complejidad, gestionar sistemas de información y responder de manera innovadora a los retos tecnológicos, organizacionales y sociales del entorno local, nacional e internacional.",
        "p1_nom": "EO1", "p1_fec": "Acuerdo 012-2015",
        "p2_nom": "EO2", "p2_fec": "Acuerdo 088-2020",
        "p3_nom": "EO3", "p3_fec": "Acuerdo 102-2024",
        #DATOS CAPÍTULO 2
        "objeto_nombre": "Sistemas de información",
        "objeto_concep": "Los sistemas de información son conjuntos organizados de personas, datos, procesos, tecnologías y recursos que interactúan de manera integrada para capturar, almacenar, procesar, analizar y distribuir información, con el fin de apoyar la toma de decisiones, la gestión operativa, el control organizacional y la generación de conocimiento. Estos sistemas permiten transformar los datos en información útil y oportuna, facilitando la eficiencia, la innovación y la competitividad en organizaciones de distintos sectores. Su diseño y gestión consideran aspectos técnicos, organizacionales y humanos, garantizando la calidad, seguridad, disponibilidad y uso ético de la información.",        
        "fund_epi": "El programa se inscribe en el racionalismo crítico y el pragmatismo tecnológico, vinculando la ciencia de la computación con la ingeniería aplicada.",
        # DATOS PARA LAS TABLAS (Se guardan como listas de diccionarios)
        "recon_data": [
            {"Año": "2024", "Nombre del premio": "Excelencia Académica", "Nombre del Ganador": "Juan Pérez", "Cargo": "Docente"}
        ],
        "tabla_cert_ej": [
            {"Nombre": "Desarrollador Web Junior", "Curso 1": "Programación I", "Créditos 1": 3, "Curso 2": "Bases de Datos", "Créditos 2": 4},
            {"Nombre": "Analista de Datos", "Curso 1": "Estadística", "Créditos 1": 4, "Curso 2": "Python para Ciencia", "Créditos 2": 4}
        ], #         
        "referencias_data": [
            {
                "Año": "2021", 
                "Autor(es)": "Sommerville, I.", 
                "Revista": "Computer science", 
                "Título del artículo/Libro": "Engineering Software Products"
            },
            {
                "Año": "2023", 
                "Autor(es)": "Pressman, R. & Maxim, B.", 
                "Revista": "Software Engineering Journal", 
                "Título del artículo/Libro": "A Practitioner's Approach"
            }
        ],
    }

    
    st.rerun()





    st.markdown("---")
    st.markdown("#### CAPÍTULO 5. Estructura curricular")
    st.info("5.1. Pertinencia Social. Complete los campos basándose en la tabla de Estructura Curricular del diseño del programa.")

# Fila 1: Objeto de Conocimiento y Sector Productivo
    col1, col2 = st.columns(2)

    with col1:
        st.subheader("5.1.1 Objeto de Conocimiento")
        st.text_area(
            "Describa el Objeto de Conocimiento",
            key="input_objeto_conocimiento",
            height=200,
            help="Defina el campo del saber."
    )

    with col2:
        st.subheader("5.1.2. Sector Social/Productivo")
        st.text_area(
            "Contexto del Sector",
            key="input_sector_productivo",
            height=200,
            help="Sectores donde impacta el programa."
    )

# Fila 2: Objeto de Formación y Competencias
    col3, col4 = st.columns(2)

    with col3:
        st.subheader("5.1.3. Objeto de Formación")
        st.text_area(
        "Perspectivas de intervención",
        key="input_objeto_formacion",
        height=200,
        help="Intención formativa."
    )

    with col4:
        st.subheader("5.1.4. Competencias de Desempeño Profesional")
        competencia_compartida = st.text_area(
        "Competencias de Desempeño",
        key="input_comp_social", # Esta es la llave principal
        height=200
    )

    st.markdown("---")
    st.markdown("#### CAPÍTULO 5. Estructura curricular")
    st.info("5.2. Pertinencia Académica.")

# Fila 1: Competencia de desempeño y Areas de formación
    col1, col2 = st.columns(2)

    with col1:
        st.subheader("5.2.1 Competencia de desempeño profesional")
        if st.session_state.get("input_comp_social"):
            st.success("✅ Texto copiado de 5.1.4. Competencias de Desempeño Profesional:")
            st.markdown(f"> {st.session_state.input_comp_social}")
        else:
            st.warning("⚠️ Primero completa la sección 5.1.4")

    with col2:
        st.subheader("5.2.2. Áreas de formación")
        st.text_area(
        "ÁREAS",
        key="input_areas",
        height=200,
        help="Áreas de formación del programa."
    )

# Fila 2: Cursos y RA
    col3, col4 = st.columns(2)

    with col3:
        st.subheader("5.2.3. Cursos")
        st.text_area(
        "Cursos del Programa",
        key="input_cursos",
        height=200,
        help="Cursos."
    )

    with col4:
        st.subheader("5.2.4. Resultados de Aprendizaje")
        st.text_area(
        "Resultados de Aprendizaje",
        key="input_ra",
        height=200,
        help="RA."
    )

    st.markdown("---")
    st.markdown("### 5.3. Plan de Estudios")
    st.info("Cargue la imagen del plan de estudios del Programa.")

# Widget para subir la imagen
    archivo_plan = st.file_uploader(
        "Seleccione la imagen del Plan de Estudios", 
        type=["png", "jpg", "jpeg"],
        key="uploader_plan_estudios"
    )

# Mostrar vista previa si el archivo existe
    if archivo_plan is not None:
        st.image(archivo_plan, caption="Vista previa del Plan de Estudios", use_container_width=True)
        # Guardamos el contenido en el session_state para el generador de Word
        st.session_state["imagen_plan"] = archivo_plan
    
    st.markdown("---")
    st.markdown("### 5.4 Perfiles del Programa")
    
        # Perfil Profesional con Experiencia
    perfil_exp = st.text_area(
            "Perfil Profesional con Experiencia :red[•]",
            value=ej.get("perfil_exp", ""),
            placeholder="Describa la trayectoria y experiencia esperada...",
            height=150,
            key="perfil_exp_input"
        )
    
        # Perfil Profesional del Egresado
    perfil_egresado = st.text_area(
            "Perfil Profesional del Egresado (Competencias) :red[•]",
            value=ej.get("perfil_egresado", ""),
            placeholder="Describa las capacidades académicas y profesionales del egresado...",
            height=150,
            key="perfil_egresado_input"
        )
    
        # Perfil Ocupacional
    perfil_ocupacional = st.text_area(
            "Perfil Ocupacional (Campos de acción) :red[•]",
            value=ej.get("perfil_ocupacional", ""),
            placeholder="Describa los cargos y sectores donde podrá desempeñarse...",
            height=150,
            key="perfil_ocupacional_input"
        )


    # --- 2. SECCIÓN MANUAL (Línea 689) ---
    # Aquí aplicamos la condición: SOLO se muestra si NO elegimos el modo automatizado
    if metodo_trabajo != "Automatizado (Cargar Documento Maestro)":
        #CAPÍTULO 2
        st.markdown("---")
        st.header("2. Referentes Conceptuales")
        # 2. Objeto de conocimiento del Programa
        val_obj_nombre = ej.get("objeto_nombre", "")
        objeto_nombre = st.text_input(
            "1. ¿Cuál es el Objeto de conocimiento del Programa? :red[•]",
                 placeholder="Ejemplo: Sistemas de información",
                 key="obj_nombre_input"
        )
          # 2.1 Conceptualización 
            #val_obj_concep = ej.get("objeto_concep", "")
        objeto_conceptualizacion = st.text_area(
                "2. Conceptualización del objeto de conocimiento del Programa :red[•]",
               # value=val_obj_concep, 
                height=150, 
                key="obj_concep_input", 
                placeholder="Ejemplo: Los sistemas de información son conjuntos organizados de personas, datos, procesos, tecnologías y recursos que interactúan de manera integrada para capturar, almacenar, procesar, analizar y distribuir información, con el fin de apoyar la toma de decisiones, la gestión operativa, el control organizacional y la generación de conocimiento. Estos sistemas permiten transformar los datos en información útil y oportuna, facilitando la eficiencia, la innovación y la competitividad en organizaciones de distintos sectores. Su diseño y gestión consideran aspectos técnicos, organizacionales y humanos, garantizando la calidad, seguridad, disponibilidad y uso ético de la información."
            )
         #2.2 
        st.write("Referencias bibliográficas que sustentan la conceptualización del Objeto de Conocimiento.")
        referencias_previa = ej.get("referencias_data", [
                {"Año": "", "Autor(es) separados por coma": "", "Revista": "", "Título del artículo/Libro": ""}
            ])
        
        referencias_data = st.data_editor(
                referencias_previa,
                num_rows="dynamic", # Permite al usuario agregar/borrar filas con el signo +
                key="editor_referencias",
                use_container_width=True,
                column_config={
                    "Año": st.column_config.TextColumn("Año", width="small"),
                    "Autor(es)": st.column_config.TextColumn("Autor(es)", width="medium"),
                    "Revista": st.column_config.TextColumn("Revista", width="medium"),
                    "Título del artículo/Libro": st.column_config.TextColumn("Título del artículo/Libro", width="large"),
                }
            )
    
      # 2.2. Fundamentación epistemológica en Pestañas ---
    st.markdown("---")
    st.subheader("2.2. Fundamentación epistemológica")
    st.info("Utilice las pestañas para completar los tres párrafos de la Fundamentación epistemológica.")
    
    # 1. Creamos las pestañas
    tab1, tab2, tab3 = st.tabs(["Párrafo 1", "Párrafo 2", "Párrafo 3"])
    
    # Configuración de columnas 
    config_columnas_ref = {
            "Año": st.column_config.TextColumn("Año", width="small"),
            "Autor(es) separados por coma": st.column_config.TextColumn("Autor(es)", width="medium"),
            "Revista": st.column_config.TextColumn("Revista", width="medium"),
            "Título del artículo/Libro": st.column_config.TextColumn("Título del artículo/Libro", width="large"),
        }
    
    # Bloque Párrafo 1
    with tab1:
            st.markdown("### Párrafo 1: Marco filósofico")
            st.text_area(
                "¿Cuál es la postura filosófica predominante (positivismo, constructivismo, teoría crítica, complejidad)?:red[•]",
                value=ej.get("fund_epi_p1", ""), 
                height=200,
                key="input_epi_p1",
                placeholder="""Ejemplo: El programa se fundamenta en el paradigma de la complejidad y la visión sistémica, donde la realidad no se percibe como un conjunto de elementos aislados, sino como una red de interacciones y procesos emergentes. Bajo esta postura, el conocimiento se valida a través de la capacidad de modelar y abstraer sistemas reales para transformarlos en sistemas de información lógico-formales que sean verificables y funcionales. Así, la "verdad" en esta disciplina no reside únicamente en el componente físico (el hardware) o en el código (el software), sino en la coherencia del flujo de información y en la eficacia del sistema para resolver problemas de organización, entropía y control en entornos dinámicos y globales..""",
            )
            st.write("Referencias bibliográficas (Párrafo 1):")
            st.data_editor(
                ej.get("referencias_epi_p1", [{"Año": "", "Autor(es) separados por coma": "", "Revista": "", "Título del artículo/Libro": ""}]),
                num_rows="dynamic",
                key="editor_refs_p1",
                use_container_width=True,
                column_config=config_columnas_ref
            )
    
    # Bloque Párrafo 2
    with tab2:
            st.markdown("### Párrafo 2: Identidad disciplinar")
            st.text_area(
                "Origen etimológico y teórias conceptuales que sustentan el Programa:red[•]",
                value=ej.get("fund_epi_p2", ""), 
                height=200,
                key="input_epi_p2",
                placeholder="""Ejemplo: La identidad de este programa se define desde la convergencia etimológica de la ingeniería —del latín ingenium, que refiere a la capacidad natural de invención y resolución de problemas— y el concepto de sistema —del griego systema, entendido como la unión de partes que forman un todo organizado—. Esta génesis conceptual establece que su objeto de estudio no es la máquina en sí misma, sino la arquitectura de procesos y la gestión de la complejidad mediante el uso de la tecnología. Sustentado en la Teoría General de Sistemas y la Cibernética, el programa se deslinda de las ingenierías tradicionales al centrarse en lo intangible —la información y la estructura—, permitiendo que el profesional no solo diseñe herramientas digitales, sino que sea capaz de integrar elementos humanos, tecnológicos y organizacionales en soluciones holísticas y escalables.""",
           )
            st.write("Referencias bibliográficas (Párrafo 2):")
            st.data_editor(
                ej.get("referencias_epi_p2", [{"Año": "", "Autor(es) separados por coma": "", "Revista": "", "Título del artículo/Libro": ""}]),
                num_rows="dynamic",
                key="editor_refs_p2",
                use_container_width=True,
                column_config=config_columnas_ref
            )
    
    # Bloque Párrafo 3
    with tab3:
            st.markdown("### Párrafo 3: Intencionalidad social")
            st.text_area(
                "¿De qué manera la forma en que se produce el conocimiento en este programa garantiza una intervención ética y transformadora en el entorno profesional?:red[•]",
                value=ej.get("fund_epi_p3", ""), 
                height=200,
                key="input_epi_p3",
                placeholder="""Ejemplo: Finalmente, la producción de conocimiento en este programa se orienta hacia una praxis ética y socialmente responsable, donde la tecnología se entiende como un medio para el desarrollo humano y no como un fin deshumanizante. La intervención del ingeniero de sistemas trasciende la ejecución técnica para convertirse en una labor de transformación digital con conciencia crítica, garantizando la seguridad, la privacidad y la integridad de los datos en una sociedad cada vez más automatizada. Este compromiso teleológico asegura que el profesional no solo responda a las demandas del mercado, sino que actúe como un gestor del cambio capaz de diseñar soluciones sostenibles que reduzcan las brechas tecnológicas y promuevan la eficiencia organizacional bajo principios de transparencia y justicia social.""",
    
            )
            st.write("Referencias bibliográficas (Párrafo 3):")
            st.data_editor(
               ej.get("referencias_epi_p3", [{"Año": "", "Autor(es) separados por coma": "", "Revista": "", "Título del artículo/Libro": ""}]),
                num_rows="dynamic",
                key="editor_refs_p3",
                use_container_width=True,
                column_config=config_columnas_ref
            )

   # 2.3. Fundamentación Académica 
    st.markdown("---")
    st.subheader("2.3. Fundamentación académica")
    # EL TEXTO FIJO (Aparece en ambos modos porque es institucional)
    texto_fijo = """La fundamentación académica del Programa responde a los Lineamientos Académicos y Curriculares (LAC) de la I.U. Pascual Bravo, garantizando la coherencia entre el diseño curricular, la metodología pedagógica y los estándares de calidad definidos por el Ministerio de Educación Nacional de Colombia; conceptualizando los principios que orientan la estructuración del plan de estudios, abarcando las áreas de formación, la política de créditos, el tiempo de trabajo presencial e independiente, y las certificaciones temáticas, entre otros aspectos clave.
    En los LAC se establece la política de créditos académicos de la Universidad, siendo ésta el conjunto de lineamientos y procedimientos que rigen la asignación de créditos a los programas de formación en cuanto a mínimos y máximos, el porcentaje de créditos para cada una de las áreas de formación que debe incluir el programa; incluyendo a su vez información sobre cómo se asignan los créditos a diferentes tipos de cursos definidos como teórico-prácticos y prácticos, el requisito de grado o las prácticas profesionales. 
    Rutas educativas: Certificaciones Temáticas Tempranas
    Las Certificaciones Temáticas Tempranas son el resultado del agrupamiento de competencias y cursos propios del currículo en diferentes rutas educativas que posibilitan que el estudiante acceda a una certificación en la medida que avanza en su proceso formativo y demuestra el alcance de las competencias, y finalizan con la expedición de las micro y macro credenciales. Las certificaciones impulsan en el estudiante el deseo particular de adquirir habilidades relevantes en áreas específicas de su interés que les posibilite insertarse en el mercado laboral tempranamente, por lo tanto, son voluntarias. Las certificaciones son revisadas, y reestructuradas de ser necesario, con base en la evaluación de los resultados académicos o los procesos de autoevaluación que realiza el programa."""
    
    st.markdown(f"> {texto_fijo}")
    

    st.write("**Certificaciones Temáticas Tempranas**")    
    cert_data = st.data_editor(
            ej.get("tabla_cert_ej", [{"Nombre": "", "Curso 1": "", "Créditos 1": 0, "Curso 2": "", "Créditos 2": 0}]),
            num_rows="dynamic",      
            key="editor_cert"
        )
    
    if metodo_trabajo == "Manual (Desde cero)":
        st.write("") 
        st.write("**Áreas de formación (Ingreso Manual)**")
        
        st.text_area(
            "Descripción del Área de Fundamentación Específica del Programa :red[•]",
            value=ej.get("fund_especifica_desc", ""),
            height=150,
            placeholder="Escriba aquí la descripción...",
            key="input_area_especifica"
        )
    else:
        # MODO AUTOMATIZADO: Solo mostramos el resultado de la extracción
        st.write("**Área de Fundamentación Específica (Extraída del DM)**")
        st.text_area(
            "Contenido detectado:",
            key="input_area_especifica", # Mismo KEY para que el Word lo encuentre
            height=150,
            help="Este campo se llena automáticamente con la extracción por rangos."
        )

 # Itinerario formativo
    st.write("") 
    st.write("**3.Itinerario formativo**")
    
    area_especifica = st.text_area("Teniendo como fundamento que, en torno a un objeto de conocimiento se pueden estructurar varios programas a diferentes niveles de complejidad, es importante expresar si el programa en la actualidad es único en torno al objeto de conocimiento al que está adscrito o hay otros de mayor o de menor complejidad.:red[•]",
        value=ej.get("fund_especifica_desc", ""),
        height=150,
        placeholder=" Ejemplo si el PEP es de Ingeniería Mecánica, determinar si hay otro programa de menor complejidad como una tecnología Mecánica o uno de mayor complejidad como una especialización o una maestría. Este itinerario debe considerar posibles programas de la misma naturaleza que se puedan desarrollar en el futuro.",
        key="input_itinerario"
    )

    st.markdown("---")
    st.markdown("### 7. Recursos Académicos")

# Sección 7.1
    st.subheader("7.1. Entornos académicos")
    st.info("A continuación se incluirá el texto institucional sobre infraestructura. Puede añadir detalles específicos del programa abajo.")

    entornos_especificos = st.text_area(
    "Detalles específicos de entornos (Laboratorios, talleres, software especializado):",
        value=ej.get("entornos_desc", ""),
        placeholder="Ejemplo: El programa cuenta con el Laboratorio de Prototipado 3D y licencias de software...",
        height=120,
        key="input_recursos_7_1"
    )

# Subsección 7.2
    st.subheader("7.2. Talento Humano")
    
    perfil_docente = st.text_area(
        "Perfil del equipo docente requerido (Funciones sustantivas) :red[•]",
        value=ej.get("perfil_docente_desc", ""),
        placeholder="Describa la formación académica, experiencia profesional e investigativa que deben tener los docentes del programa...",
        height=150,
        key="input_talento_humano"
        )

    st.markdown("---")
    st.markdown("### 8. Investigación, Tecnología e Innovación")
    
    # Usamos la variable 'denom' para que el título sea dinámico
    texto_ayuda_inv = f"Describa cómo se desarrolla la investigación en el programa de {denom if denom else 'Nombre del Programa'}..."
    
    investigacion_desc = st.text_area(
        "Organización de la Investigación (Líneas, Grupos y Semilleros) :red[•]",
        value=ej.get("investigacion_desc", ""),
        placeholder=texto_ayuda_inv,
        height=250,
        key="input_investigacion"
    )
    st.caption("💡 Tip: Mencione el nombre de los grupos categorizados en MinCiencias y los semilleros activos vinculados al programa.")

    st.markdown("---")
    st.markdown("### 9. Vinculación Nacional e Internacional")
    
    st.info("Describa las estrategias de visibilidad del programa: convenios de doble titulación, redes académicas, movilidad de docentes/estudiantes y proyectos conjuntos con instituciones externas.")
    
    vinculacion_desc = st.text_area(
        "9.1 Estrategias de internacionalización. :red[•]",
        value=ej.get("vinculacion_desc", ""),
        placeholder="Ejemplo: El programa pertenece a la red de facultades de ingeniería nacional (ACOFI) y cuenta con convenios de movilidad con la Universidad de Politécnica de Valencia...",
        height=200,
        key="input_vinculacion"
    )

    st.markdown("---")
    st.markdown("### 10. Bienestar del Programa")
    
    st.info("Describa cómo las políticas de Bienestar Institucional impactan directamente al programa (permanencia, graduación estudiantil, clima organizacional y apoyo integral).")
    
    bienestar_desc = st.text_area(
        "Acciones y estrategias de bienestar para el programa :red[•]",
        value=ej.get("bienestar_desc", ""),
        placeholder="Ejemplo: El programa articula con Bienestar Universitario el seguimiento a la deserción mediante el sistema de alertas tempranas, además de promover la participación en actividades culturales y deportivas...",
        height=200,
        key="input_bienestar"
    )
    st.caption("💡 Tip: Mencione programas específicos como tutorías, apoyos socioeconómicos o estrategias de salud mental.")


    st.markdown("---")
    st.markdown("### 11. Estructura Administrativa")
    
    st.info("Suba el organigrama o esquema de la estructura administrativa del programa.")
    
    # Campo para subir la imagen
    archivo_organigrama = st.file_uploader(
        "Subir imagen del organigrama (JPG, PNG)", 
        type=["png", "jpg", "jpeg"],
        key="uplo_organigrama"
    )

    st.subheader("11.2 Órganos de Decisión")
    
    # Creamos dos columnas paralelas
    col_cc, col_cf = st.columns(2)
    
    with col_cc:
        comite_curricular = st.text_area(
            "Comité Curricular",
            value=ej.get("comite_curricular", ""),
            placeholder="Describa la composición y funciones principales del Comité Curricular en el programa...",
            height=200,
            key="input_comite_curr"
        )
        
    with col_cf:
        consejo_facultad = st.text_area(
            "Consejo de Facultad",
            value=ej.get("consejo_facultad", ""),
            placeholder="Describa el rol del Consejo de Facultad respecto a las decisiones estratégicas del programa...",
            height=200,
            key="input_consejo_fac"
        )

     # 4. Justificación del Programa
    if metodo_trabajo != "Automatizado (Cargar Documento Maestro)":
        st.write("") 
        st.write("**Justificación del Programa**")
        
        area_especifica = st.text_area("Demostrar la relevancia del programa en el contexto actual, resaltando su impacto en la solución de problemáticas sociales y productivas. Se debe enfatizar cómo la formación impartida contribuye al desarrollo del entorno local, regional y global, alineándose con las necesidades del sector productivo, las políticas educativas y las tendencias del mercado laboral.:red[•]",
            value=ej.get("fund_especifica_desc", ""),
            height=150,
            placeholder="Fundamentar la relevancia del programa con datos actualizados, referencias normativas y estudios sectoriales. Evidenciar su alineación con los Objetivos de Desarrollo Sostenible (ODS), planes de desarrollo nacionales y políticas de educación superior. Incorporar análisis de tendencias internacionales que justifiquen su pertinencia en un contexto globalizado.",
            key="input_just"
        )

    st.markdown("---")
    st.markdown("### 12. Evaluación y Mejoramiento Continuo")
    
    # Subsección 12.1
    st.subheader("12.1. Sistema de Aseguramiento de la calidad del Programa")
    
    aseguramiento_desc = st.text_area(
        "Descripción del Sistema de Aseguramiento de la Calidad :red[•]",
        value=ej.get("aseguramiento_desc", ""),
        placeholder="Describa los procesos de autoevaluación, seguimiento a planes de mejoramiento y cómo el programa utiliza los resultados para la toma de decisiones...",
        height=200,
        key="input_aseguramiento"
    )
    st.caption("💡 Tip: Mencione la articulación con el Modelo de Autoevaluación Institucional y la periodicidad de las revisiones curriculares.") 

        #  LÓGICA DE GENERACIÓN DEL WORD 
    generar = st.form_submit_button("🚀 GENERAR DOCUMENTO PEP", type="primary")

    if generar:
        if not denom or not reg1:
            st.error("⚠️ Falta información obligatoria (Denominación o Registro Calificado).")
        else:
            doc = Document()
            # Estilo base
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(11)
                # 1.1 Historia del Programa
            doc.add_heading("1.1. Historia del Programa", level=1)
                                
                            # PÁRRAFO 1. Datos creación
            texto_historia = (
                                f"El Programa de {denom} fue creado mediante el {acuerdo} del {instancia} "
                                f"y aprobado mediante la resolución de Registro Calificado {reg1} del Ministerio de Educación Nacional "
                                f"con código SNIES {snies}."
                            )
            doc.add_paragraph(texto_historia)
                        
                            # PÁRRAFO 2. Motivo de creación
            if motivo.strip():
                        # El usuario ya escribió empezando con "La creación del programa..."
                               doc.add_paragraph(motivo) 
            else:
                                doc.add_paragraph("No se suministró información sobre el motivo de creación.")
                         
                            # PÁRRAFO 3. Acreditación 1 y/o 2
            if acred1 and not acred2:
                        # Caso: Solo una acreditación
                                texto_acred = (
                                f"El programa obtuvo la Acreditación en alta calidad otorgada por el "
                                f"Consejo Nacional de Acreditación (CNA) a través de la resolución {acred1}, "
                                f"como reconocimiento a su solidez académica, administrativa y de impacto social."
                            )
                                doc.add_paragraph(texto_acred)
                    
            elif acred1 and acred2:
                        # Caso: Dos acreditaciones (Primera vez + Renovación)
                                texto_acred = (
                                f"El programa obtuvo por primera vez la Acreditación en alta calidad otorgada por el "
                                f"Consejo Nacional de Acreditación (CNA) a través de la resolución {acred1}, "
                                f"esta le fue renovada mediante resolución {acred2}, reafirmando la solidez "
                                f"académica, administrativa y de impacto social del Programa."
                            )
                                doc.add_paragraph(texto_acred)    
                    
                            # PÁRRAFO 4: Modificaciones curriculares
            planes_nom = [n for n in [p1_nom, p2_nom, p3_nom] if n]
            planes_fec_lista = [f for f in [p1_fec, p2_fec, p3_fec] if f]
                            
            if planes_fec_lista and planes_nom:
                                # A. Formatear nombres de planes (lo que antes era "lista")
                               if len(planes_nom) > 1:
                                    txt_planes_lista = ", ".join(planes_nom[:-1]) + f" y {planes_nom[-1]}"
            else:
                                    txt_planes_lista = planes_nom[0]
                    
                                # B. Formatear fechas/acuerdos
            if len(planes_fec_lista) > 1:
                                    txt_acuerdos_formateado = ", ".join(planes_fec_lista[:-1]) + f" y {planes_fec_lista[-1]}"
            else:
                                    txt_acuerdos_formateado = planes_fec_lista[0]
                    
            texto_planes = (
                                     f"El plan de estudios del Programa de {denom} ha sido objeto de procesos periódicos de evaluación, "
                                     f"con el fin de asegurar su pertinencia académica y su alineación con los avances tecnológicos "
                                     f"y las demandas del entorno. Como resultado, se han realizado las modificaciones curriculares "
                                     f"{txt_planes_lista}, aprobadas mediante el {txt_acuerdos_formateado}, respectivamente."
                                )
            p_planes = doc.add_paragraph(texto_planes)
            p_planes.alignment = 3  # Justificado
                        
                            # PÁRRAFO 5: Reconocimientos
            recons_validos = [r for r in recon_data if r.get("Nombre del premio", "").strip()]
                            
            if recons_validos:
                                 doc.add_paragraph(
                                     f"El Programa de {denom} ha alcanzado importantes logros académicos e institucionales "
                                     f"que evidencian su calidad y compromiso con la excelencia. Entre ellos se destacan:"
                                 )
            for r in recons_validos:
                                     premio = r.get("Nombre del premio", "N/A")
                                     anio = r.get("Año", "N/A")
                                     ganador = r.get("Nombre del Ganador", "N/A")
                                     cargo = r.get("Cargo", "N/A")
                                     doc.add_paragraph(
                                 f" {premio} ({anio}): Otorgado a {ganador}, en su calidad de {cargo}.", 
                                 style='List Bullet')
                    
                            # Línea de tiempo
            doc.add_heading("Línea de Tiempo del Programa", level=2)
                        # Función interna para extraer solo el año (4 dígitos)
            def extraer_anio(texto):
                                 if not texto: return "N/A"
                                 match = re.search(r'20\d{2}', str(texto)) # Busca "20" seguido de dos números
                                 return match.group(0) if match else str(texto).split()[-1]
                                
                        # 1. Creación (Usando el año del primer plan o acuerdo)
            if p1_fec:
                                 anio = extraer_anio(p1_fec)
                                 doc.add_paragraph(f"{anio}: Creación del Programa")
                    
                    
                        # 2. Registros Calificados
            if reg1:
                                        # Intenta extraer el año (asumiendo formato "Res XXX de 20XX")
                                 anio_reg1 = reg1.split()[-1] if len(reg1.split()) > 0 else "Fecha N/A"
                                 doc.add_paragraph(f"{anio_reg1}: Obtención del Registro Calificado inicial")
                    
            if reg2:
                                 anio_reg2 = reg2.split()[-1] if len(reg2.split()) > 0 else "Fecha N/A"
                                 doc.add_paragraph(f"{anio_reg2}: Renovación del Registro Calificado")
                    
                        # 3. Modificaciones Curriculares (Planes de estudio)
            if p2_fec:
                                  anio = extraer_anio(p2_fec)
                                  doc.add_paragraph(f"{anio}: Modificación curricular 1 (Actualización del plan de estudios)")
                            
            if p3_fec:
                                  anio = extraer_anio(p3_fec)
                                  doc.add_paragraph(f"{anio}: Modificación curricular 2")
                    
                        # 4. Acreditaciones de Alta Calidad
            if acred1:
                                  anio_acred1 = acred1.split()[-1] if len(acred1.split()) > 0 else "Fecha N/A"
                                  doc.add_paragraph(f"{anio_acred1}: Obtención de la Acreditación en Alta Calidad")
                            
            if acred2:
                                  anio_acred2 = acred2.split()[-1] if len(acred2.split()) > 0 else "Fecha N/A"
                                  doc.add_paragraph(f"{anio_acred2}: Renovación de la Acreditación en Alta Calidad")
                    
                            # 5. Reconocimientos (Si existen en la tabla)
            if recons_validos:
                                        # Tomamos los años únicos de los reconocimientos para no repetir
                                 anios_recon = sorted(list(set([r['Año'] for r in recons_validos if r['Año']])))
            for a in anios_recon:
                                     doc.add_paragraph(f"{a}: Reconocimientos académicos destacados")
                                    
                            # 1.2 GENERALIDADES (Tabla de datos)
            doc.add_page_break() 
            doc.add_heading("1.2 Generalidades del Programa", level=1)
                            # --- EXTRACCIÓN DE VALORES PARA LA TABLA ---
                            # Sacamos los datos del estado de la sesión para que las variables existan
            denom = st.session_state.get("denom_input", "N/A")
            titulo = st.session_state.get("titulo_input", "N/A")
            nivel = st.session_state.get("nivel_input", "N/A")
            area = st.session_state.get("area_input", "N/A")
            modalidad = st.session_state.get("modalidad_input", "N/A")
            acuerdo = st.session_state.get("acuerdo_input", "N/A")
            reg1 = st.session_state.get("reg_calificado_input", "N/A")
            creditos = st.session_state.get("creditos_input", "N/A")
            periodicidad = st.session_state.get("periodicidad_input", "N/A")
            lugares = st.session_state.get("lugares_input", "N/A")
            snies = st.session_state.get("snies_input", "N/A")
            reg1_val = st.session_state.get("reg1", "N/A")
            creditos_val = st.session_state.get("creditos", "N/A")
            lugares_val = st.session_state.get("lugar", "N/A")
                   
            items_gen = [
                                                ("Denominación del programa", denom),
                                                ("Título otorgado", titulo),
                                               ("Nivel de formación", nivel),
                                                ("Área de formación", area),
                                                ("Modalidad de oferta", modalidad),
                                                ("Acuerdo de creación", acuerdo),
                                                ("Registro calificado", reg1),
                                                ("Créditos académicos", creditos),
                                                ("Periodicidad de admisión", periodicidad),
                                                ("Lugares de desarrollo", lugares),
                                                ("Código SNIES", snies)
                                            ]
                            
            for k, v in items_gen:
                                                p = doc.add_paragraph()
                                                p.add_run(f"{k}: ").bold = True
                                                p.add_run(str(v))
                    
                      # 2.1 Referentes conceptuales 
            doc.add_heading("2.1. Referentes conceptuales", level=2)
                    
            obj_nom = st.session_state.get("obj_concep_input", "No definido")
            obj_con = st.session_state.get("obj_concep_input", "")
                    
                            # Bloque: Objeto + Enter + Conceptualización
            p_obj = doc.add_paragraph()
            p_obj.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY 
            p_obj.add_run("Objeto de conocimiento del Programa: ").bold = True
            p_obj.add_run(str(obj_nom)) # Forzamos a string para evitar errores
                    
            obj_con = st.session_state.get("full_problemas_input", st.session_state.get("problemas_input", ""))
                    
            if obj_con.strip():
                        p_concep = doc.add_paragraph(obj_con)
                        p_concep.alignment = 3
                    
            fundamentacion = st.session_state.get("full_input_epi_p1", st.session_state.get("input_epi_p1", ""))
                    
            if fundamentacion.strip():
                            doc.add_heading("Fundamentación Epistemológica", level=3)
                            p_fund = doc.add_paragraph(fundamentacion)
                            p_fund.alignment = 3
            else:
                            doc.add_paragraph("\n(Sección de fundamentación no suministrada)")
                    
                            # ---PERFILES ---
            doc.add_heading("1.2. Perfiles del Programa", level=1)
                    
            doc.add_heading("1.2.1. Perfil Profesional con Experiencia", level=2)
            doc.add_paragraph(perfil_exp if perfil_exp.strip() else "No definido.")
                    
            doc.add_heading("1.2.2. Perfil Profesional del Egresado", level=2)
            doc.add_paragraph(perfil_egresado if perfil_egresado.strip() else "No definido.")
                    
            doc.add_heading("1.2.3. Perfil Ocupacional", level=2)
            doc.add_paragraph(perfil_ocupacional if perfil_ocupacional.strip() else "No definido.")
                    
                            
                    
                            # --- MANEJO DE CITAS (Sincronizado) ---
                            # 1. Obtenemos los datos de la tabla (Modo Manual)
                            # Si no existe la key, devolvemos una lista vacía por defecto
            raw_concep = st.session_state.get("editor_referencias", [])
                           
            citas_c = []
            datos_lista = []
                    
                            # 2. Normalizar los datos según cómo vengan del st.data_editor
            if isinstance(raw_concep, dict):
                                # Si el usuario editó la tabla, Streamlit a veces devuelve un dict con 'edited_rows'
                                datos_lista = list(raw_concep.get("edited_rows", {}).values())
            elif isinstance(raw_concep, list):
                                # Si es la lista inicial cargada desde el ejemplo o BD
                                datos_lista = raw_concep
                            
                            # 3. Extraer Autor y Año de cada fila válida
            for fila in datos_lista:
                             if isinstance(fila, dict):
                                    aut = ""
                                    ani = ""
                                    # Buscamos de forma flexible (no importa si es "Autor" o "autor")
            for k, v in fila.items():
                                        k_low = str(k).lower()
                                        if "autor" in k_low: aut = str(v).strip()
                                        if "año" in k_low or "anio" in k_low: ani = str(v).strip()
                                    
                                    # Solo agregamos si hay un autor real (evitamos campos vacíos o "None")
            if aut and aut.lower() != "none" and aut != "":
                                        # Si el año está vacío, solo ponemos el autor
                                        citas_c.append(f"{aut}, {ani}" if ani else aut)
                    
                            # 4. Pegar las citas al párrafo del Objeto de Conocimiento
            if citas_c:
                                # Asegúrate de que p_obj o p_concep existan antes de esta línea
                                p_obj.add_run(" (Sustentado en: " + "; ".join(citas_c) + ").")
                               
                        
                            
                       
                            # --- 2.2 FUNDAMENTACIÓN EPISTEMOLÓGICA ---
            doc.add_heading("2.2. Fundamentación epistemológica", level=2)
                            
                            # Iteramos los 3 bloques de las pestañas
            for i in range(1, 4):
                                llave_full = f"full_input_epi_p{i}"
                                llave_normal = f"input_epi_p{i}"
                                texto_p = st.session_state.get(llave_full, st.session_state.get(llave_normal, ""))
                                if texto_p:
                                    if "[... " in texto_p and " PÁRRAFOS INTERMEDIOS" in texto_p:
                                        st.warning(f"Aviso: El bloque {i} de epistemología parece estar incompleto.")
                                    p_f = doc.add_paragraph(texto_p)
                                    p_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    
                                    raw_f = st.session_state.get(f"editor_refs_p{i}", [])
                                    
                                    # Normalizar datos de la tabla de la pestaña
                                    if isinstance(raw_f, dict):
                                        datos_f = raw_f.get("data", list(raw_f.get("edited_rows", {}).values()))
                                    else:
                                        datos_f = raw_f
                                        
                                    citas_p = []
                                    for f in datos_f:
                                        if isinstance(f, dict):
                                            a_f, n_f = "", ""
                                            for k, v in f.items():
                                                k_l = str(k).lower()
                                                if "autor" in k_l: a_f = str(v).strip()
                                                if "año" in k_l or "anio" in k_l: n_f = str(v).strip()
                                            if a_f and n_f and a_f.lower() != "none" and a_f != "":
                                                citas_p.append(f"{a_f}, {n_f}")
                                    
                                    if citas_p:
                                        p_f.add_run(" (Ref: " + "; ".join(citas_p) + ").")
                                    
                        # 2.3 Fundamentación Académica (TEXTO FIJO PASCUAL BRAVO)
                         
            doc.add_heading("2.3. Fundamentación académica", level=2)
            doc.add_paragraph("La fundamentación académica del Programa responde a los Lineamientos Académicos y Curriculares (LAC) de la I.U. Pascual Bravo, garantizando la coherencia entre el diseño curricular, la metodología pedagógica y los estándares de calidad definidos por el Ministerio de Educación Nacional de Colombia; conceptualizando los principios que orientan la estructuración del plan de estudios, abarcando las áreas de formación, la política de créditos, el tiempo de trabajo presencial e independiente, y las certificaciones temáticas, entre otros aspectos clave.")
                            #p1_fa.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY   
            doc.add_paragraph("En los LAC se establece la política de créditos académicos de la Universidad, siendo ésta el conjunto de lineamientos y procedimientos que rigen la asignación de créditos a los programas de formación en cuanto a mínimos y máximos, el porcentaje de créditos para cada una de las áreas de formación que debe incluir el programa; incluyendo a su vez información sobre cómo se asignan los créditos a diferentes tipos de cursos definidos como teórico-prácticos y prácticos, el requisito de grado o las prácticas profesionales.")
                       
            doc.add_heading("Rutas educativas: Certificaciones Temáticas Tempranas", level=3)
            doc.add_paragraph("Las Certificaciones Temáticas Tempranas son el resultado del agrupamiento de competencias y cursos propios del currículo en diferentes rutas educativas que posibilitan que el estudiante acceda a una certificación en la medida que avanza en su proceso formativo y demuestra el alcance de las competencias, y finalizan con la expedición de las micro y macro credenciales. Las certificaciones impulsan en el estudiante el deseo particular de adquirir habilidades relevantes en áreas específicas de su interés que les posibilite insertarse en el mercado laboral tempranamente, por lo tanto, son voluntarias. Las certificaciones son revisadas, y reestructuradas de ser necesario, con base en la evaluación de los resultados académicos o los procesos de autoevaluación que realiza el programa.")
                    # Subsección 7.2: Talento Humano
            doc.add_heading("7.2. Talento Humano", level=2)
                    
            if perfil_docente.strip():
                        # El texto que el usuario redactó en el formulario
                        p_talento = doc.add_paragraph(perfil_docente)
                        p_talento.alignment = 3  # Justificado
            else:
                        doc.add_paragraph("No se suministró información sobre el perfil del equipo docente.")
    
            doc.add_heading("8. Investigación, Tecnología e Innovación", level=1)
                            
            if investigacion_desc.strip():
                                # Añadimos el contenido redactado por el usuario
                        p_inv = doc.add_paragraph(investigacion_desc)
                        p_inv.alignment = 3  # Justificado
            else:
                        doc.add_paragraph("Pendiente por describir la articulación de grupos y semilleros de investigación.")
                            
                    
                         
                        # Tabla de Certificaciones
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text = 'Certificación', 'Cursos', 'Créditos Totales'
                        
            for c in cert_data:
                            if c["Nombre"]:
                                                  row = table.add_row().cells
                                                  row[0].text = c["Nombre"]
                                                  row[1].text = f"{c['Curso 1']}, {c['Curso 2']}"
                                                  row[2].text = str(c["Créditos 1"] + c["Créditos 2"])
    
    
                    # --- SECCIÓN 7: RECURSOS ACADÉMICOS ---
            doc.add_heading("7. Recursos académicos", level=1)
                    
                    # Párrafo fijo institucional
            parrafo_fijo_recursos = (
                        "La I.U. Pascual Bravo dispone de una infraestructura y una serie de recursos que garantizan el "
                        "adecuado desarrollo académico, investigativo y administrativo de sus programas. Estos recursos "
                        "están diseñados para apoyar la formación integral de los estudiantes y fortalecer la gestión "
                        "docente, asegurando la calidad y pertinencia del Programa en concordancia con los lineamientos "
                        "institucionales."
                    )
            p_fijo = doc.add_paragraph(parrafo_fijo_recursos)
            p_fijo.alignment = 3  # Justificado
                    
                   # Subsección 7.1
            doc.add_heading("7.1. Entornos académicos", level=2)
                    
                    # Si el usuario escribió algo en la casilla, se añade al Word
            if entornos_especificos.strip():
                        doc.add_paragraph(entornos_especificos)
            else:
                        doc.add_paragraph("El programa hace uso de los entornos académicos generales dispuestos por la institución.")
    
                        # --- SECCIÓN 9: VINCULACIÓN ---
            doc.add_heading("9. Vinculación Nacional e Internacional", level=1)
                    
            if vinculacion_desc.strip():
                        # Añadimos el contenido redactado por el usuario
                        p_vinc = doc.add_paragraph(vinculacion_desc)
                        p_vinc.alignment = 3  # Justificado
            else:
                        doc.add_paragraph("No se ha registrado información sobre convenios o redes de cooperación.")
    
            # --- SECCIÓN 10: BIENESTAR ---
            doc.add_heading("10. Bienestar del Programa", level=1)
                    
            if bienestar_desc.strip():
                        # Añadimos el contenido redactado por el usuario
                        p_bien = doc.add_paragraph(bienestar_desc)
                        p_bien.alignment = 3  # Justificado (3 corresponde a WD_ALIGN_PARAGRAPH.JUSTIFY)
            else:
                        doc.add_paragraph("Se aplican las políticas generales de bienestar institucional enfocadas en la permanencia y el desarrollo integral.")
    
            # --- SECCIÓN 11: ESTRUCTURA ADMINISTRATIVA ---
            doc.add_heading("11. Estructura Administrativa", level=1)
                    
            doc.add_paragraph("A continuación se presenta la estructura administrativa y organizacional que soporta la gestión del programa:")
            
            if archivo_organigrama is not None:
                        # Insertar la imagen subida por el usuario
                        doc.add_picture(archivo_organigrama, width=Inches(6.0))
                        # Opcional: Centrar la imagen
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = 1 # 1 es para Centrado
            else:
                        doc.add_paragraph("[Pendiente: Insertar organigrama del programa]")
    
                # --- SUBSECCIÓN 11.2: ÓRGANOS DE DECISIÓN ---
            doc.add_heading("11.2. Órganos de decisión", level=2)
                    
                    # Comité Curricular
            doc.add_heading("11.2.1. Comité Curricular", level=3)
            if comite_curricular.strip():
                        p_cc = doc.add_paragraph(comite_curricular)
                        p_cc.alignment = 3  # Justificado
            else:
                        doc.add_paragraph("Información pendiente sobre el Comité Curricular.")
            
                    # Consejo de Facultad
            doc.add_heading("11.2.2. Consejo de Facultad", level=3)
            if consejo_facultad.strip():
                        p_cf = doc.add_paragraph(consejo_facultad)
                        p_cf.alignment = 3  # Justificado
            else:
                        doc.add_paragraph("Información pendiente sobre el Consejo de Facultad.")
    
                        # --- SECCIÓN 12: EVALUACIÓN Y MEJORAMIENTO ---
            doc.add_heading("12. Evaluación y Mejoramiento continuo", level=1)
                    
                    # Subsección 12.1
            doc.add_heading("12.1. Sistema de Aseguramiento de la calidad del Programa", level=2)
                   
            if aseguramiento_desc.strip():
                        # Añadimos el contenido redactado por el usuario
                        p_aseg = doc.add_paragraph(aseguramiento_desc)
                        p_aseg.alignment = 3  # Justificado
            else:
                        doc.add_paragraph("El programa se acoge al sistema institucional de aseguramiento de la calidad, realizando procesos periódicos de autoevaluación y actualización curricular.")
                        
                            # Guardar archivo3
            bio = io.BytesIO()
            doc.save(bio)
            bio.seek(0)
                                                
            st.success("✅ ¡Documento PEP generado!")
            st.download_button(
                                                     label="📥 Descargar Documento PEP en Word",
                                                      data=bio.getvalue(),
                                                      file_name=f"PEP_Modulo1_{denom.replace(' ', '_')}.docx",
                                                      mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                                      )
