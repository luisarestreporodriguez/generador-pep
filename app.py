import streamlit as st
from google import genai
from docx import Document
from docx.shared import Pt
import io
import time
import re
import pandas as pd 
from streamlit_gsheets import GSheetsConnection 

# 1. ESTABLECER CONEXIÓN
conn = st.connection("gsheets", type=GSheetsConnection)

# --- BARRA LATERAL PARA CARGA ---
with st.sidebar:
    st.header("💾 Identificación del Programa")
    email_usuario = st.text_input("Correo electrónico")
    snies_input = st.text_input("SNIES")
    
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("🔍 Cargar Datos"):
            if email_usuario and snies_input:
                # Aquí irá la lógica de búsqueda después
                st.success("Datos recuperados")
            else:
                st.warning("Ingresa Email y SNIES")

    with col2:
        if st.button("💾 Guardar Progreso"):
            if email_usuario and snies_input:
                # 1. Recopilar datos (Asegúrate que 'denom' exista en tu código arriba)
                datos_a_guardar = {
                    "SNIES": snies_input,
                    "Email": email_usuario,
                    "Denominacion": denom if 'denom' in locals() else "", 
                    "Fecha": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M")
                }
                
                # 2. Convertir a DataFrame y enviar a Sheets
                # TODO ESTO DEBE IR CON LA MISMA SANGRÍA QUE EL PASO 1
                df_nuevo = pd.DataFrame([datos_a_guardar])
            
                # Leer lo que ya hay para no borrarlo
                try:
                    df_actual = conn.read()
                    df_final = pd.concat([df_actual, df_nuevo], ignore_index=True)
                    
                    # Actualizar la hoja
                    conn.update(data=df_final)
                    st.info("✅ Progreso guardado en la nube (Google Sheets)")
                except Exception as e:
                    st.error(f"Error al conectar con Sheets: {e}")
            
            else: # Este else ahora sí está alineado con 'if email_usuario...'
                st.error("Faltan datos de identificación")

# --- CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(page_title="Generador PEP", page_icon="📚", layout="wide")

st.title("Generador PEP - Módulo 1: Información del Programa")

# --- LÓGICA DE API KEY (Nube + Local) ---
# Intentamos leer la clave desde los Secrets de Streamlit
if "GEMINI_API_KEY" in st.secrets:
    api_key = st.secrets["GEMINI_API_KEY"]
else:
    with st.sidebar:
        st.header("Configuración Local")
        api_key = st.text_input("Ingresa tu Google API Key", type="password")
        if not api_key:
            st.warning("⚠️ Sin API Key la IA no podrá redactar textos largos.")
            
# --- FUNCIÓN DE REDACCIÓN IA ---
def redactar_seccion_ia(titulo_seccion, datos_seccion):
    if not api_key: return "Error: No hay API Key configurada."
    respuestas_reales = {k: v for k, v in datos_seccion.items() if str(v).strip()}
    contexto = "\n".join([f"- {k}: {v}" for k, v in respuestas_reales.items()])
    
    try:
        client = genai.Client(api_key=api_key)
        prompt = f"""
        Actúa como un Vicerrector Académico experto en aseguramiento de la calidad.
        Tarea: Redactar el motivo de creación del Programa
        DATOS SUMINISTRADOS:{contexto}
        
        REGLAS CRÍTICAS DE FORMATO:
        1. Responde ÚNICAMENTE con UN SOLO PÁRRAFO de texto corrido.
        2. NO incluyas títulos, ni subtítulos (prohibido usar "##" o "Contexto").
        3. NO uses negritas, ni corchetes, ni nombres de la institución entre etiquetas.
        4. Empieza directamente con la redacción (ej: "La pertinencia de este programa se fundamenta...").
        5. El tono debe ser muy formal, académico y fluido. Máximo 150 palabras.
       """
        response = client.models.generate_content(model="gemini-flash-latest", contents=prompt)
        # Limpiar posibles espacios en blanco extras o saltos de línea al inicio/final
        return response.text
    except Exception as e:
        return f"Error en redacción: {str(e)}"

# --- ESTRUCTURA DE CONTENIDOS ---
estructura_pep = {
    "1. Información del Programa": {
        "1.1. Historia del Programa": {"tipo": "especial_historia"},
        "1.2. Generalidades del Programa": {"tipo": "directo"}
    },
    "2. Referentes Conceptuales": {
    "2.1. Naturaleza del Programa": {
        "tipo": "ia",
        "campos": [
            {"label": "Objeto de conocimiento del Programa", "req": True, "help": "¿Qué conoce, investiga y transforma este programa?"}
        ]
    },
    "2.2. Fundamentación epistemológica": {
        "tipo": "ia",
        "campos": [
            {"label": "Naturaleza epistemológica e identidad académica", "req": True},
            {"label": "Campo del saber y relación con ciencia/tecnología", "req": True}
        ]
    },
    "2.3. Fundamentación académica": {
        "tipo": "especial_pascual", # Nueva lógica para textos fijos + tabla
        "campos": [] 
    }
}

}


# --- BOTÓN DE DATOS DE EJEMPLO ---
# Usamos session_state para persistir los datos al hacer clic
if st.button("🧪 Llenar con datos de ejemplo"):
    st.session_state.ejemplo = {
        "denom": "Ingeniería de Sistemas",
        "titulo": "Ingeniero de Sistemas",
        "nivel_idx": 2, # Profesional universitario
        "area": "Ingeniería, Arquitectura y Urbanismo",
        "modalidad_idx": 4, # Presencial y Virtual
        "acuerdo": "Acuerdo 012 de 2015",
        "instancia": "Consejo Académico",
        "reg1": "Res. 4567 de 2016",
        "reg2": "Res. 8901 de 2023",
        "acred1": "Res. 00234 de 2024",
        "creditos": "165",
        "periodo_idx": 0, # Semestral
        "lugar": "Sede Principal (Medellín)",
        "snies": "1234",
        "motivo": "El programa se fundamenta en la necesidad regional de formar profesionales capaces de liderar la transformación digital y el desarrollo de software de alta complejidad.",
        "p1_nom": "EO01", "p1_fec": "Acuerdo 012-2015",
        "p2_nom": "EO02", "p2_fec": "Acuerdo 088-2020",
        "p3_nom": "EO03", "p3_fec": "Acuerdo 089-2024",
        "recon_data": [
                    {
                        "Nombre del premio": "Orden al Mérito Académico", 
                        "Año": "2022", 
                        "Nombre del Ganador": "Juan Pérez", 
                        "Cargo": "Docente Investigador"
                    },
                    {
                        "Nombre del premio": "Mejor Puntaje Saber Pro", 
                        "Año": "2023", 
                        "Nombre del Ganador": "María López", 
                        "Cargo": "Estudiante"  
                        }
                        ],
        #DATOS CAPÍTULO 2
        "objeto_con": "El programa investiga el ciclo de vida del software, la arquitectura de sistemas complejos y la integración de IA para transformar procesos industriales.",
        "fund_epi": "El programa se inscribe en el racionalismo crítico y el pragmatismo tecnológico, vinculando la ciencia de la computación con la ingeniería aplicada.",
        # DATOS PARA LAS TABLAS (Se guardan como listas de diccionarios)
        "tabla_recon_ej": [
            {"Año": "2024", "Nombre del premio": "Excelencia Académica", "Nombre del Ganador": "Juan Pérez", "Cargo": "Docente"}
        ],
        "tabla_cert_ej": [
            {"Nombre": "Desarrollador Web Junior", "Curso 1": "Programación I", "Créditos 1": 3, "Curso 2": "Bases de Datos", "Créditos 2": 4},
            {"Nombre": "Analista de Datos", "Curso 1": "Estadística", "Créditos 1": 4, "Curso 2": "Python para Ciencia", "Créditos 2": 4}
        ]
    }
    st.rerun()

# --- FORMULARIO DE ENTRADA ---
with st.form("pep_form"):
    ej = st.session_state.get("ejemplo", {})

    st.markdown("### 📋 1. Identificación General")
    col1, col2 = st.columns(2)
    with col1:
        denom = st.text_input("Denominación del programa (Obligatorio)", value=ej.get("denom", ""))
        titulo = st.text_input("Título otorgado (Obligatorio)", value=ej.get("titulo", ""))
        nivel = st.selectbox("Nivel de formación (Obligatorio)", 
                             ["Técnico", "Tecnológico", "Profesional universitario", "Especialización", "Maestría", "Doctorado"], 
                             index=ej.get("nivel_idx", 2))
        area = st.text_input("Área de formación (Obligatorio)", value=ej.get("area", ""))
    
    with col2:
        modalidad = st.selectbox("Modalidad de oferta (Obligatorio)", 
                                 ["Presencial", "Virtual", "A Distancia", "Dual", "Presencial y Virtual", "Presencial y a Distancia", "Presencial y Dual"],
                                 index=ej.get("modalidad_idx", 0))
        acuerdo = st.text_input("Acuerdo de creación / Norma interna (Obligatorio)", value=ej.get("acuerdo", ""))
        instancia = st.text_input("Instancia interna que aprueba (Obligatorio)", value=ej.get("instancia", ""))
        snies = st.text_input("Código SNIES (Obligatorio)", value=ej.get("snies", ""))

    st.markdown("---")
    st.markdown("### 📄 2. Registros, Acreditaciones y Tiempos")
    col3, col4 = st.columns(2)
    with col3:
        reg1 = st.text_input("Resolución Registro calificado 1 (Obligatorio)", value=ej.get("reg1", ""), placeholder="Número y año")
        reg2 = st.text_input("Registro calificado 2 (Opcional)", value=ej.get("reg2", ""))
        acred1 = st.text_input("Resolución Acreditación en alta calidad 1 (Opcional)", value=ej.get("acred1", ""))
        acred2 = st.text_input("Resolución Acreditación en alta calidad 2 (Opcional)", value="")

    with col4:
        creditos = st.text_input("Créditos académicos (Obligatorio)", value=ej.get("creditos", ""))
        periodicidad = st.selectbox("Periodicidad de admisión (Obligatorio)", ["Semestral", "Anual"], index=ej.get("periodo_idx", 0))
        lugares = st.text_input("Lugares de desarrollo (Obligatorio)", value=ej.get("lugar", ""))

    motivo = st.text_area("Motivo de creación del Programa (Obligatorio)", value=ej.get("motivo", ""), height=100)

    st.markdown("---")
    st.markdown("### 🧬 3. Planes de Estudios")
    p_col1, p_col2, p_col3 = st.columns(3)
    with p_col1:
        p1_nom = st.text_input("Nombre Plan v1 (Obligatorio)", value=ej.get("p1_nom", ""))
        p1_fec = st.text_input("Acuerdo aprobación Plan v1 (Obligatorio)", value=ej.get("p1_fec", ""))
    with p_col2:
        p2_nom = st.text_input("Nombre Plan v2 (Opcional)", value=ej.get("p2_nom", ""))
        p2_fec = st.text_input("Acuerdo aprobación Plan v2 (Opcional)", value=ej.get("p2_fec", ""))
    with p_col3:
        p3_nom = st.text_input("Nombre Plan v3 (Opcional)", value=ej.get("p3_nom", ""))
        p3_fec = st.text_input("Acuerdo aprobación Plan v3 (Opcional)", value=ej.get("p3_fec", ""))

    st.markdown("---")
    st.markdown("### 🏆 4. Reconocimientos (Opcional)")

    datos_finales = st.session_state.get("ejemplo", {}).get("recon_data", 
    [{"Año": "", "Nombre del premio": "", "Nombre del Ganador": "", "Cargo": "Estudiante"}] # <--- TU CÓDIGO AQUÍ
)
    recon_data = st.data_editor(
        datos_finales, # <--- Se conecta aquí
    num_rows="dynamic",
    key="editor_recon",
    column_config={
        "Cargo": st.column_config.SelectboxColumn(
            options=["Docente", "Líder", "Decano", "Estudiante", "Docente Investigador"]
        )
    }
)
    
# --- CAPÍTULO 2 ---
    st.markdown("---")
    st.header("2. Referentes Conceptuales")

    # 2.1. Naturaleza del Programa
    objeto_con = st.text_area(
        "Objeto de conocimiento del Programa (Obligatorio)", 
        value=ej.get("objeto_con", ""), 
        help="¿Qué conoce, investiga y transforma?",
        key="input_objeto"
    ) 

    # 2.2. Fundamentación epistemológica
    fund_epi = st.text_area(
        "Fundamentación epistemológica (Instrucciones 1 y 2)",
        value=ej.get("fund_epi", ""), 
        key="input_epi"
    )

    # 2.3. Fundamentación académica 
    st.subheader("Certificaciones Temáticas Tempranas")
    cert_data = st.data_editor(
        ej.get("tabla_cert_ej", [{"Nombre": "", "Curso 1": "", "Créditos 1": 0, "Curso 2": "", "Créditos 2": 0}]),
        num_rows="dynamic",      
        key="editor_cert"
    )

    generar = st.form_submit_button("🚀 GENERAR DOCUMENTO PEP", type="primary")

# --- LÓGICA DE GENERACIÓN DEL WORD ---
if generar:
    if not denom or not reg1:
        st.error("⚠️ Falta información obligatoria (Denominación o Registro Calificado).")
    else:
        #1. Crear el documento
        doc = Document()
        # Estilo base
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        
        # --- BLOQUE IA CON MEMORIA (SESSION STATE) ---
        # Solo llamamos a la API si el texto no existe en memoria
        
    if "motivo_ia_cache" not in st.session_state:
            with st.spinner("🤖 La IA está redactando el motivo (esto solo se hace una vez)..."):
                st.session_state.motivo_ia_cache = redactar_seccion_ia("Motivo de Creación", {"Motivo": motivo})
        
    if "naturaleza_ia_cache" not in st.session_state:
            with st.spinner("🤖 Redactando Naturaleza del Programa..."):
                st.session_state.naturaleza_ia_cache = redactar_seccion_ia("Naturaleza", {"Objeto": objeto_con})

# 1.1 Historia del Programa
    doc.add_heading("1.1. Historia del Programa", level=1)
        
        # PÁRRAFO 1. Datos de creación
    texto_historia = (
            f"El Programa de {denom} fue creado mediante el {acuerdo} del {instancia} "
            f"y aprobado mediante la resolución de Registro Calificado {reg1} del Ministerio de Educación Nacional "
            f"con código SNIES {snies}."
        )
    doc.add_paragraph(texto_historia)

        # PÁRRAFO 2. Motivo de creación (Desde la memoria de la IA)
    p_motivo = doc.add_paragraph(st.session_state.motivo_ia_cache)
    p_motivo.alignment = 3  # Justificado

        # PÁRRAFO 3. Acreditación 1 y/o 2
    if acred1:
            if not acred2:
                texto_acred = (
                    f"El programa obtuvo la Acreditación en alta calidad otorgada por el "
                    f"Consejo Nacional de Acreditación (CNA) a través de la resolución {acred1}, "
                    f"como reconocimiento a su solidez académica, administrativa y de impacto social."
                )
            else:
                texto_acred = (
                    f"El programa obtuvo por primera vez la Acreditación en alta calidad otorgada por el "
                    f"Consejo Nacional de Acreditación (CNA) a través de la resolución {acred1}, "
                    f"esta le fue renovada mediante resolución {acred2}, reafirmando la solidez "
                    f"académica, administrativa y de impacto social del Programa."
                )
            doc.add_paragraph(texto_acred)

        # PÁRRAFO 4. Evolución Curricular
    planes_nom = [n for n in [p1_nom, p2_nom, p3_nom] if n]
    planes_fec = [f for f in [p1_fec, p2_fec, p3_fec] if f]

    if planes_nom and planes_fec:
            # 1. Lógica para los ACUERDOS (txt_acuerdos)
            if len(planes_nom) == 1:
                txt_acuerdos = planes_nom[0]
            elif len(planes_nom) == 2:
                txt_acuerdos = f"{planes_nom[0]} y {planes_nom[1]}"
            else:
                txt_acuerdos = ", ".join(planes_nom[:-1]) + f" y {planes_nom[-1]}"

            # 2. Lógica para los AÑOS/PLANES (txt_anios)
            if len(planes_fec) == 1:
                txt_anios = planes_fec[0]
            elif len(planes_fec) == 2:
                txt_anios = f"{planes_fec[0]} y {planes_fec[1]}"
            else:
                txt_anios = ", ".join(planes_fec[:-1]) + f" y {planes_fec[-1]}"

            # 3. Redacción final (Variables sincronizadas)
    texto_planes = (
            f"El plan de estudios del Programa de {denom} ha sido objeto de procesos periódicos de evaluación, "
            f"con el fin de asegurar su pertinencia académica y su alineación con los avances tecnológicos "
            f"y las demandas del entorno. Como resultado, se han realizado las modificaciones curriculares "
            f"{txt_acuerdos}, aprobadas mediante el {txt_anios}, respectivamente."
        )
    doc.add_paragraph(texto_planes)

        # PÁRRAFO 5: Reconocimientos
    recons_validos = [r for r in recon_data if r.get("Nombre del premio", "").strip()]
        
    if recons_validos:
            doc.add_paragraph(
                f"El Programa de {denom} ha alcanzado importantes logros académicos e institucionales "
                f"que evidencian su calidad y compromiso con la excelencia. Entre ellos se destacan:"
            )
            for r in recons_validos:
                premio = r.get("Nombre del premio", "N/A")
                anio = r.get("Año", "N/A")
                ganador = r.get("Nombre del Ganador", "N/A")
                cargo = r.get("Cargo", "N/A")
                
                doc.add_paragraph(
                    f" {premio} ({anio}): Otorgado a {ganador}, en su calidad de {cargo}.", 
                    style='List Bullet'
                )

        # Línea de tiempo


# --- SECCIÓN: LÍNEA DE TIEMPO ---
doc.add_heading('Línea de Tiempo del Programa', level=2)

# 1. Función para extraer solo el año (busca 4 números seguidos)
def limpiar_anio(texto):
    if not texto: return None
    match = re.search(r'(19|20)\d{2}', str(texto))
    return match.group(0) if match else None

# 2. Recopilar todos los hitos en una lista para poder ordenarlos
lista_hitos = []

if p1_fec:
    anio = limpiar_anio(p1_fec)
    if anio: lista_hitos.append((anio, "Creación del Programa"))

if reg1:
    anio = limpiar_anio(reg1)
    if anio: lista_hitos.append((anio, "Obtención del Registro Calificado inicial"))

if reg2:
    anio = limpiar_anio(reg2)
    if anio: lista_hitos.append((anio, "Renovación del Registro Calificado"))

if p2_fec:
    anio = limpiar_anio(p2_fec)
    if anio: lista_hitos.append((anio, "Modificación curricular 1 (Actualización del plan de estudios)"))

if p3_fec:
    anio = limpiar_anio(p3_fec)
    if anio: lista_hitos.append((anio, "Modificación curricular 2"))

if acred1:
    anio = limpiar_anio(acred1)
    if anio: lista_hitos.append((anio, "Obtención de la Acreditación en Alta Calidad"))

if acred2:
    anio = limpiar_anio(acred2)
    if anio: lista_hitos.append((anio, "Renovación de la Acreditación en Alta Calidad"))

# Agregar también los reconocimientos a la línea de tiempo
for r in recons_validos:
    anio = limpiar_anio(r.get("Año"))
    if anio:
        lista_hitos.append((anio, f"Reconocimiento académico: {r.get('Nombre del premio')}"))

# 3. Ordenar cronológicamente por el año
lista_hitos.sort(key=lambda x: x[0])

# 4. Escribir en el documento
for anio, descripcion in lista_hitos:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f"{anio}: ")
    run.bold = True  # El año sale en negrita
    p.add_run(descripcion)


    

        # 5. Reconocimientos (Si existen en la tabla)
    if recons_validos:
            # Tomamos los años únicos de los reconocimientos para no repetir
            anios_recon = sorted(list(set([r['Año'] for r in recons_validos if r['Año']])))
            for a in anios_recon:
                doc.add_paragraph(f"{a}: Reconocimientos académicos destacados")
                
        # 1.2 GENERALIDADES (Tabla de datos)
    doc.add_page_break()
    doc.add_heading("1.2 Generalidades del Programa", level=1)
        
    items_gen = [
                ("Denominación del programa", denom),
                ("Título otorgado", titulo),
                ("Nivel de formación", nivel),
                ("Área de formación", area),
                ("Modalidad de oferta", modalidad),
                ("Acuerdo de creación", acuerdo),
                ("Registro calificado", reg1),
                ("Créditos académicos", creditos),
                ("Periodicidad de admisión", periodicidad),
                ("Lugares de desarrollo", lugares),
                ("Código SNIES", snies)
            ]
        
    for k, v in items_gen:
                p = doc.add_paragraph()
                p.add_run(f"{k}: ").bold = True
                p.add_run(str(v))

# 2.1 Naturaleza
    doc.add_heading("2.1. Naturaleza del Programa", level=2)
    doc.add_paragraph(redactar_seccion_ia("Naturaleza del Programa", {"Objeto": objeto_con}))

    # 2.2 Epistemología
    doc.add_heading("2.2. Fundamentación epistemológica", level=2)
    doc.add_paragraph(redactar_seccion_ia("Fundamentación Epistemológica", {"Datos": fund_epi}))

    # 2.3 Fundamentación Académica (TEXTO FIJO PASCUAL BRAVO)
    doc.add_heading("2.3. Fundamentación académica", level=2)
    doc.add_paragraph("La fundamentación académica del Programa responde a los Lineamientos Académicos y Curriculares (LAC) de la I.U. Pascual Bravo...")
    doc.add_paragraph("Dentro de los LAC se establece la política de créditos académicos...")
    
    doc.add_heading("Rutas educativas: Certificaciones Temáticas Tempranas", level=3)
    doc.add_paragraph("Las Certificaciones Temáticas Tempranas son el resultado del agrupamiento de competencias...")
    
    # Tabla de Certificaciones
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = 'Certificación', 'Cursos', 'Créditos Totales'
    
    for c in cert_data:
        if c["Nombre"]:
                row = table.add_row().cells
                row[0].text = c["Nombre"]
                row[1].text = f"{c['Curso 1']}, {c['Curso 2']}"
                row[2].text = str(c["Créditos 1"] + c["Créditos 2"])
            

# --- LÓGICA DE GENERACIÓN Y GUARDADO ---
if generar:
    # (Toda tu lógica anterior de crear el 'doc'...)
    
    # AL FINAL, DESPUÉS DE GENERAR EL WORD:
    try:
        # 1. Leer datos actuales
        df_actual = conn.read()
        
        # 2. Crear nueva fila con la info del formulario
        nueva_data = pd.DataFrame([{
            "SNIES": snies,
            "Email": email_usuario,
            "Denominacion": denom,
            "Acuerdo": acuerdo,
            "Fecha_Registro": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M")
        }])
        
        # 3. Concatenar y actualizar el Sheet
        df_final = pd.concat([df_actual, nueva_data], ignore_index=True)
        conn.update(data=df_final)
        
        st.info("📊 Información guardada en Google Sheets.")
    except Exception as e:
        st.error(f"No se pudo guardar en el Excel: {e}")

    # (Botón de descarga de Word...)

    
        # Guardar archivo
    bio = io.BytesIO()
    doc.save(bio)
        
    st.success("✅ ¡Documento PEP generado!")
    st.download_button(
        label="📥 Descargar Documento PEP en Word",
        data=bio.getvalue(),
        file_name=f"PEP_Modulo1_{denom.replace(' ', '_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        # --- FINALIZACIÓN Y DESCARGA ---
#bio = io.BytesIO()
#doc.save(bio)
#st.success("✅ ¡Documento PEP generado con éxito!")
#st.download_button(
 #       label="📥 Descargar Documento PEP en Word",
  #      data=bio.getvalue(),
   #     file_name=f"PEP_{denom.replace(' ', '_')}.docx",
    #    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#)























































































