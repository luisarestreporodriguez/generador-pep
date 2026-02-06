import streamlit as st
from google import genai
from docx import Document
from docx.shared import Pt
import requests
import io
import time
import re 
import os
from huggingface_hub import InferenceClient
import pandas as pd
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(page_title="Generador PEP", page_icon="📚", layout="wide")

st.title("Generador PEP - Módulo 1: Información del Programa")

# --- LÓGICA DE API KEY (Nube + Local) ---
# Intentamos leer la clave desde los Secrets de Streamlit
#if "GEMINI_API_KEY" in st.secrets:
 #   api_key = st.secrets["GEMINI_API_KEY"]
  #  else:
# --- LÓGICA DE API KEYS Y SELECTOR (Nube + Local) ---
with st.sidebar:
    st.header("⚙️ Configuración de IA")
    
    # 1. Selector de motor de IA
modelo_ia = st.radio(
        "Selecciona el motor de redacción:",
        ["Google Gemini (Recomendado)", "Hugging Face (Gratuito)"],
        help="Gemini requiere una API Key. Hugging Face usa el token de Secrets o ingreso manual."
    )

    # Inicializamos las variables para que existan en todo el código
api_key = None
hf_token = None

    # 2. Lógica para Gemini
if "Gemini" in modelo_ia:
    if "GEMINI_API_KEY" in st.secrets:
        api_key = st.secrets["GEMINI_API_KEY"]
        st.success("✅ Gemini API Key cargada")
    else:
        api_key = st.text_input("Ingresa tu Google API Key", type="password")
        if not api_key:
           st.warning("⚠️ Introduce la API Key para usar Gemini.")
    
    # 3. Lógica para Hugging Face (Solo un 'else', sin duplicados)
else:
    if "HF_TOKEN" in st.secrets:
          hf_token = st.secrets["HF_TOKEN"]
          st.success("✅ HF Token cargado desde Secrets")
    else:
          hf_token = st.text_input("Ingresa tu HF Token", type="password")
          if not hf_token:
             st.warning("⚠️ Introduce el Token de Hugging Face.")
            
# --- FUNCIÓN DE REDACCIÓN IA ---
def redactar_seccion_ia(titulo_seccion, datos_seccion, llave_api):
    # Ahora usamos 'llave_api' que viene desde el sidebar
    if not llave_api: 
        return "Error: No hay API Key configurada en el sidebar."
    respuestas_reales = {k: v for k, v in datos_seccion.items() if str(v).strip()}
    contexto = "\n".join([f"- {k}: {v}" for k, v in respuestas_reales.items()])
    
    try:
        client = genai.Client(api_key=llave_api)
        prompt = f"""
        Actúa como un Vicerrector Académico experto en aseguramiento de la calidad.
        Tarea: Redactar el motivo de creación del Programa
        DATOS SUMINISTRADOS:{contexto}
        
        REGLAS CRÍTICAS DE FORMATO:
        1. Responde ÚNICAMENTE con UN SOLO PÁRRAFO de texto corrido.
        2. NO incluyas títulos, ni subtítulos (prohibido usar "##" o "Contexto").
        3. NO uses negritas, ni corchetes, ni nombres de la institución entre etiquetas.
        4. Empieza directamente con la redacción (ej: "La pertinencia de este programa se fundamenta...").
        5. El tono debe ser muy formal, académico y fluido. Máximo 150 palabras.
       """
        response = client.models.generate_content(model="gemini-flash-latest", contents=prompt)
        # Limpiar posibles espacios en blanco extras o saltos de línea al inicio/final
        return response.text
    except Exception as e:
        return f"Error en redacción: {str(e)}"

# --- CONFIGURACIÓN HUGGING FACE (Alternativa Gratuita) ---
def redactar_seccion_ia_hf(titulo_seccion, datos_seccion, hf_token):
    """Función alternativa usando modelos gratuitos de Hugging Face"""
    if not hf_token:
        return "Error: No hay Token de Hugging Face configurado"

 # Usamos Zephyr directamente aquí para evitar confusiones
    client = InferenceClient(api_key=hf_token)
    respuestas_reales = {k: v for k, v in datos_seccion.items() if str(v).strip()}
    contexto = "\n".join([f"- {k}: {v}" for k, v in respuestas_reales.items()])
    if not respuestas_reales:
        return f"No hay información suficiente para redactar la sección {titulo_seccion}."

   
    try:
        # Usamos el modelo Qwen 2.5
        completion = client.chat.completions.create(
            model="HuggingFaceH4/zephyr-7b-beta",
            messages=[
                {
                    "role": "system",
                    "content": "Eres un Vicerrector Académico experto. Redacta párrafos formales, académicos y fluidos. No uses negritas ni títulos."
                },
                {
                    "role": "user",
                    "content": f"Redacta un párrafo para la sección '{titulo_seccion}' con esta información:\n{contexto}"
                }
            ],
            max_tokens=400,
            temperature=0.5
        )
        # Extraemos el texto de la respuesta
        return completion.choices[0].message.content.strip()

    except Exception as e:
        if "503" in str(e) or "loading" in str(e).lower():
            return "⏳ El modelo está cargando en el servidor. Reintenta en 15 segundos."
        return f"Error con la IA: {str(e)}"
     
# --- ESTRUCTURA DE CONTENIDOS ---
estructura_pep = {
    "1. Información del Programa": {
        "1.1. Historia del Programa": {"tipo": "especial_historia"},
        "1.2. Generalidades del Programa": {"tipo": "directo"}
    },
    "2. Referentes Conceptuales": {
    "2.1. Naturaleza del Programa": {
        "tipo": "ia",
        "campos": [
            {"label": "Objeto de conocimiento del Programa", "req": True, "help": "¿Qué conoce, investiga y transforma este programa?"}
        ]
    },
    "2.2. Fundamentación epistemológica": {
        "tipo": "ia",
        "campos": [
            {"label": "Naturaleza epistemológica e identidad académica", "req": True},
            {"label": "Campo del saber y relación con ciencia/tecnología", "req": True}
        ]
    },
    "2.3. Fundamentación académica": {
        "tipo": "especial_pascual", # Nueva lógica para textos fijos + tabla
        "campos": [] 
    }
}

}


# --- BOTÓN DE DATOS DE EJEMPLO ---
# Usamos session_state para persistir los datos al hacer clic
if st.button("📎 Llenar con datos de ejemplo"):
    st.session_state.ejemplo = {
        "denom": "Ingeniería de Sistemas",
        "titulo": "Ingeniero de Sistemas",
        "nivel_idx": 2, # Profesional universitario
        "area": "Ingeniería, Arquitectura y Urbanismo",
        "modalidad_idx": 4, # Presencial y Virtual
        "acuerdo": "Acuerdo 012 de 2015",
        "instancia": "Consejo Académico",
        "reg1": "Res. 4567 de 2016",
        "reg2": "Res. 8901 de 2023",
        "acred1": "Res. 00234 de 2024",
        "creditos": "165",
        "periodo_idx": 0, # Semestral
        "lugar": "Sede Principal (Cali)",
        "snies": "54321",
        "motivo": "La creación del Programa se fundamenta en la necesidad de formar profesionales capaces de liderar la transformación digital, diseñar y desarrollar soluciones de software de alta complejidad, gestionar sistemas de información y responder de manera innovadora a los retos tecnológicos, organizacionales y sociales del entorno local, nacional e internacional.",
        "p1_nom": "EO1", "p1_fec": "Acuerdo 012-2015",
        "p2_nom": "EO2", "p2_fec": "Acuerdo 088-2020",
        "p3_nom": "EO3", "p3_fec": "Acuerdo 102-2024",
        #DATOS CAPÍTULO 2
        "objeto_nombre": "Sistemas de información",
        "objeto_concep": "Los sistemas de información son conjuntos organizados de personas, datos, procesos, tecnologías y recursos que interactúan de manera integrada para capturar, almacenar, procesar, analizar y distribuir información, con el fin de apoyar la toma de decisiones, la gestión operativa, el control organizacional y la generación de conocimiento. Estos sistemas permiten transformar los datos en información útil y oportuna, facilitando la eficiencia, la innovación y la competitividad en organizaciones de distintos sectores. Su diseño y gestión consideran aspectos técnicos, organizacionales y humanos, garantizando la calidad, seguridad, disponibilidad y uso ético de la información.",        
        "fund_epi": "El programa se inscribe en el racionalismo crítico y el pragmatismo tecnológico, vinculando la ciencia de la computación con la ingeniería aplicada.",
        # DATOS PARA LAS TABLAS (Se guardan como listas de diccionarios)
        "recon_data": [
            {"Año": "2024", "Nombre del premio": "Excelencia Académica", "Nombre del Ganador": "Juan Pérez", "Cargo": "Docente"}
        ],
        "tabla_cert_ej": [
            {"Nombre": "Desarrollador Web Junior", "Curso 1": "Programación I", "Créditos 1": 3, "Curso 2": "Bases de Datos", "Créditos 2": 4},
            {"Nombre": "Analista de Datos", "Curso 1": "Estadística", "Créditos 1": 4, "Curso 2": "Python para Ciencia", "Créditos 2": 4}
        ], # <--- ESTE CORCHETE Y COMA FALTABAN PARA CERRAR LA LISTA ANTERIOR
        
        "referencias_data": [
            {
                "Año": "2021", 
                "Autor(es)": "Sommerville, I.", 
                "Revista": "Computer science", 
                "Título del artículo/Libro": "Engineering Software Products"
            },
            {
                "Año": "2023", 
                "Autor(es)": "Pressman, R. & Maxim, B.", 
                "Revista": "Software Engineering Journal", 
                "Título del artículo/Libro": "A Practitioner's Approach"
            }
        ],
    }
    st.rerun()

# --- FORMULARIO DE ENTRADA ---
with st.form("pep_form"):
    ej = st.session_state.get("ejemplo", {})

    st.markdown("### 📋 1. Identificación General")
    col1, col2 = st.columns(2)
    with col1:
        denom = st.text_input("Denominación del programa :red[•] ", value=ej.get("denom", ""))
        titulo = st.text_input("Título otorgado :red[•]", value=ej.get("titulo", ""))
        nivel = st.selectbox("Nivel de formación :red[•]", 
                             ["Técnico", "Tecnológico", "Profesional universitario", "Especialización", "Maestría", "Doctorado"], 
                             index=ej.get("nivel_idx", 2))
        area = st.text_input("Área de formación :red[•]", value=ej.get("area", ""))
    
    with col2:
        modalidad = st.selectbox("Modalidad de oferta :red[•]", 
                                 ["Presencial", "Virtual", "A Distancia", "Dual", "Presencial y Virtual", "Presencial y a Distancia", "Presencial y Dual"],
                                 index=ej.get("modalidad_idx", 0))
        acuerdo = st.text_input("Acuerdo de creación / Norma interna :red[•]", value=ej.get("acuerdo", ""))
        instancia = st.text_input("Instancia interna que aprueba :red[•]", value=ej.get("instancia", ""))
        snies = st.text_input("Código SNIES :red[•]", value=ej.get("snies", ""))

    st.markdown("---")
    st.markdown("### 📄 2. Registros y Acreditaciones")
    col3, col4 = st.columns(2)
    with col3:
        reg1 = st.text_input("Resolución Registro calificado 1 :red[•]", value=ej.get("reg1", ""), placeholder="Número y año")
        reg2 = st.text_input("Registro calificado 2 (Opcional)", value=ej.get("reg2", ""))
        acred1 = st.text_input("Resolución Acreditación en alta calidad 1 (Opcional)", value=ej.get("acred1", ""))
        acred2 = st.text_input("Resolución Acreditación en alta calidad 2 (Opcional)", value="")

    with col4:
        creditos = st.text_input("Créditos académicos :red[•]", value=ej.get("creditos", ""))
        periodicidad = st.selectbox("Periodicidad de admisión :red[•]", ["Semestral", "Anual"], index=ej.get("periodo_idx", 0))
        lugares = st.text_input("Lugares de desarrollo :red[•]", value=ej.get("lugar", ""))

    frase_auto = f"La creación del Programa {denom} se fundamenta en la necesidad de "
    val_motivo = ej.get("motivo", frase_auto)
    motivo = st.text_area("Motivo de creación :red[•]", value=val_motivo, height=150)
      
    st.markdown("---")
    st.markdown("### 📚 3. Modificaciones al Plan de Estudios")
    p_col1, p_col2, p_col3 = st.columns(3)
    with p_col1:
        p1_nom = st.text_input("Nombre Plan v1:red[•]", value=ej.get("p1_nom", ""))
        p1_fec = st.text_input("Acuerdo aprobación Plan v1 :red[•]", value=ej.get("p1_fec", ""))
    with p_col2:
        p2_nom = st.text_input("Nombre Plan v2 (Opcional)", value=ej.get("p2_nom", ""))
        p2_fec = st.text_input("Acuerdo aprobación Plan v2 (Opcional)", value=ej.get("p2_fec", ""))
    with p_col3:
        p3_nom = st.text_input("Nombre Plan v3 (Opcional)", value=ej.get("p3_nom", ""))
        p3_fec = st.text_input("Acuerdo aprobación Plan v3 (Opcional)", value=ej.get("p3_fec", ""))

    st.markdown("---")
    st.markdown("### 🏆 4. Reconocimientos (Opcional)")
    recon_data = st.data_editor(
        ej.get("recon_data", [{"Año": "", "Nombre del premio": "", "Nombre del Ganador": "", "Cargo": "Estudiante"}]),
        num_rows="dynamic",
        key="editor_recon", # Es vital tener una key única
        column_config={
            "Cargo": st.column_config.SelectboxColumn(options=["Docente", "Líder", "Decano", "Estudiante,Docente Investigador, Investigador"])
        },
        use_container_width=True
        )
# --- CAPÍTULO 2 ---
    st.markdown("---")
    st.header("2. Referentes Conceptuales")
   # 2. Objeto de conocimiento del Programa
    val_obj_nombre = ej.get("objeto_nombre", "")
    objeto_nombre = st.text_input(
    "1. ¿Cuál es el Objeto de conocimiento del Programa? :red[•]",
         value=val_obj_nombre, 
         placeholder="Ejemplo: Sistemas de información",
         key="obj_nombre_input"
)

  # 2.1 Conceptualización 
    val_obj_concep = ej.get("objeto_concep", "")
    objeto_conceptualizacion = st.text_area(
        "2. Conceptualización del objeto de conocimiento del Programa :red[•]",
        value=val_obj_concep, 
        height=150, 
        key="obj_concep_input", 
        placeholder="Ejemplo: Los sistemas de información son conjuntos organizados de personas, datos, procesos, tecnologías y recursos que interactúan de manera integrada para capturar, almacenar, procesar, analizar y distribuir información, con el fin de apoyar la toma de decisiones, la gestión operativa, el control organizacional y la generación de conocimiento. Estos sistemas permiten transformar los datos en información útil y oportuna, facilitando la eficiencia, la innovación y la competitividad en organizaciones de distintos sectores. Su diseño y gestión consideran aspectos técnicos, organizacionales y humanos, garantizando la calidad, seguridad, disponibilidad y uso ético de la información."
    )
 #2.2 
    st.write("Referencias bibliográficas que sustentan la conceptualización del Objeto de Conocimiento.")
    referencias_previa = ej.get("referencias_data", [
        {"Año": "", "Autor(es) separados por coma": "", "Revista": "", "Título del artículo/Libro": ""}
    ])

    referencias_data = st.data_editor(
        referencias_previa,
        num_rows="dynamic", # Permite al usuario agregar/borrar filas con el signo +
        key="editor_referencias",
        use_container_width=True,
        column_config={
            "Año": st.column_config.TextColumn("Año", width="small"),
            "Autor(es)": st.column_config.TextColumn("Autor(es)", width="medium"),
            "Revista": st.column_config.TextColumn("Revista", width="medium"),
            "Título del artículo/Libro": st.column_config.TextColumn("Título del artículo/Libro", width="large"),
        }
    )

  # 2.2. Fundamentación epistemológica en Pestañas ---
    st.markdown("---")
    st.subheader("2.2. Fundamentación epistemológica")
    st.info("Utilice las pestañas para completar los tres párrafos de la Fundamentación epistemológica.")

# 1. Creamos las pestañas
    tab1, tab2, tab3 = st.tabs(["Párrafo 1", "Párrafo 2", "Párrafo 3"])

# Configuración de columnas 
    config_columnas_ref = {
        "Año": st.column_config.TextColumn("Año", width="small"),
        "Autor(es) separados por coma": st.column_config.TextColumn("Autor(es)", width="medium"),
        "Revista": st.column_config.TextColumn("Revista", width="medium"),
        "Título del artículo/Libro": st.column_config.TextColumn("Título del artículo/Libro", width="large"),
    }

# Bloque Párrafo 1
    with tab1:
        st.markdown("### Párrafo 1: Marco filósofico")
        st.text_area(
            "¿Cuál es la postura filosófica predominante (positivismo, constructivismo, teoría crítica, complejidad)?:red[•]",
            value=ej.get("fund_epi_p1", ""), 
            height=200,
            key="input_epi_p1",
            placeholder="""Ejemplo: El programa se fundamenta en el paradigma de la complejidad y la visión sistémica, donde la realidad no se percibe como un conjunto de elementos aislados, sino como una red de interacciones y procesos emergentes. Bajo esta postura, el conocimiento se valida a través de la capacidad de modelar y abstraer sistemas reales para transformarlos en sistemas de información lógico-formales que sean verificables y funcionales. Así, la "verdad" en esta disciplina no reside únicamente en el componente físico (el hardware) o en el código (el software), sino en la coherencia del flujo de información y en la eficacia del sistema para resolver problemas de organización, entropía y control en entornos dinámicos y globales..""",
        )
        st.write("Referencias bibliográficas (Párrafo 1):")
        st.data_editor(
            ej.get("referencias_epi_p1", [{"Año": "", "Autor(es) separados por coma": "", "Revista": "", "Título del artículo/Libro": ""}]),
            num_rows="dynamic",
            key="editor_refs_p1",
            use_container_width=True,
            column_config=config_columnas_ref
        )

# Bloque Párrafo 2
    with tab2:
        st.markdown("### Párrafo 2: Identidad disciplinar")
        st.text_area(
            "Origen etimológico y teórias conceptuales que sustentan el Programa:red[•]",
            value=ej.get("fund_epi_p2", ""), 
            height=200,
            key="input_epi_p2",
            placeholder="""Ejemplo: La identidad de este programa se define desde la convergencia etimológica de la ingeniería —del latín ingenium, que refiere a la capacidad natural de invención y resolución de problemas— y el concepto de sistema —del griego systema, entendido como la unión de partes que forman un todo organizado—. Esta génesis conceptual establece que su objeto de estudio no es la máquina en sí misma, sino la arquitectura de procesos y la gestión de la complejidad mediante el uso de la tecnología. Sustentado en la Teoría General de Sistemas y la Cibernética, el programa se deslinda de las ingenierías tradicionales al centrarse en lo intangible —la información y la estructura—, permitiendo que el profesional no solo diseñe herramientas digitales, sino que sea capaz de integrar elementos humanos, tecnológicos y organizacionales en soluciones holísticas y escalables.""",
       )
        st.write("Referencias bibliográficas (Párrafo 2):")
        st.data_editor(
            ej.get("referencias_epi_p2", [{"Año": "", "Autor(es) separados por coma": "", "Revista": "", "Título del artículo/Libro": ""}]),
            num_rows="dynamic",
            key="editor_refs_p2",
            use_container_width=True,
            column_config=config_columnas_ref
        )

# Bloque Párrafo 3
    with tab3:
        st.markdown("### Párrafo 3: Intencionalidad social")
        st.text_area(
            "¿De qué manera la forma en que se produce el conocimiento en este programa garantiza una intervención ética y transformadora en el entorno profesional?:red[•]",
            value=ej.get("fund_epi_p3", ""), 
            height=200,
            key="input_epi_p3",
            placeholder="""Ejemplo: Finalmente, la producción de conocimiento en este programa se orienta hacia una praxis ética y socialmente responsable, donde la tecnología se entiende como un medio para el desarrollo humano y no como un fin deshumanizante. La intervención del ingeniero de sistemas trasciende la ejecución técnica para convertirse en una labor de transformación digital con conciencia crítica, garantizando la seguridad, la privacidad y la integridad de los datos en una sociedad cada vez más automatizada. Este compromiso teleológico asegura que el profesional no solo responda a las demandas del mercado, sino que actúe como un gestor del cambio capaz de diseñar soluciones sostenibles que reduzcan las brechas tecnológicas y promuevan la eficiencia organizacional bajo principios de transparencia y justicia social.""",

        )
        st.write("Referencias bibliográficas (Párrafo 3):")
        st.data_editor(
           ej.get("referencias_epi_p3", [{"Año": "", "Autor(es) separados por coma": "", "Revista": "", "Título del artículo/Libro": ""}]),
            num_rows="dynamic",
            key="editor_refs_p3",
            use_container_width=True,
            column_config=config_columnas_ref
        )

    # 2.3. Fundamentación académica 
    st.subheader("Certificaciones Temáticas Tempranas")
    cert_data = st.data_editor(
        ej.get("tabla_cert_ej", [{"Nombre": "", "Curso 1": "", "Créditos 1": 0, "Curso 2": "", "Créditos 2": 0}]),
        num_rows="dynamic",      
        key="editor_cert"
    )

    generar = st.form_submit_button("🚀 GENERAR DOCUMENTO PEP", type="primary")

# --- LÓGICA DE GENERACIÓN DEL WORD ---
if generar:
    if not denom or not reg1:
        st.error("⚠️ Falta información obligatoria (Denominación o Registro Calificado).")
    else:
        doc = Document()
        # Estilo base
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
            # 1.1 Historia del Programa
        doc.add_heading("1.1. Historia del Programa", level=1)
        
        # PÁRRAFO 1. Datos creación
        texto_historia = (
            f"El Programa de {denom} fue creado mediante el {acuerdo} del {instancia} "
            f"y aprobado mediante la resolución de Registro Calificado {reg1} del Ministerio de Educación Nacional "
            f"con código SNIES {snies}."
        )
        doc.add_paragraph(texto_historia)
        
        # PÁRRAFO 2. Motivo de creación
        if motivo.strip():
    # El usuario ya escribió empezando con "La creación del programa..."
           doc.add_paragraph(motivo) 
        else:
            doc.add_paragraph("No se suministró información sobre el motivo de creación.")
     
        # PÁRRAFO 3. Acreditación 1 y/o 2
        if acred1 and not acred2:
    # Caso: Solo una acreditación
            texto_acred = (
            f"El programa obtuvo la Acreditación en alta calidad otorgada por el "
            f"Consejo Nacional de Acreditación (CNA) a través de la resolución {acred1}, "
            f"como reconocimiento a su solidez académica, administrativa y de impacto social."
        )
            doc.add_paragraph(texto_acred)

        elif acred1 and acred2:
    # Caso: Dos acreditaciones (Primera vez + Renovación)
            texto_acred = (
            f"El programa obtuvo por primera vez la Acreditación en alta calidad otorgada por el "
            f"Consejo Nacional de Acreditación (CNA) a través de la resolución {acred1}, "
            f"esta le fue renovada mediante resolución {acred2}, reafirmando la solidez "
            f"académica, administrativa y de impacto social del Programa."
        )
            doc.add_paragraph(texto_acred)    

        # PÁRRAFO 4: Modificaciones curriculares
        planes_nom = [n for n in [p1_nom, p2_nom, p3_nom] if n]
        planes_fec_lista = [f for f in [p1_fec, p2_fec, p3_fec] if f]
        
        if planes_fec_lista and planes_nom:
            # A. Formatear nombres de planes (lo que antes era "lista")
            if len(planes_nom) > 1:
                txt_planes_lista = ", ".join(planes_nom[:-1]) + f" y {planes_nom[-1]}"
            else:
                txt_planes_lista = planes_nom[0]

            # B. Formatear fechas/acuerdos
            if len(planes_fec_lista) > 1:
                txt_acuerdos_formateado = ", ".join(planes_fec_lista[:-1]) + f" y {planes_fec_lista[-1]}"
            else:
                txt_acuerdos_formateado = planes_fec_lista[0]

            texto_planes = (
                 f"El plan de estudios del Programa de {denom} ha sido objeto de procesos periódicos de evaluación, "
                 f"con el fin de asegurar su pertinencia académica y su alineación con los avances tecnológicos "
                 f"y las demandas del entorno. Como resultado, se han realizado las modificaciones curriculares "
                 f"{txt_planes_lista}, aprobadas mediante el {txt_acuerdos_formateado}, respectivamente."
            )
            p_planes = doc.add_paragraph(texto_planes)
            p_planes.alignment = 3  # Justificado
    
        # PÁRRAFO 5: Reconocimientos
        recons_validos = [r for r in recon_data if r.get("Nombre del premio", "").strip()]
        
        if recons_validos:
             doc.add_paragraph(
                 f"El Programa de {denom} ha alcanzado importantes logros académicos e institucionales "
                 f"que evidencian su calidad y compromiso con la excelencia. Entre ellos se destacan:"
             )
             for r in recons_validos:
                 premio = r.get("Nombre del premio", "N/A")
                 anio = r.get("Año", "N/A")
                 ganador = r.get("Nombre del Ganador", "N/A")
                 cargo = r.get("Cargo", "N/A")
                 doc.add_paragraph(
             f" {premio} ({anio}): Otorgado a {ganador}, en su calidad de {cargo}.", 
             style='List Bullet')

        # Línea de tiempo
        doc.add_heading("Línea de Tiempo del Programa", level=2)
    # Función interna para extraer solo el año (4 dígitos)
        def extraer_anio(texto):
             if not texto: return "N/A"
             match = re.search(r'20\d{2}', str(texto)) # Busca "20" seguido de dos números
             return match.group(0) if match else str(texto).split()[-1]
            
    # 1. Creación (Usando el año del primer plan o acuerdo)
        if p1_fec:
             anio = extraer_anio(p1_fec)
             doc.add_paragraph(f"{anio}: Creación del Programa")
             p_obj.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # <--- JUSTIFICADO


    # 2. Registros Calificados
        if reg1:
                    # Intenta extraer el año (asumiendo formato "Res XXX de 20XX")
             anio_reg1 = reg1.split()[-1] if len(reg1.split()) > 0 else "Fecha N/A"
             doc.add_paragraph(f"{anio_reg1}: Obtención del Registro Calificado inicial")
             p_obj.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # <--- JUSTIFICADO

        if reg2:
             anio_reg2 = reg2.split()[-1] if len(reg2.split()) > 0 else "Fecha N/A"
             doc.add_paragraph(f"{anio_reg2}: Renovación del Registro Calificado")
             p_obj.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # <--- JUSTIFICADO


    # 3. Modificaciones Curriculares (Planes de estudio)
        if p2_fec:
              anio = extraer_anio(p2_fec)
              doc.add_paragraph(f"{anio}: Modificación curricular 1 (Actualización del plan de estudios)")
              p_obj.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # <--- JUSTIFICADO

        
        if p3_fec:
              anio = extraer_anio(p3_fec)
              doc.add_paragraph(f"{anio}: Modificación curricular 2")
              p_obj.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # <--- JUSTIFICADO


    # 4. Acreditaciones de Alta Calidad
        if acred1:
              anio_acred1 = acred1.split()[-1] if len(acred1.split()) > 0 else "Fecha N/A"
              doc.add_paragraph(f"{anio_acred1}: Obtención de la Acreditación en Alta Calidad")
              p_obj.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # <--- JUSTIFICADO

        
        if acred2:
              anio_acred2 = acred2.split()[-1] if len(acred2.split()) > 0 else "Fecha N/A"
              doc.add_paragraph(f"{anio_acred2}: Renovación de la Acreditación en Alta Calidad")
              p_obj.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # <--- JUSTIFICADO


        # 5. Reconocimientos (Si existen en la tabla)
        if recons_validos:
                    # Tomamos los años únicos de los reconocimientos para no repetir
             anios_recon = sorted(list(set([r['Año'] for r in recons_validos if r['Año']])))
             for a in anios_recon:
                 doc.add_paragraph(f"{a}: Reconocimientos académicos destacados")
                
        # 1.2 GENERALIDADES (Tabla de datos)
        doc.add_page_break() 
        doc.add_heading("1.2 Generalidades del Programa", level=1)
        
        items_gen = [
                            ("Denominación del programa", denom),
                            ("Título otorgado", titulo),
                           ("Nivel de formación", nivel),
                            ("Área de formación", area),
                            ("Modalidad de oferta", modalidad),
                            ("Acuerdo de creación", acuerdo),
                            ("Registro calificado", reg1),
                            ("Créditos académicos", creditos),
                            ("Periodicidad de admisión", periodicidad),
                            ("Lugares de desarrollo", lugares),
                            ("Código SNIES", snies)
                        ]
        
        for k, v in items_gen:
                            p = doc.add_paragraph()
                            p.add_run(f"{k}: ").bold = True
                            p.add_run(str(v))

  # 2.1 Referentes conceptuales 
        doc.add_heading("2.1. Referentes conceptuales", level=2)

        obj_nom = st.session_state.get("obj_nombre_input", "No definido")
        obj_con = st.session_state.get("obj_concep_input", "")

        # Bloque: Objeto + Enter + Conceptualización
        p_obj = doc.add_paragraph()
        p_obj.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # <--- JUSTIFICADO
        p_obj.add_run("Objeto de conocimiento del Programa: ").bold = True
        p_obj.add_run(str(obj_nom)) # Forzamos a string para evitar errores

        p_concep = doc.add_paragraph(obj_con)
        p_concep.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # <--- JUSTIFICADO
        
        # 3. Referencias de la tabla
       # --- EXTRACCIÓN ROBUSTA DE REFERENCIAS ---
        raw_concep = st.session_state.get("editor_referencias", [])
        
        citas_c = []
        
        # Caso 1: Los datos vienen en un diccionario (Común en st.form)
        if isinstance(raw_concep, dict):
            # Intentamos obtener la lista de 'data' o los valores de 'edited_rows'
            datos_lista = raw_concep.get("data", list(raw_concep.get("edited_rows", {}).values()))
        elif isinstance(raw_concep, list):
            datos_lista = raw_concep
        else:
            datos_lista = []
        
        for fila in datos_lista:
            # Verificamos que 'fila' sea realmente un diccionario antes de usar .get()
            if isinstance(fila, dict):
                aut = ""
                ani = ""
                # Buscamos en las llaves del diccionario de forma flexible
                for k, v in fila.items():
                    k_low = str(k).lower()
                    if "autor" in k_low: aut = str(v).strip()
                    if "año" in k_low or "anio" in k_low: ani = str(v).strip()
                
                if aut and ani and aut.lower() != "none" and aut != "":
                    citas_c.append(f"{aut}, {ani}")
        
        if citas_c:
            p_concep.add_run(" (Sustentado en: " + "; ".join(citas_c) + ").")
   
        # --- 2.2 FUNDAMENTACIÓN EPISTEMOLÓGICA ---
        doc.add_heading("2.2. Fundamentación epistemológica", level=2)
        
        # Iteramos los 3 bloques de las pestañas
        for i in range(1, 4):
            texto_p = st.session_state.get(f"input_epi_p{i}", "")
            if texto_p:
                p_f = doc.add_paragraph(texto_p)
                p_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                raw_f = st.session_state.get(f"editor_refs_p{i}", [])
                
                # Normalizar datos de la tabla de la pestaña
                if isinstance(raw_f, dict):
                    datos_f = raw_f.get("data", list(raw_f.get("edited_rows", {}).values()))
                else:
                    datos_f = raw_f
                    
                citas_p = []
                for f in datos_f:
                    if isinstance(f, dict):
                        a_f, n_f = "", ""
                        for k, v in f.items():
                            k_l = str(k).lower()
                            if "autor" in k_l: a_f = str(v).strip()
                            if "año" in k_l or "anio" in k_l: n_f = str(v).strip()
                        if a_f and n_f and a_f.lower() != "none" and a_f != "":
                            citas_p.append(f"{a_f}, {n_f}")
                
                if citas_p:
                    p_f.add_run(" (Ref: " + "; ".join(citas_p) + ").")
                
    # 2.2 Epistemología
    #    doc.add_heading("2.2. Fundamentación epistemológica", level=2)
     #   doc.add_paragraph(redactar_seccion_ia("Fundamentación Epistemológica", {"Datos": fund_epi}))

    # 2.3 Fundamentación Académica (TEXTO FIJO PASCUAL BRAVO)
        #doc.add_heading("2.3. Fundamentación académica", level=2)
        #doc.add_paragraph("La fundamentación académica del Programa responde a los Lineamientos Académicos y Curriculares (LAC) de la I.U. Pascual Bravo...")
        #doc.add_paragraph("Dentro de los LAC se establece la política de créditos académicos...")
   
        #doc.add_heading("Rutas educativas: Certificaciones Temáticas Tempranas", level=3)
        #doc.add_paragraph("Las Certificaciones Temáticas Tempranas son el resultado del agrupamiento de competencias...")
    
    # Tabla de Certificaciones
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = 'Certificación', 'Cursos', 'Créditos Totales'
    
        for c in cert_data:
         if c["Nombre"]:
                              row = table.add_row().cells
                              row[0].text = c["Nombre"]
                              row[1].text = f"{c['Curso 1']}, {c['Curso 2']}"
                              row[2].text = str(c["Créditos 1"] + c["Créditos 2"])
            
        # Guardar archivo
    bio = io.BytesIO()
    doc.save(bio)
        
    st.success("✅ ¡Documento PEP generado!")
    st.download_button(
        label="📥 Descargar Documento PEP en Word",
        data=bio.getvalue(),
        file_name=f"PEP_Modulo1_{denom.replace(' ', '_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
   )
